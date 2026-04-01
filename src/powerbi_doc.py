# -*- coding: utf-8 -*-
"""
Gera documentação técnica de relatórios Power BI (`.pbix`).

Fluxo principal:
- lê um único arquivo `.pbix` por execução;
- inspeciona a estrutura interna do pacote;
- tenta extrair layout, páginas, visuais, filtros, bookmarks, temas e metadados textuais;
- tenta localizar metadados de modelo acessíveis no pacote, quando disponíveis;
- gera artefatos de documentação em Markdown, RTF, DOCX, JSON e/ou Excel.

Observação importante:
- arquivos `.pbix` não oferecem o mesmo nível de transparência estrutural de um projeto `.pbip`;
- por isso, a extração do modelo semântico em `.pbix` é feita em regime de melhor esforço.
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import sys
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path
from typing import Any

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import olefile
except ImportError:
    olefile = None

try:
    from docx import Document
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:
    Document = None
    WD_TAB_ALIGNMENT = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "data"
TEMPORARY_OUTPUT_NAMES = {
    "package_contents",
    ".tmp",
    "tmp",
    "temp",
}

RTF_BODY_FONT_NAME = "Arial"
RTF_MONO_FONT_NAME = "Courier New"
DOCX_BODY_FONT_NAME = RTF_BODY_FONT_NAME
DOCX_MONO_FONT_NAME = RTF_MONO_FONT_NAME
ENABLE_PROGRESS_PRINTS = True

TEXTUAL_PACKAGE_MEMBERS = {
    "version",
    "metadata",
    "settings",
    "connections",
    "diagramlayout",
    "diagramstate",
    "linguisticschema",
    "report/layout",
    "report/metadata",
}


def sanitize_filename(value: str) -> str:
    safe = re.sub(r"[^\w\-. ]+", "_", value, flags=re.UNICODE).strip()
    return safe or "arquivo"


def unique_ordered(values: list[Any]) -> list[Any]:
    output = []
    seen = set()
    for value in values:
        marker = json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list)) else value
        if marker in seen:
            continue
        seen.add(marker)
        output.append(value)
    return output


def compact_json(value: Any) -> str:
    if value is None:
        return "-"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def normalize_whitespace(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = value.replace("\r\n", "\n").replace("\r", "\n").strip()
    return normalized or None


def clean_display_label(value: str | None) -> str | None:
    if value is None:
        return None
    text = value.replace("\xa0", " ").strip()
    text = re.sub(r"\s{2,}", " ", text)
    return text or None


def _log_progress(message: str) -> None:
    if ENABLE_PROGRESS_PRINTS:
        print(f"[PowerBIDoc][progresso] {message}")


class PowerBIDoc:
    """Carrega, extrai e documenta um relatório Power BI."""

    def __init__(self, source_path: str | Path, output_format: str = "all") -> None:
        self.source_path = Path(source_path).expanduser().resolve()
        if not self.source_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {self.source_path}")
        if self.source_path.suffix.lower() != ".pbix":
            raise ValueError("O arquivo informado deve ter extensão .pbix.")

        self.output_format = output_format.lower()
        self.base_name = self.source_path.stem
        self.output_dir = DEFAULT_OUTPUT_ROOT / self.base_name
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.generated_at = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.package_manifest: list[dict[str, Any]] = []
        self.text_payloads: dict[str, str] = {}
        self.package_member_bytes: dict[str, bytes] = {}
        self.layout_payload: dict[str, Any] | None = None
        _log_progress(
            f"Iniciando documentação de `{self.source_path.name}` com formato `{self.output_format}`."
        )
        _log_progress(f"Diretório de saída preparado em `{self.output_dir}`.")

        self._load_package_contents()
        self.metadata = self._build_metadata()

    def _load_package_contents(self) -> None:
        """Lê o pacote `.pbix`, registra seus membros e extrai arquivos úteis."""
        _log_progress(f"Lendo arquivo de origem `{self.source_path}`.")
        try:
            with zipfile.ZipFile(self.source_path, "r") as archive:
                extracted_dir = self.output_dir / "package_contents"
                if extracted_dir.exists():
                    shutil.rmtree(extracted_dir)
                extracted_dir.mkdir(parents=True, exist_ok=True)

                for info in archive.infolist():
                    member_name = info.filename
                    self.package_manifest.append(
                        {
                            "path": member_name,
                            "kind": self._classify_package_member(member_name),
                            "size_bytes": info.file_size,
                        }
                    )

                    if info.is_dir():
                        continue

                    lower = member_name.lower()
                    should_capture = (
                        lower in TEXTUAL_PACKAGE_MEMBERS
                        or lower.endswith(".json")
                        or lower.endswith(".xml")
                        or lower.endswith(".txt")
                    )
                    if not should_capture and info.file_size > 5_000_000:
                        continue

                    raw_bytes = archive.read(member_name)
                    self.package_member_bytes[member_name] = raw_bytes
                    destination = extracted_dir / member_name
                    destination.parent.mkdir(parents=True, exist_ok=True)

                    if self._is_binary_package_member(member_name, raw_bytes):
                        destination.write_bytes(raw_bytes)
                        continue

                    decoded = self._decode_package_text(raw_bytes)
                    if decoded is not None:
                        self.text_payloads[member_name] = decoded
                        destination.write_text(decoded, encoding="utf-8")
                    else:
                        destination.write_bytes(raw_bytes)
        except zipfile.BadZipFile as exc:
            raise ValueError(
                "O arquivo `.pbix` não pôde ser aberto como pacote ZIP legível. "
                "Alguns arquivos PBIX podem exigir ferramentas adicionais ou outra estratégia de leitura."
            ) from exc

        self.layout_payload = self._load_layout_payload()

    def _classify_package_member(self, package_path: str) -> str:
        lower = package_path.lower()
        if lower == "report/layout":
            return "report_layout"
        if lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg", ".webp")):
            return "image"
        if lower in {"datamodel", "datamashup"}:
            return "binary_model_or_mashup"
        if lower.startswith("report/"):
            return "report_asset"
        if lower.endswith((".json", ".txt", ".xml", ".m", ".dax")):
            return "textual"
        return "other"

    def _is_binary_package_member(self, member_name: str, raw_bytes: bytes) -> bool:
        lower = member_name.lower()
        if lower in {"datamodel", "datamashup"}:
            return True
        if raw_bytes[:4] == b"\x00\x00\x00\x00":
            return True
        return False

    def _decode_package_text(self, raw_bytes: bytes) -> str | None:
        for encoding in ("utf-8-sig", "utf-16-le", "utf-16", "latin-1"):
            try:
                text = raw_bytes.decode(encoding)
                if "\x00" in text and encoding not in {"utf-16-le", "utf-16"}:
                    continue
                return text
            except UnicodeDecodeError:
                continue
        return None

    def _safe_json_loads(self, payload: str | None) -> Any:
        if not payload:
            return None
        try:
            return json.loads(payload)
        except json.JSONDecodeError:
            return None

    def _load_layout_payload(self) -> dict[str, Any] | None:
        for member_name in ("Report/Layout", "report/layout"):
            payload = self.text_payloads.get(member_name)
            data = self._safe_json_loads(payload)
            if isinstance(data, dict):
                _log_progress("Layout principal do relatório Power BI carregado.")
                return data
        return None

    def _build_metadata(self) -> dict[str, Any]:
        """Consolida o relatório Power BI em uma estrutura única para saída."""
        _log_progress("Iniciando extração de metadados do relatório Power BI.")
        report_info = {
            "source_file_name": self.source_path.name,
            "source_file_path": str(self.source_path),
            "source_file_type": self.source_path.suffix.lower(),
            "source_file_last_modified": datetime.fromtimestamp(
                self.source_path.stat().st_mtime
            ).isoformat(timespec="seconds"),
            "output_directory": str(self.output_dir),
            "version": normalize_whitespace(self.text_payloads.get("Version") or self.text_payloads.get("version")),
            "has_layout": self.layout_payload is not None,
        }

        pages = self._extract_pages()
        visuals = self._extract_visuals(pages)
        bookmarks = self._extract_bookmarks()
        report_filters = self._extract_report_filters()
        themes = self._extract_themes()
        model = self._extract_model_metadata(visuals)
        queries = self._extract_query_artifacts()
        custom_visuals = self._extract_custom_visuals(visuals)
        package_settings = self._extract_package_settings()
        package_metadata = self._extract_package_metadata()
        security = self._extract_security_bindings()
        content_types = self._extract_content_types()
        mashup = self._extract_data_mashup()

        _log_progress(
            f"Páginas identificadas: {len(pages)}. Visuais identificados: {len(visuals)}."
        )

        return {
            "report": report_info,
            "package_manifest": self.package_manifest,
            "pages": pages,
            "visuals": visuals,
            "bookmarks": bookmarks,
            "report_filters": report_filters,
            "themes": themes,
            "custom_visuals": custom_visuals,
            "model": model,
            "queries": queries,
            "package_settings": package_settings,
            "package_metadata": package_metadata,
            "security": security,
            "content_types": content_types,
            "mashup": mashup,
            "text_payloads_available": sorted(self.text_payloads.keys(), key=str.lower),
            "summary": {
                "page_count": len(pages),
                "hidden_page_count": sum(1 for item in pages if item.get("hidden")),
                "visual_count": len(visuals),
                "bookmark_count": len(bookmarks),
                "report_filter_count": len(report_filters),
                "theme_color_count": len(themes.get("colors", [])),
                "custom_visual_count": len(custom_visuals),
                "table_count": len(model.get("tables", [])),
                "column_count": sum(len(table.get("columns", [])) for table in model.get("tables", [])),
                "measure_count": sum(len(table.get("measures", [])) for table in model.get("tables", [])),
                "relationship_count": len(model.get("relationships", [])),
                "query_count": len(queries),
                "mashup_query_count": len(mashup.get("queries", [])),
                "content_type_count": len(content_types),
                "package_member_count": len(self.package_manifest),
            },
        }

    def _extract_pages(self) -> list[dict[str, Any]]:
        if not self.layout_payload:
            return []

        pages = []
        for index, section in enumerate(self.layout_payload.get("sections", []), start=1):
            display_name = section.get("displayName") or section.get("name") or f"Página {index}"
            pages.append(
                {
                    "name": section.get("name"),
                    "display_name": display_name,
                    "ordinal": index,
                    "width": section.get("width"),
                    "height": section.get("height"),
                    "hidden": bool(section.get("isHidden") or section.get("hidden")),
                    "visual_count": len(section.get("visualContainers", [])),
                    "filters": self._extract_page_filters(section),
                    "config": self._safe_json_loads(section.get("config")) if isinstance(section.get("config"), str) else section.get("config"),
                }
            )
        return pages

    def _extract_page_filters(self, section: dict[str, Any]) -> list[dict[str, Any]]:
        raw_filters = section.get("filters")
        filters_payload = self._safe_json_loads(raw_filters) if isinstance(raw_filters, str) else raw_filters
        if not isinstance(filters_payload, list):
            return []
        return [self._normalize_filter_payload(item) for item in filters_payload if isinstance(item, dict)]

    def _extract_visuals(self, pages: list[dict[str, Any]]) -> list[dict[str, Any]]:
        if not self.layout_payload:
            return []

        page_lookup = {
            item.get("name"): item.get("display_name") or item.get("name")
            for item in pages
        }
        visuals = []
        for section in self.layout_payload.get("sections", []):
            page_name = section.get("name")
            page_label = page_lookup.get(page_name, page_name or "(sem página)")
            for index, container in enumerate(section.get("visualContainers", []), start=1):
                config_payload = self._safe_json_loads(container.get("config")) if isinstance(container.get("config"), str) else {}
                query_payload = self._safe_json_loads(container.get("query")) if isinstance(container.get("query"), str) else {}
                data_transforms = self._safe_json_loads(container.get("dataTransforms")) if isinstance(container.get("dataTransforms"), str) else {}
                visual_type = self._extract_visual_type(config_payload)
                title = self._extract_visual_title(config_payload)
                fields = self._extract_visual_fields(config_payload, query_payload, data_transforms)
                filters = self._extract_visual_filters(container)
                visuals.append(
                    {
                        "page_name": page_label,
                        "visual_index": index,
                        "visual_type": visual_type,
                        "title": title,
                        "x": container.get("x"),
                        "y": container.get("y"),
                        "z": container.get("z"),
                        "width": container.get("width"),
                        "height": container.get("height"),
                        "fields": fields,
                        "filter_count": len(filters),
                        "filters": filters,
                        "query": query_payload if isinstance(query_payload, dict) else {},
                    }
                )
        return visuals

    def _extract_visual_type(self, config_payload: dict[str, Any] | None) -> str:
        if not isinstance(config_payload, dict):
            return "(não identificado)"
        single_visual = config_payload.get("singleVisual") or {}
        return single_visual.get("visualType") or config_payload.get("visualType") or "(não identificado)"

    def _extract_visual_title(self, config_payload: dict[str, Any] | None) -> str | None:
        if not isinstance(config_payload, dict):
            return None
        single_visual = config_payload.get("singleVisual") or {}
        vc_objects = single_visual.get("vcObjects") or {}
        title = vc_objects.get("title")
        if isinstance(title, list):
            for candidate in title:
                properties = candidate.get("properties") or {}
                text_payload = properties.get("text") or {}
                expr = text_payload.get("expr") or {}
                literal = expr.get("Literal") or {}
                value = literal.get("Value")
                if isinstance(value, str):
                    return value.strip("'")
        return None

    def _extract_visual_fields(
        self,
        config_payload: dict[str, Any] | None,
        query_payload: dict[str, Any] | None,
        data_transforms: dict[str, Any] | None,
    ) -> list[str]:
        fields: list[str] = []
        if isinstance(config_payload, dict):
            projections = ((config_payload.get("singleVisual") or {}).get("projections")) or {}
            for projection_values in projections.values():
                if isinstance(projection_values, list):
                    for item in projection_values:
                        query_ref = item.get("queryRef")
                        if query_ref:
                            fields.append(str(query_ref))
        if isinstance(query_payload, dict):
            for command in query_payload.get("Commands", []) or []:
                semantic_query = ((command.get("SemanticQueryDataShapeCommand") or {}).get("Query")) or {}
                for select in semantic_query.get("Select", []) or []:
                    for key in ("Name", "NativeReferenceName"):
                        value = select.get(key)
                        if value:
                            fields.append(str(value))
        if isinstance(data_transforms, dict):
            projection_order = data_transforms.get("projectionOrdering") or {}
            if isinstance(projection_order, dict):
                for items in projection_order.values():
                    if isinstance(items, list):
                        fields.extend(str(item) for item in items if item)
        return sorted(unique_ordered([field for field in fields if field]), key=str.lower)

    def _extract_visual_filters(self, container: dict[str, Any]) -> list[dict[str, Any]]:
        raw_filters = container.get("filters")
        filters_payload = self._safe_json_loads(raw_filters) if isinstance(raw_filters, str) else raw_filters
        if not isinstance(filters_payload, list):
            return []
        return [self._normalize_filter_payload(item) for item in filters_payload if isinstance(item, dict)]

    def _extract_bookmarks(self) -> list[dict[str, Any]]:
        if not self.layout_payload:
            return []
        config_payload = self._layout_config_payload()
        if not isinstance(config_payload, dict):
            return []
        bookmarks = []
        for item in config_payload.get("bookmarks", []) or []:
            exploration_state = item.get("explorationState") or {}
            report_filters = ((exploration_state.get("filters") or {}).get("byExpr")) or []
            sections = exploration_state.get("sections") or {}
            section_names = sorted(unique_ordered(list(sections.keys())), key=str.lower)
            section_filter_count = sum(
                len((((section_state or {}).get("filters") or {}).get("byExpr")) or [])
                for section_state in sections.values()
            )
            bookmarks.append(
                {
                    "name": item.get("name"),
                    "display_name": item.get("displayName") or item.get("name"),
                    "active_section": exploration_state.get("activeSection"),
                    "report_filter_count": len(report_filters),
                    "section_filter_count": section_filter_count,
                    "sections": section_names,
                }
            )
        return bookmarks

    def _extract_report_filters(self) -> list[dict[str, Any]]:
        if not self.layout_payload:
            return []
        raw_filters = self.layout_payload.get("filters")
        filters_payload = self._safe_json_loads(raw_filters) if isinstance(raw_filters, str) else raw_filters
        if not isinstance(filters_payload, list):
            return []
        return [self._normalize_filter_payload(item) for item in filters_payload if isinstance(item, dict)]

    def _layout_config_payload(self) -> dict[str, Any] | None:
        if not isinstance(self.layout_payload, dict):
            return None
        config_payload = self.layout_payload.get("config")
        if isinstance(config_payload, str):
            config_payload = self._safe_json_loads(config_payload)
        return config_payload if isinstance(config_payload, dict) else None

    def _normalize_filter_payload(self, filter_payload: dict[str, Any]) -> dict[str, Any]:
        expression = filter_payload.get("expression")
        return {
            "name": filter_payload.get("name"),
            "type": filter_payload.get("filterType") or filter_payload.get("type"),
            "field": self._powerbi_expression_label(expression),
            "expression": compact_json(expression),
        }

    def _powerbi_expression_label(self, expression: dict[str, Any] | None) -> str | None:
        if not isinstance(expression, dict):
            return None
        if "Column" in expression:
            column = expression.get("Column") or {}
            source_ref = ((column.get("Expression") or {}).get("SourceRef")) or {}
            entity = source_ref.get("Entity") or source_ref.get("Source")
            prop = column.get("Property")
            if entity and prop:
                return f"{entity}.{prop}"
            return prop or entity
        if "Measure" in expression:
            measure = expression.get("Measure") or {}
            source_ref = ((measure.get("Expression") or {}).get("SourceRef")) or {}
            entity = source_ref.get("Entity") or source_ref.get("Source")
            prop = measure.get("Property")
            if entity and prop:
                return f"{entity}.{prop}"
            return prop or entity
        return None

    def _extract_themes(self) -> dict[str, Any]:
        theme_colors: list[str] = []
        theme_fonts: list[str] = []

        for member_name, payload in self.text_payloads.items():
            lowered = member_name.lower()
            if "basethemes" not in lowered and "registeredresources" not in lowered:
                continue
            theme_payload = self._safe_json_loads(payload)
            if not isinstance(theme_payload, dict):
                continue
            data_colors = theme_payload.get("dataColors") or []
            theme_colors.extend(str(color) for color in data_colors if color)
            for key in (
                "foreground", "foregroundNeutralSecondary", "foregroundNeutralTertiary",
                "background", "backgroundLight", "backgroundNeutral", "tableAccent",
                "good", "neutral", "bad", "maximum", "center", "minimum", "null",
                "hyperlink", "visitedHyperlink",
            ):
                value = theme_payload.get(key)
                if value:
                    theme_colors.append(f"{key}: {value}")
            text_classes = theme_payload.get("textClasses") or {}
            for key, value in text_classes.items():
                font_family = (value or {}).get("fontFace")
                if font_family:
                    theme_fonts.append(f"{key}: {font_family}")

        config_payload = self._layout_config_payload()
        if isinstance(config_payload, dict):
            theme_collection = config_payload.get("themeCollection") or {}
            base_theme = theme_collection.get("baseTheme") or {}
            for key, value in (base_theme.get("colors") or {}).items():
                if key and value:
                    theme_colors.append(f"{key}: {value}")
            for key, value in (base_theme.get("textClasses") or {}).items():
                font_family = ((value or {}).get("fontFace"))
                if font_family:
                    theme_fonts.append(f"{key}: {font_family}")

        return {
            "colors": sorted(unique_ordered(theme_colors), key=str.lower),
            "fonts": sorted(unique_ordered(theme_fonts), key=str.lower),
        }

    def _extract_model_metadata(self, visuals: list[dict[str, Any]]) -> dict[str, Any]:
        """Tenta localizar estruturas textuais de modelo no pacote PBIX."""
        candidates = []
        for member_name, payload in self.text_payloads.items():
            lowered = member_name.lower()
            if "schema" in lowered or "metadata" in lowered or "diagram" in lowered:
                data = self._safe_json_loads(payload)
                if isinstance(data, dict):
                    candidates.append((member_name, data))

        for member_name, data in candidates:
            model = data.get("model") if isinstance(data.get("model"), dict) else data
            tables = self._extract_tables_from_model_payload(model)
            relationships = self._extract_relationships_from_model_payload(model)
            if tables or relationships:
                return {
                    "source_member": member_name,
                    "tables": tables,
                    "relationships": relationships,
                    "accessible": True,
                    "notes": [],
                }

        diagram_tables = self._extract_tables_from_diagram_layout()
        inferred_tables = self._extract_tables_from_visual_queries()
        merged_tables = self._merge_inferred_tables(diagram_tables, inferred_tables)
        if merged_tables:
            return {
                "source_member": "DiagramLayout + Report/Layout",
                "tables": merged_tables,
                "relationships": [],
                "accessible": True,
                "notes": [
                    "As tabelas, colunas e medidas abaixo foram inferidas a partir do `DiagramLayout`, das queries dos visuais e dos filtros acessíveis no `Report/Layout`.",
                    "Relacionamentos completos ainda não estavam disponíveis em formato textual acessível dentro deste `.pbix`.",
                ],
            }

        if diagram_tables:
            return {
                "source_member": "DiagramLayout",
                "tables": diagram_tables,
                "relationships": [],
                "accessible": True,
                "notes": [
                    "As tabelas abaixo foram inferidas a partir do `DiagramLayout` do modelo.",
                    "Colunas, medidas e relacionamentos completos não estavam disponíveis em formato textual acessível dentro deste `.pbix`.",
                ],
            }

        return {
            "source_member": None,
            "tables": [],
            "relationships": [],
            "accessible": False,
            "notes": [
                "O modelo semântico completo não foi encontrado em formato textual acessível dentro do `.pbix`.",
                "Para cobertura mais profunda, o formato `.pbip` costuma oferecer estrutura mais adequada à documentação automática.",
            ],
        }

    def _extract_tables_from_diagram_layout(self) -> list[dict[str, Any]]:
        payload = self.text_payloads.get("DiagramLayout") or self.text_payloads.get("diagramlayout")
        data = self._safe_json_loads(payload)
        if not isinstance(data, dict):
            return []
        tables = []
        for diagram in data.get("diagrams", []) or []:
            for node in diagram.get("nodes", []) or []:
                node_name = clean_display_label(node.get("nodeIndex"))
                if not node_name:
                    continue
                tables.append(
                    {
                        "name": node_name,
                        "description": None,
                        "hidden": False,
                        "column_count": 0,
                        "measure_count": 0,
                        "columns": [],
                        "measures": [],
                    }
                )
        return sorted(unique_ordered(tables), key=lambda item: (item.get("name") or "").lower())

    def _extract_tables_from_visual_queries(self) -> list[dict[str, Any]]:
        if not self.layout_payload:
            return []

        table_map: dict[str, dict[str, Any]] = {}

        def ensure_table(table_name: str) -> dict[str, Any]:
            record = table_map.setdefault(
                table_name,
                {
                    "name": table_name,
                    "description": None,
                    "hidden": False,
                    "column_count": 0,
                    "measure_count": 0,
                    "columns": [],
                    "measures": [],
                },
            )
            return record

        def ensure_column(table_name: str, column_name: str) -> None:
            table = ensure_table(table_name)
            if column_name not in {item.get("name") for item in table["columns"]}:
                table["columns"].append(
                    {
                        "name": column_name,
                        "data_type": None,
                        "hidden": None,
                        "expression": None,
                    }
                )

        def ensure_measure(table_name: str, measure_name: str) -> None:
            table = ensure_table(table_name)
            if measure_name not in {item.get("name") for item in table["measures"]}:
                table["measures"].append(
                    {
                        "name": measure_name,
                        "expression": None,
                        "hidden": None,
                        "format_string": None,
                    }
                )

        for report_filter in self._extract_report_filters():
            field = report_filter.get("field")
            if field and "." in field:
                table_name, column_name = field.split(".", 1)
                ensure_column(table_name, column_name)

        for page in self.layout_payload.get("sections", []) or []:
            for page_filter in self._extract_page_filters(page):
                field = page_filter.get("field")
                if field and "." in field:
                    table_name, column_name = field.split(".", 1)
                    ensure_column(table_name, column_name)

            for container in page.get("visualContainers", []) or []:
                query_payload = self._safe_json_loads(container.get("query")) if isinstance(container.get("query"), str) else container.get("query")
                if not isinstance(query_payload, dict):
                    continue
                for command in query_payload.get("Commands", []) or []:
                    semantic_query = ((command.get("SemanticQueryDataShapeCommand") or {}).get("Query")) or {}
                    alias_map = {
                        str(item.get("Name")): item.get("Entity")
                        for item in semantic_query.get("From", []) or []
                        if item.get("Name") and item.get("Entity")
                    }
                    for select in semantic_query.get("Select", []) or []:
                        self._collect_select_into_model(select, alias_map, ensure_column, ensure_measure)

                for visual_filter in self._extract_visual_filters(container):
                    field = visual_filter.get("field")
                    if field and "." in field:
                        table_name, column_name = field.split(".", 1)
                        ensure_column(table_name, column_name)

        tables = []
        for table in table_map.values():
            table["columns"] = sorted(table["columns"], key=lambda item: (item.get("name") or "").lower())
            table["measures"] = sorted(table["measures"], key=lambda item: (item.get("name") or "").lower())
            table["column_count"] = len(table["columns"])
            table["measure_count"] = len(table["measures"])
            tables.append(table)
        return sorted(tables, key=lambda item: (item.get("name") or "").lower())

    def _collect_select_into_model(
        self,
        select: dict[str, Any],
        alias_map: dict[str, str],
        ensure_column: Any,
        ensure_measure: Any,
    ) -> None:
        if not isinstance(select, dict):
            return

        if isinstance(select.get("Column"), dict):
            table_name, field_name = self._resolve_table_field_reference(select["Column"], alias_map)
            if table_name and field_name:
                ensure_column(table_name, field_name)
            return

        if isinstance(select.get("Measure"), dict):
            table_name, field_name = self._resolve_table_field_reference(select["Measure"], alias_map)
            if table_name and field_name:
                ensure_measure(table_name, field_name)
            return

        if isinstance(select.get("Aggregation"), dict):
            expression = select["Aggregation"].get("Expression")
            if isinstance(expression, dict) and isinstance(expression.get("Column"), dict):
                table_name, field_name = self._resolve_table_field_reference(expression["Column"], alias_map)
                if table_name and field_name:
                    ensure_column(table_name, field_name)
            return

        if isinstance(select.get("HierarchyLevel"), dict):
            expression = select["HierarchyLevel"].get("Expression")
            if isinstance(expression, dict) and isinstance(expression.get("Hierarchy"), dict):
                table_name, field_name = self._resolve_table_field_reference(expression["Hierarchy"], alias_map)
                if table_name and field_name:
                    ensure_column(table_name, field_name)

    def _resolve_table_field_reference(
        self,
        payload: dict[str, Any],
        alias_map: dict[str, str],
    ) -> tuple[str | None, str | None]:
        expression = payload.get("Expression") or {}
        source_ref = expression.get("SourceRef") or {}
        alias = source_ref.get("Source")
        entity = source_ref.get("Entity") or alias_map.get(str(alias))
        prop = payload.get("Property")
        return (entity, prop)

    def _merge_inferred_tables(
        self,
        base_tables: list[dict[str, Any]],
        inferred_tables: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        merged: dict[str, dict[str, Any]] = {}
        for source in [base_tables, inferred_tables]:
            for table in source:
                name = table.get("name")
                if not name:
                    continue
                record = merged.setdefault(
                    name,
                    {
                        "name": name,
                        "description": table.get("description"),
                        "hidden": table.get("hidden"),
                        "columns": [],
                        "measures": [],
                    },
                )
                record["description"] = record.get("description") or table.get("description")
                for column in table.get("columns", []):
                    if column.get("name") not in {item.get("name") for item in record["columns"]}:
                        record["columns"].append(column)
                for measure in table.get("measures", []):
                    if measure.get("name") not in {item.get("name") for item in record["measures"]}:
                        record["measures"].append(measure)

        output = []
        for table in merged.values():
            table["columns"] = sorted(table["columns"], key=lambda item: (item.get("name") or "").lower())
            table["measures"] = sorted(table["measures"], key=lambda item: (item.get("name") or "").lower())
            table["column_count"] = len(table["columns"])
            table["measure_count"] = len(table["measures"])
            output.append(table)
        return sorted(output, key=lambda item: (item.get("name") or "").lower())

    def _extract_tables_from_model_payload(self, model_payload: dict[str, Any]) -> list[dict[str, Any]]:
        tables = []
        for table in model_payload.get("tables", []) if isinstance(model_payload, dict) else []:
            columns = []
            for column in table.get("columns", []) or []:
                columns.append(
                    {
                        "name": column.get("name"),
                        "data_type": column.get("dataType") or column.get("type"),
                        "hidden": column.get("isHidden") or column.get("hidden"),
                        "expression": normalize_whitespace(column.get("expression")),
                    }
                )
            measures = []
            for measure in table.get("measures", []) or []:
                measures.append(
                    {
                        "name": measure.get("name"),
                        "expression": normalize_whitespace(measure.get("expression")),
                        "hidden": measure.get("isHidden") or measure.get("hidden"),
                        "format_string": measure.get("formatString"),
                    }
                )
            tables.append(
                {
                    "name": table.get("name"),
                    "description": table.get("description"),
                    "hidden": table.get("isHidden") or table.get("hidden"),
                    "column_count": len(columns),
                    "measure_count": len(measures),
                    "columns": columns,
                    "measures": measures,
                }
            )
        return tables

    def _extract_relationships_from_model_payload(self, model_payload: dict[str, Any]) -> list[dict[str, Any]]:
        relationships = []
        for relationship in model_payload.get("relationships", []) if isinstance(model_payload, dict) else []:
            relationships.append(
                {
                    "from_table": relationship.get("fromTable"),
                    "from_column": relationship.get("fromColumn"),
                    "to_table": relationship.get("toTable"),
                    "to_column": relationship.get("toColumn"),
                    "cross_filtering_behavior": relationship.get("crossFilteringBehavior"),
                    "active": relationship.get("isActive"),
                }
            )
        return relationships

    def _extract_query_artifacts(self) -> list[dict[str, Any]]:
        artifacts = []
        for member_name, payload in self.text_payloads.items():
            lowered = member_name.lower()
            if lowered in {"connections", "report/layout"}:
                continue
            if "query" in lowered or "mashup" in lowered or "connection" in lowered:
                artifacts.append(
                    {
                        "member_name": member_name,
                        "preview": normalize_whitespace(payload[:1000]),
                    }
                )
        connections_payload = self.text_payloads.get("Connections") or self.text_payloads.get("connections")
        if connections_payload:
            artifacts.append(
                {
                    "member_name": "Connections",
                    "preview": normalize_whitespace(connections_payload[:1000]),
                }
            )
        return artifacts

    def _extract_package_metadata(self) -> dict[str, Any]:
        payload = self._safe_json_loads(self.text_payloads.get("Metadata") or self.text_payloads.get("metadata"))
        return payload if isinstance(payload, dict) else {}

    def _extract_package_settings(self) -> dict[str, Any]:
        payload = self._safe_json_loads(self.text_payloads.get("Settings") or self.text_payloads.get("settings"))
        return payload if isinstance(payload, dict) else {}

    def _extract_security_bindings(self) -> dict[str, Any]:
        member_name = "SecurityBindings" if "SecurityBindings" in self.package_member_bytes else "securitybindings"
        raw_bytes = self.package_member_bytes.get(member_name)
        if not raw_bytes:
            return {"present": False, "size_bytes": 0, "is_binary": False}
        return {
            "present": True,
            "size_bytes": len(raw_bytes),
            "is_binary": True,
            "preview_hex": raw_bytes[:64].hex(),
        }

    def _extract_content_types(self) -> list[dict[str, Any]]:
        payload = self.text_payloads.get("[Content_Types].xml")
        if not payload:
            return []
        default_matches = re.findall(
            r'<Default\s+Extension="([^"]*)"\s+ContentType="([^"]*)"\s*/?>',
            payload,
            flags=re.IGNORECASE,
        )
        override_matches = re.findall(
            r'<Override\s+PartName="([^"]*)"\s+ContentType="([^"]*)"\s*/?>',
            payload,
            flags=re.IGNORECASE,
        )
        rows = [
            {"kind": "default", "name": extension, "content_type": content_type}
            for extension, content_type in default_matches
        ]
        rows.extend(
            {"kind": "override", "name": part_name, "content_type": content_type}
            for part_name, content_type in override_matches
        )
        return rows

    def _extract_data_mashup(self) -> dict[str, Any]:
        """Tenta extrair consultas e textos úteis de `DataMashup`, quando presente."""
        member_name = None
        for candidate in self.package_member_bytes.keys():
            if candidate.lower() == "datamashup":
                member_name = candidate
                break
        if not member_name:
            return {
                "present": False,
                "queries": [],
                "parameters": [],
                "functions": [],
                "notes": ["O arquivo `.pbix` analisado não contém o membro `DataMashup` no pacote."],
            }

        raw_bytes = self.package_member_bytes.get(member_name, b"")
        if not raw_bytes:
            return {
                "present": True,
                "queries": [],
                "parameters": [],
                "functions": [],
                "notes": ["O membro `DataMashup` foi encontrado, mas não pôde ser lido."],
            }

        notes: list[str] = []
        texts: list[str] = []

        zip_texts = self._extract_texts_from_embedded_zip(raw_bytes)
        if zip_texts:
            texts.extend(zip_texts)
            notes.append("Conteúdo textual extraído de um pacote interno ZIP em `DataMashup`.")

        ole_texts = self._extract_texts_from_ole_container(raw_bytes)
        if ole_texts:
            texts.extend(ole_texts)
            notes.append("Conteúdo textual extraído de streams OLE em `DataMashup`.")

        if not texts:
            best_effort_text = self._extract_text_candidates_from_binary(raw_bytes)
            if best_effort_text:
                texts.extend(best_effort_text)
                notes.append("Conteúdo textual recuperado por varredura heurística do binário de `DataMashup`.")

        texts = unique_ordered([text for text in texts if text])
        queries = self._extract_mashup_queries_from_texts(texts)
        parameters = self._extract_mashup_parameters_from_texts(texts)
        functions = self._extract_mashup_functions_from_texts(texts)

        if not notes:
            notes.append("O membro `DataMashup` foi encontrado, mas nenhuma estrutura textual reconhecível foi extraída.")

        return {
            "present": True,
            "queries": queries,
            "parameters": parameters,
            "functions": functions,
            "notes": notes,
        }

    def _extract_texts_from_embedded_zip(self, raw_bytes: bytes) -> list[str]:
        texts: list[str] = []
        try:
            with zipfile.ZipFile(BytesIO(raw_bytes), "r") as archive:
                for info in archive.infolist():
                    if info.is_dir():
                        continue
                    content = archive.read(info.filename)
                    decoded = self._decode_package_text(content)
                    if decoded:
                        texts.append(decoded)
        except zipfile.BadZipFile:
            return []
        return texts

    def _extract_texts_from_ole_container(self, raw_bytes: bytes) -> list[str]:
        if olefile is None:
            return []
        texts: list[str] = []
        try:
            ole = olefile.OleFileIO(BytesIO(raw_bytes))
        except Exception:
            return []
        try:
            for stream_name in ole.listdir(streams=True, storages=False):
                try:
                    stream = ole.openstream(stream_name).read()
                except Exception:
                    continue
                decoded = self._decode_package_text(stream)
                if decoded:
                    texts.append(decoded)
        finally:
            ole.close()
        return texts

    def _extract_text_candidates_from_binary(self, raw_bytes: bytes) -> list[str]:
        texts: list[str] = []
        for encoding in ("utf-8", "utf-16-le", "latin-1"):
            try:
                decoded = raw_bytes.decode(encoding, errors="ignore")
            except Exception:
                continue
            candidates = re.findall(
                r"(section\s+[^\n]{0,120}.*?shared\s+[^\n]{0,120}.*?in\s+[^\n]{0,120})",
                decoded,
                flags=re.IGNORECASE | re.DOTALL,
            )
            texts.extend(candidate[:12000] for candidate in candidates if candidate.strip())
        return texts

    def _extract_mashup_queries_from_texts(self, texts: list[str]) -> list[dict[str, Any]]:
        records: list[dict[str, Any]] = []
        seen = set()
        pattern = re.compile(
            r"shared\s+([A-Za-zÀ-ÿ0-9_ ]+)\s*=\s*(.*?)(?=\nshared\s+[A-Za-zÀ-ÿ0-9_ ]+\s*=|\Z)",
            flags=re.IGNORECASE | re.DOTALL,
        )
        for text in texts:
            for match in pattern.finditer(text):
                name = clean_display_label(match.group(1))
                body = normalize_whitespace(match.group(2))
                if not name or not body:
                    continue
                key = (name.lower(), body[:500])
                if key in seen:
                    continue
                seen.add(key)
                records.append(
                    {
                        "name": name,
                        "m_code": body[:12000],
                    }
                )
        return sorted(records, key=lambda item: (item.get("name") or "").lower())

    def _extract_mashup_parameters_from_texts(self, texts: list[str]) -> list[dict[str, Any]]:
        records: list[dict[str, Any]] = []
        seen = set()
        for item in self._extract_mashup_queries_from_texts(texts):
            code = item.get("m_code") or ""
            if "IsParameterQuery=true" not in code and "IsParameterQuery = true" not in code:
                continue
            key = item.get("name")
            if not key or key in seen:
                continue
            seen.add(key)
            records.append(
                {
                    "name": item.get("name"),
                    "preview": code[:2000],
                }
            )
        return sorted(records, key=lambda item: (item.get("name") or "").lower())

    def _extract_mashup_functions_from_texts(self, texts: list[str]) -> list[dict[str, Any]]:
        records: list[dict[str, Any]] = []
        seen = set()
        for item in self._extract_mashup_queries_from_texts(texts):
            code = item.get("m_code") or ""
            if "=>" not in code:
                continue
            key = item.get("name")
            if not key or key in seen:
                continue
            seen.add(key)
            records.append(
                {
                    "name": item.get("name"),
                    "preview": code[:2000],
                }
            )
        return sorted(records, key=lambda item: (item.get("name") or "").lower())

    def _extract_custom_visuals(self, visuals: list[dict[str, Any]]) -> list[str]:
        visuals_types = [item.get("visual_type") for item in visuals if item.get("visual_type")]
        base_visuals = {
            "barChart", "columnChart", "lineChart", "pieChart", "tableEx", "matrix", "card",
            "slicer", "map", "filledMap", "scatterChart", "clusteredColumnChart", "clusteredBarChart",
        }
        return sorted(
            unique_ordered(
                [item for item in visuals_types if item not in base_visuals and item != "(não identificado)"]
            ),
            key=str.lower,
        )

    def generate_package_structure_map(self) -> tuple[Path, Path]:
        map_md_path = self.output_dir / "package_structure_map.md"
        map_json_path = self.output_dir / "package_structure_map.json"
        rows = sorted(self.package_manifest, key=lambda item: (item.get("kind") or "", item.get("path") or ""))
        markdown_lines = [
            "# Mapa da Estrutura do PBIX",
            "",
            "| Caminho | Tipo | Tamanho (bytes) |",
            "| --- | --- | ---: |",
        ]
        for row in rows:
            markdown_lines.append(
                f"| `{row['path']}` | {row['kind']} | {row['size_bytes']} |"
            )
        map_md_path.write_text("\n".join(markdown_lines) + "\n", encoding="utf-8")
        map_json_path.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
        return map_md_path, map_json_path

    def write_outputs(self) -> list[Path]:
        written_files: list[Path] = []

        package_summary_path = self.output_dir / f"{self.base_name}_package_manifest.json"
        package_summary_path.write_text(
            json.dumps(self.package_manifest, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        written_files.append(package_summary_path)

        structure_md_path, structure_json_path = self.generate_package_structure_map()
        written_files.extend([structure_md_path, structure_json_path])

        if self.output_format in {"all", "json"}:
            written_files.append(self._write_json())
        if self.output_format in {"all", "markdown"}:
            written_files.append(self._write_markdown())
            written_files.append(self._write_rtf())
        if self.output_format == "rtf":
            written_files.append(self._write_rtf())
        if self.output_format in {"all", "docx"}:
            written_files.append(self._write_docx())
        if self.output_format in {"all", "excel"}:
            written_files.append(self._write_excel())

        manifest_path = self.output_dir / f"{self.base_name}_manifest.json"
        manifest_path.write_text(
            json.dumps(
                {
                    "source_file": str(self.source_path),
                    "output_directory": str(self.output_dir),
                    "generated_files": [str(path) for path in written_files],
                    "package_manifest": self.package_manifest,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        written_files.append(manifest_path)
        self._cleanup_temporary_outputs()
        return written_files

    def _cleanup_temporary_outputs(self) -> None:
        _log_progress("Removendo artefatos temporários de processamento.")
        for name in TEMPORARY_OUTPUT_NAMES:
            path = self.output_dir / name
            if not path.exists():
                continue
            if path.is_dir():
                shutil.rmtree(path, ignore_errors=True)
            else:
                try:
                    path.unlink()
                except FileNotFoundError:
                    pass

    def _write_json(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.json"
        output_path.write_text(
            json.dumps(self.metadata, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return output_path

    def _write_markdown(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.md"
        summary = self.metadata["summary"]
        lines = [
            f"# Documentação do Relatório Power BI - {self.source_path.name}",
            "",
            f"Relatório gerado em {self.generated_at}",
            "",
            "## Resumo",
            "",
            f"- Caminho de origem: `{self.source_path}`",
            f"- Tipo do arquivo: `{self.source_path.suffix.lower()}`",
            f"- Última alteração do arquivo: `{self.metadata['report']['source_file_last_modified']}`",
            f"- Páginas: {summary['page_count']} ({summary['hidden_page_count']} ocultas)",
            f"- Visuais: {summary['visual_count']}",
            f"- Bookmarks: {summary['bookmark_count']}",
            f"- Filtros do relatório: {summary['report_filter_count']}",
            f"- Tabelas do modelo: {summary['table_count']}",
            f"- Medidas: {summary['measure_count']}",
            f"- Relacionamentos: {summary['relationship_count']}",
            "",
        ]
        lines.extend(self._build_package_markdown())
        lines.extend(self._build_pages_markdown())
        lines.extend(self._build_bookmarks_and_filters_markdown())
        lines.extend(self._build_model_markdown())
        lines.extend(self._build_visual_tokens_markdown())
        lines.extend(self._build_query_artifacts_markdown())
        lines.extend(self._build_mashup_markdown())
        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return output_path

    def _write_rtf(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.rtf"
        body = self._build_rtf_document()
        header = (
            r"{\rtf1\ansi\deff0"
            rf"{{\fonttbl{{\f0 {self._rtf_escape(RTF_BODY_FONT_NAME)};}}{{\f1 {self._rtf_escape(RTF_MONO_FONT_NAME)};}}}}"
            r"\viewkind4\uc1"
        )
        output_path.write_text(header + body + "}", encoding="utf-8")
        return output_path

    def _write_docx(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.docx"
        document = Document()
        self._configure_docx_document(document)
        blocks = self._build_document_blocks()
        title_block_count = 2 if len(blocks) >= 2 else len(blocks)
        for block in blocks[:title_block_count]:
            self._append_docx_block(document, block)
        if blocks:
            spacer = document.add_paragraph()
            self._apply_docx_paragraph_format(spacer)
            self._append_docx_toc(document)
            document.add_page_break()
        for block in blocks[title_block_count:]:
            self._append_docx_block(document, block)
        document.save(output_path)
        return output_path

    def _write_excel(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.xlsx"
        with pd.ExcelWriter(output_path) as writer:
            self._to_frame(self.metadata["pages"], stringify_nested=True).to_excel(writer, sheet_name="Pages", index=False)
            self._to_frame(self.metadata["visuals"], stringify_nested=True).to_excel(writer, sheet_name="Visuals", index=False)
            self._to_frame(self.metadata["report_filters"], stringify_nested=True).to_excel(writer, sheet_name="ReportFilters", index=False)
            self._to_frame(self.metadata["bookmarks"]).to_excel(writer, sheet_name="Bookmarks", index=False)
            self._to_frame(self.metadata["package_manifest"]).to_excel(writer, sheet_name="PackageManifest", index=False)
            self._to_frame(self.metadata["queries"], stringify_nested=True).to_excel(writer, sheet_name="Queries", index=False)
            self._to_frame([self.metadata["package_metadata"]], stringify_nested=True).to_excel(writer, sheet_name="PackageMetadata", index=False)
            self._to_frame([self.metadata["package_settings"]], stringify_nested=True).to_excel(writer, sheet_name="Settings", index=False)
            self._to_frame([self.metadata["security"]], stringify_nested=True).to_excel(writer, sheet_name="Security", index=False)
            self._to_frame(self.metadata["content_types"]).to_excel(writer, sheet_name="ContentTypes", index=False)
            self._to_frame(self.metadata["mashup"]["queries"], stringify_nested=True).to_excel(writer, sheet_name="MashupQueries", index=False)
            self._to_frame(self.metadata["mashup"]["parameters"], stringify_nested=True).to_excel(writer, sheet_name="MashupParameters", index=False)
            self._to_frame(self.metadata["mashup"]["functions"], stringify_nested=True).to_excel(writer, sheet_name="MashupFunctions", index=False)

            model_tables = self.metadata["model"]["tables"]
            tables_flat = [
                {
                    "table_name": table.get("name"),
                    "description": table.get("description"),
                    "hidden": table.get("hidden"),
                    "column_count": table.get("column_count"),
                    "measure_count": table.get("measure_count"),
                }
                for table in model_tables
            ]
            columns_flat = [
                {"table_name": table.get("name"), **column}
                for table in model_tables
                for column in table.get("columns", [])
            ]
            measures_flat = [
                {"table_name": table.get("name"), **measure}
                for table in model_tables
                for measure in table.get("measures", [])
            ]
            self._to_frame(tables_flat).to_excel(writer, sheet_name="ModelTables", index=False)
            self._to_frame(columns_flat, stringify_nested=True).to_excel(writer, sheet_name="ModelColumns", index=False)
            self._to_frame(measures_flat, stringify_nested=True).to_excel(writer, sheet_name="ModelMeasures", index=False)
            self._to_frame(self.metadata["model"]["relationships"], stringify_nested=True).to_excel(writer, sheet_name="Relationships", index=False)
        return output_path

    def _to_frame(self, records: list[dict[str, Any]], stringify_nested: bool = False) -> pd.DataFrame:
        if not records:
            return pd.DataFrame()
        df = pd.DataFrame(records)
        if stringify_nested:
            for column in df.columns:
                df[column] = df[column].apply(
                    lambda value: json.dumps(value, ensure_ascii=False, indent=2)
                    if isinstance(value, (dict, list))
                    else value
                )
        return df

    def _build_document_blocks(self) -> list[dict[str, Any]]:
        parts: list[dict[str, Any]] = []
        summary = self.metadata["summary"]
        report = self.metadata["report"]
        model = self.metadata["model"]

        parts.append(self._doc_paragraph(f"Documentação do Relatório Power BI - {self.source_path.name}", style="title"))
        parts.append(self._doc_paragraph(f"Relatório gerado em {self.generated_at}", style="subtitle"))

        parts.append(self._doc_paragraph("Seção de dados gerais", style="section"))
        parts.append(self._doc_paragraph("Resumo", style="subsection"))
        parts.append(self._doc_bullet(f"Caminho de origem: {self.source_path}"))
        parts.append(self._doc_bullet(f"Tipo do arquivo: {self.source_path.suffix.lower()}"))
        parts.append(self._doc_bullet(f"Última alteração do arquivo: {report['source_file_last_modified']}"))
        parts.append(self._doc_bullet(f"Páginas: {summary['page_count']} ({summary['hidden_page_count']} ocultas)"))
        parts.append(self._doc_bullet(f"Visuais: {summary['visual_count']}"))
        parts.append(self._doc_bullet(f"Bookmarks: {summary['bookmark_count']}"))
        parts.append(self._doc_bullet(f"Filtros do relatório: {summary['report_filter_count']}"))
        parts.append(self._doc_bullet(f"Tabelas do modelo: {summary['table_count']}"))
        parts.append(self._doc_bullet(f"Medidas: {summary['measure_count']}"))
        parts.append(self._doc_bullet(f"Relacionamentos: {summary['relationship_count']}"))
        parts.append(self._doc_bullet(f"Itens do pacote: {summary['package_member_count']}"))

        parts.append(self._doc_paragraph("Pacote PBIX", style="section"))
        parts.append(self._doc_paragraph("Manifesto do pacote", style="subsection"))
        parts.append(self._doc_bullet(f"Arquivo de versão: {report.get('version') or '-'}"))
        if self.metadata["package_metadata"]:
            parts.append(self._doc_bullet(f"Origem declarada: {self.metadata['package_metadata'].get('CreatedFrom') or '-'}"))
            parts.append(self._doc_bullet(f"Release de origem: {self.metadata['package_metadata'].get('CreatedFromRelease') or '-'}"))
        if self.metadata["package_settings"]:
            report_settings = self.metadata["package_settings"].get("ReportSettings") or {}
            query_settings = self.metadata["package_settings"].get("QueriesSettings") or {}
            parts.append(self._doc_bullet(f"Mostrar campos ocultos: {'sim' if report_settings.get('ShowHiddenFields') else 'não'}"))
            parts.append(self._doc_bullet(f"Detecção de tipo habilitada: {'sim' if query_settings.get('TypeDetectionEnabled') else 'não'}"))
            parts.append(self._doc_bullet(f"Importação automática de relacionamentos: {'sim' if query_settings.get('RelationshipImportEnabled') else 'não'}"))
        if self.metadata["security"].get("present"):
            parts.append(self._doc_bullet(f"SecurityBindings presente: sim ({self.metadata['security'].get('size_bytes')} bytes)"))
        else:
            parts.append(self._doc_bullet("SecurityBindings presente: não"))
        parts.append(self._doc_bullet(f"Entradas em [Content_Types].xml: {len(self.metadata['content_types'])}"))
        package_sample = self.metadata["package_manifest"][:20]
        for item in package_sample:
            parts.append(
                self._doc_bullet(
                    f"{item.get('path')} | tipo={item.get('kind')} | tamanho={item.get('size_bytes')}",
                    level=1,
                    mono=True,
                )
            )
        if len(self.metadata["package_manifest"]) > len(package_sample):
            parts.append(self._doc_bullet("Demais itens listados no manifesto JSON e no mapa de estrutura.", level=1))

        parts.append(self._doc_paragraph("Páginas e visuais", style="section"))
        parts.append(self._doc_paragraph("Páginas", style="subsection"))
        if not self.metadata["pages"]:
            parts.append(self._doc_bullet("Nenhuma página identificada."))
        for page in self.metadata["pages"]:
            parts.append(self._doc_paragraph(page.get("display_name") or page.get("name") or "(sem nome)", style="subsubsection"))
            parts.append(self._doc_bullet(f"Oculta: {'sim' if page.get('hidden') else 'não'}", level=1))
            parts.append(self._doc_bullet(f"Visuais na página: {page.get('visual_count')}", level=1))
            parts.append(self._doc_bullet(f"Dimensões: {page.get('width')} x {page.get('height')}", level=1))
            if page.get("filters"):
                parts.extend(
                    self._doc_list_block(
                        "Filtros da página",
                        [
                            f"{item.get('field') or '-'} | tipo={item.get('type') or '-'}"
                            for item in page.get("filters", [])
                        ],
                        level=1,
                        empty_text="nenhum filtro",
                    )
                )
            page_visuals = [
                item for item in self.metadata["visuals"]
                if item.get("page_name") == (page.get("display_name") or page.get("name"))
            ]
            if page_visuals:
                parts.append(self._doc_paragraph("Visuais", style="body_bold", level=1))
                for visual in page_visuals:
                    visual_label = visual.get("title") or f"Visual {visual.get('visual_index')}"
                    parts.append(
                        self._doc_bullet(
                            f"{visual_label} | tipo={visual.get('visual_type')} | campos={', '.join(visual.get('fields', [])) or '-'}",
                            level=2,
                        )
                    )

        parts.append(self._doc_paragraph("Bookmarks e filtros", style="section"))
        parts.append(self._doc_paragraph("Bookmarks", style="subsection"))
        if self.metadata["bookmarks"]:
            for bookmark in self.metadata["bookmarks"]:
                parts.append(self._doc_paragraph(bookmark.get("display_name") or bookmark.get("name") or "(sem nome)", style="subsubsection"))
                parts.append(self._doc_bullet(f"Página ativa: {bookmark.get('active_section') or '-'}", level=1))
                parts.append(self._doc_bullet(f"Filtros de relatório no bookmark: {bookmark.get('report_filter_count')}", level=1))
                parts.append(self._doc_bullet(f"Filtros de seção no bookmark: {bookmark.get('section_filter_count')}", level=1))
        else:
            parts.append(self._doc_bullet("Nenhum bookmark identificado."))

        parts.append(self._doc_paragraph("Filtros do relatório", style="subsection"))
        if self.metadata["report_filters"]:
            for item in self.metadata["report_filters"]:
                parts.append(
                    self._doc_bullet(
                        f"{item.get('field') or '-'} | tipo={item.get('type') or '-'}",
                        level=1,
                    )
                )
        else:
            parts.append(self._doc_bullet("Nenhum filtro de relatório identificado."))

        parts.append(self._doc_paragraph("Modelo semântico", style="section"))
        parts.append(self._doc_paragraph("Tabelas e medidas", style="subsection"))
        if model.get("tables"):
            for note in model.get("notes", []):
                parts.append(self._doc_bullet(note))
            for table in model["tables"]:
                parts.append(self._doc_paragraph(table.get("name") or "(sem nome)", style="subsubsection"))
                parts.append(self._doc_bullet(f"Oculta: {'sim' if table.get('hidden') else 'não'}", level=1))
                parts.append(self._doc_bullet(f"Colunas: {table.get('column_count')}", level=1))
                parts.append(self._doc_bullet(f"Medidas: {table.get('measure_count')}", level=1))
                if table.get("columns"):
                    parts.extend(
                        self._doc_list_block(
                            "Colunas inferidas",
                            [item.get("name") for item in table.get("columns", []) if item.get("name")],
                            level=1,
                            empty_text="nenhuma coluna",
                        )
                    )
                if table.get("measures"):
                    parts.extend(
                        self._doc_list_block(
                            "Medidas inferidas",
                            [item.get("name") for item in table.get("measures", []) if item.get("name")],
                            level=1,
                            empty_text="nenhuma medida",
                        )
                    )
        else:
            for note in model.get("notes", []):
                parts.append(self._doc_bullet(note))

        parts.append(self._doc_paragraph("Tokens visuais e temas", style="section"))
        parts.append(self._doc_paragraph("Tema", style="subsection"))
        parts.extend(
            self._doc_list_block(
                "Cores do tema",
                self.metadata["themes"].get("colors", []),
                level=0,
                empty_text="nenhuma cor identificada",
            )
        )
        parts.extend(
            self._doc_list_block(
                "Fontes do tema",
                self.metadata["themes"].get("fonts", []),
                level=0,
                empty_text="nenhuma fonte identificada",
            )
        )

        parts.append(self._doc_paragraph("Artefatos de consulta", style="section"))
        parts.append(self._doc_paragraph("Metadados acessíveis", style="subsection"))
        if self.metadata["queries"]:
            for artifact in self.metadata["queries"]:
                parts.append(self._doc_paragraph(artifact.get("member_name") or "(sem nome)", style="subsubsection"))
                parts.append(self._doc_code_block(artifact.get("preview") or "-", level=1))
        else:
            parts.append(self._doc_bullet("Nenhum artefato textual de consulta foi extraído."))

        parts.append(self._doc_paragraph("DataMashup", style="subsection"))
        mashup = self.metadata["mashup"]
        parts.append(self._doc_bullet(f"Presente no pacote: {'sim' if mashup.get('present') else 'não'}"))
        parts.append(self._doc_bullet(f"Queries extraídas: {len(mashup.get('queries', []))}"))
        parts.append(self._doc_bullet(f"Parâmetros extraídos: {len(mashup.get('parameters', []))}"))
        parts.append(self._doc_bullet(f"Funções extraídas: {len(mashup.get('functions', []))}"))
        for note in mashup.get("notes", []):
            parts.append(self._doc_bullet(note, level=1))
        for query in mashup.get("queries", [])[:10]:
            parts.append(self._doc_paragraph(query.get("name") or "(sem nome)", style="subsubsection"))
            parts.append(self._doc_code_block(query.get("m_code") or "-", level=1))
        return parts

    def _build_rtf_document(self) -> str:
        return "".join(self._render_rtf_block(block) for block in self._build_document_blocks())

    def _build_package_markdown(self) -> list[str]:
        lines = ["## Pacote PBIX", ""]
        lines.append(f"- Itens no pacote: {len(self.metadata['package_manifest'])}")
        lines.append(f"- Layout principal acessível: {'sim' if self.metadata['report']['has_layout'] else 'não'}")
        lines.append(f"- Versão declarada: `{self.metadata['report'].get('version') or '-'}`")
        if self.metadata["package_metadata"]:
            lines.append(f"- Origem declarada: `{self.metadata['package_metadata'].get('CreatedFrom') or '-'}`")
            lines.append(f"- Release de origem: `{self.metadata['package_metadata'].get('CreatedFromRelease') or '-'}`")
        if self.metadata["package_settings"]:
            report_settings = self.metadata["package_settings"].get("ReportSettings") or {}
            query_settings = self.metadata["package_settings"].get("QueriesSettings") or {}
            lines.append(f"- Mostrar campos ocultos: `{'sim' if report_settings.get('ShowHiddenFields') else 'não'}`")
            lines.append(f"- Detecção de tipo habilitada: `{'sim' if query_settings.get('TypeDetectionEnabled') else 'não'}`")
            lines.append(f"- Importação automática de relacionamentos: `{'sim' if query_settings.get('RelationshipImportEnabled') else 'não'}`")
        lines.append(
            f"- SecurityBindings: `{'presente' if self.metadata['security'].get('present') else 'não identificado'}`"
        )
        lines.append(f"- Entradas em `[Content_Types].xml`: {len(self.metadata['content_types'])}")
        lines.append("")
        lines.append("### Amostra do manifesto")
        lines.append("")
        for item in self.metadata["package_manifest"][:25]:
            lines.append(f"- `{item['path']}` | tipo: `{item['kind']}` | tamanho: `{item['size_bytes']}`")
        lines.append("")
        return lines

    def _build_pages_markdown(self) -> list[str]:
        lines = ["## Páginas e Visuais", ""]
        if not self.metadata["pages"]:
            lines.append("- Nenhuma página identificada.")
            return lines
        for page in self.metadata["pages"]:
            page_name = page.get("display_name") or page.get("name") or "(sem nome)"
            lines.append(f"### {page_name}")
            lines.append("")
            lines.append(f"- Oculta: `{'sim' if page.get('hidden') else 'não'}`")
            lines.append(f"- Visuais: {page.get('visual_count')}")
            lines.append(f"- Dimensões: `{page.get('width')}` x `{page.get('height')}`")
            if page.get("filters"):
                lines.append("- Filtros da página:")
                for filter_item in page.get("filters", []):
                    lines.append(
                        f"  - {filter_item.get('field') or '-'} | tipo: `{filter_item.get('type') or '-'}`"
                    )
            page_visuals = [item for item in self.metadata["visuals"] if item.get("page_name") == page_name]
            if page_visuals:
                lines.append("- Visuais encontrados:")
                for visual in page_visuals:
                    visual_title = visual.get("title") or f"Visual {visual.get('visual_index')}"
                    lines.append(
                        f"  - {visual_title} | tipo: `{visual.get('visual_type')}` | campos: {', '.join(visual.get('fields', [])) or '-'}"
                    )
            lines.append("")
        return lines

    def _build_bookmarks_and_filters_markdown(self) -> list[str]:
        lines = ["## Bookmarks e Filtros", ""]
        lines.append("### Bookmarks")
        lines.append("")
        if self.metadata["bookmarks"]:
            for bookmark in self.metadata["bookmarks"]:
                lines.append(f"- {bookmark.get('display_name') or bookmark.get('name')}")
                lines.append(f"  - Página ativa: `{bookmark.get('active_section') or '-'}`")
                lines.append(f"  - Filtros de relatório: {bookmark.get('report_filter_count')}")
                lines.append(f"  - Filtros de seção: {bookmark.get('section_filter_count')}")
        else:
            lines.append("- Nenhum bookmark identificado.")

        lines.append("")
        lines.append("### Filtros do relatório")
        lines.append("")
        if self.metadata["report_filters"]:
            for item in self.metadata["report_filters"]:
                lines.append(
                    f"- {item.get('field') or '-'} | tipo: `{item.get('type') or '-'}`"
                )
        else:
            lines.append("- Nenhum filtro de relatório identificado.")
        lines.append("")
        return lines

    def _build_model_markdown(self) -> list[str]:
        lines = ["## Modelo Semântico", ""]
        model = self.metadata["model"]
        if not model.get("tables"):
            for note in model.get("notes", []):
                lines.append(f"- {note}")
            lines.append("")
            return lines
        for note in model.get("notes", []):
            lines.append(f"- {note}")
        if model.get("notes"):
            lines.append("")
        for table in model["tables"]:
            lines.append(f"### {table.get('name') or '(sem nome)'}")
            lines.append("")
            lines.append(f"- Oculta: `{'sim' if table.get('hidden') else 'não'}`")
            lines.append(f"- Colunas: {table.get('column_count')}")
            lines.append(f"- Medidas: {table.get('measure_count')}")
            if table.get("columns"):
                lines.append("- Colunas:")
                for column in table["columns"]:
                    lines.append(
                        f"  - {column.get('name')} | tipo: `{column.get('data_type') or '-'}` | oculta: `{column.get('hidden')}`"
                    )
            if table.get("measures"):
                lines.append("- Medidas:")
                for measure in table["measures"]:
                    lines.append(f"  - {measure.get('name')}")
                    lines.append("```dax")
                    lines.append(measure.get("expression") or "-")
                    lines.append("```")
            lines.append("")
        if model.get("relationships"):
            lines.append("### Relacionamentos")
            lines.append("")
            for relationship in model["relationships"]:
                lines.append(
                    f"- {relationship.get('from_table')}.{relationship.get('from_column')} -> "
                    f"{relationship.get('to_table')}.{relationship.get('to_column')}"
                )
            lines.append("")
        return lines

    def _build_visual_tokens_markdown(self) -> list[str]:
        lines = ["## Tokens Visuais e Tema", ""]
        colors = self.metadata["themes"].get("colors", [])
        fonts = self.metadata["themes"].get("fonts", [])
        lines.append(f"- Cores do tema: {len(colors)}")
        lines.append(f"- Fontes do tema: {len(fonts)}")
        lines.append(f"- Visuais customizados: {len(self.metadata['custom_visuals'])}")
        if colors:
            lines.append("")
            lines.append("### Cores")
            lines.append("")
            for item in colors:
                lines.append(f"- `{item}`")
        if fonts:
            lines.append("")
            lines.append("### Fontes")
            lines.append("")
            for item in fonts:
                lines.append(f"- `{item}`")
        if self.metadata["custom_visuals"]:
            lines.append("")
            lines.append("### Visuais Customizados")
            lines.append("")
            for item in self.metadata["custom_visuals"]:
                lines.append(f"- `{item}`")
        lines.append("")
        return lines

    def _build_query_artifacts_markdown(self) -> list[str]:
        lines = ["## Artefatos de Consulta", ""]
        if not self.metadata["queries"]:
            lines.append("- Nenhum artefato textual de consulta foi extraído.")
            lines.append("")
            return lines
        for artifact in self.metadata["queries"]:
            lines.append(f"### {artifact.get('member_name')}")
            lines.append("")
            lines.append("```text")
            lines.append(artifact.get("preview") or "-")
            lines.append("```")
            lines.append("")
        return lines

    def _build_mashup_markdown(self) -> list[str]:
        lines = ["## DataMashup", ""]
        mashup = self.metadata["mashup"]
        lines.append(f"- Presente no pacote: `{'sim' if mashup.get('present') else 'não'}`")
        lines.append(f"- Queries extraídas: {len(mashup.get('queries', []))}")
        lines.append(f"- Parâmetros extraídos: {len(mashup.get('parameters', []))}")
        lines.append(f"- Funções extraídas: {len(mashup.get('functions', []))}")
        for note in mashup.get("notes", []):
            lines.append(f"- {note}")
        if mashup.get("queries"):
            lines.append("")
            lines.append("### Queries M")
            lines.append("")
            for query in mashup.get("queries", [])[:10]:
                lines.append(f"#### {query.get('name') or '(sem nome)'}")
                lines.append("")
                lines.append("```powerquery")
                lines.append(query.get("m_code") or "-")
                lines.append("```")
                lines.append("")
        return lines

    def _doc_paragraph(self, text: str, style: str = "body", level: int = 0, mono: bool = False) -> dict[str, Any]:
        return {"type": "paragraph", "text": text, "style": style, "level": level, "mono": mono}

    def _doc_bullet(self, text: str, level: int = 0, mono: bool = False) -> dict[str, Any]:
        return {"type": "bullet", "text": text, "level": level, "mono": mono}

    def _doc_code_block(self, text: str, level: int = 0) -> dict[str, Any]:
        return {"type": "code", "text": text, "level": level, "mono": True}

    def _doc_list_block(
        self,
        label: str,
        values: list[Any],
        level: int = 0,
        empty_text: str = "-",
        mono: bool = False,
    ) -> list[dict[str, Any]]:
        parts = [self._doc_bullet(f"{label}:", level=level)]
        if values:
            for value in values:
                parts.append(self._doc_bullet(str(value), level=level + 1, mono=mono))
        else:
            parts.append(self._doc_bullet(empty_text, level=level + 1, mono=mono))
        return parts

    def _render_rtf_block(self, block: dict[str, Any]) -> str:
        if block["type"] == "paragraph":
            return self._rtf_paragraph(
                block["text"],
                style=block.get("style", "body"),
                level=block.get("level", 0),
                mono=block.get("mono", False),
            )
        if block["type"] == "bullet":
            return self._rtf_bullet(
                block["text"],
                level=block.get("level", 0),
                mono=block.get("mono", False),
            )
        return self._rtf_code_block(block["text"], level=block.get("level", 0))

    def _configure_docx_document(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        self._set_docx_font_family(normal_style.font, DOCX_BODY_FONT_NAME)
        normal_style.font.size = Pt(10)
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.line_spacing = 1.0
        normal_paragraph_format.space_before = Pt(0)
        normal_paragraph_format.space_after = Pt(6)
        document.core_properties.title = f"Documentação do Relatório Power BI - {self.source_path.name}"
        self._configure_docx_header(document)
        self._configure_docx_footer(document)

    def _append_docx_block(self, document: Document, block: dict[str, Any]) -> None:
        block_type = block["type"]
        if block_type == "bullet":
            paragraph = document.add_paragraph()
            self._apply_docx_paragraph_format(paragraph)
            level = block.get("level", 0)
            paragraph.paragraph_format.left_indent = Pt(18 * max(level, 0))
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            run = paragraph.add_run(f"• {block['text']}")
            self._apply_docx_run_style(run, mono=block.get("mono", False), bold=False, italic=False, font_size=10)
            return
        if block_type == "code":
            paragraph = document.add_paragraph()
            self._apply_docx_paragraph_format(paragraph)
            paragraph.paragraph_format.left_indent = Pt(18 * max(block.get("level", 0), 0))
            run = paragraph.add_run(block["text"])
            self._apply_docx_run_style(run, mono=True, bold=False, italic=False, font_size=10)
            return

        paragraph = document.add_paragraph()
        self._apply_docx_paragraph_format(paragraph)
        style = block.get("style", "body")
        level = block.get("level", 0)
        heading_style = self._docx_heading_style_for_block(style)
        if heading_style:
            paragraph.style = heading_style
        paragraph.paragraph_format.left_indent = Pt(18 * max(level, 0))
        font_size = 10
        bold = False
        italic = False
        if style == "title":
            font_size = 16
            bold = True
        elif style == "subtitle":
            font_size = 10
            italic = True
        elif style == "section":
            font_size = 14
            bold = True
            paragraph.paragraph_format.left_indent = Cm(-0.8)
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "subsection":
            font_size = 12
            bold = True
            paragraph.paragraph_format.left_indent = Cm(-0.7)
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "subsubsection":
            font_size = 11
            bold = True
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "body_bold":
            font_size = 10.5
            bold = True
        run = paragraph.add_run(block["text"])
        self._apply_docx_run_style(run, mono=block.get("mono", False), bold=bold, italic=italic, font_size=font_size)

    def _apply_docx_run_style(
        self,
        run: Any,
        mono: bool,
        bold: bool,
        italic: bool,
        font_size: float,
    ) -> None:
        run.bold = bold
        run.italic = italic
        self._set_docx_font_family(run.font, DOCX_MONO_FONT_NAME if mono else DOCX_BODY_FONT_NAME)
        run.font.size = Pt(font_size)

    def _set_docx_font_family(self, font: Any, font_name: str) -> None:
        font.name = font_name
        if getattr(font, "element", None) is None:
            return
        r_pr = font.element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), font_name)

    def _apply_docx_paragraph_format(self, paragraph: Any) -> None:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.0
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)

    def _configure_docx_header(self, document: Document) -> None:
        for section in document.sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.clear()
            self._apply_docx_paragraph_format(paragraph)
            paragraph.alignment = 1
            self._apply_docx_bottom_border(paragraph)
            run = paragraph.add_run("Documentação Power BI")
            self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=8)

    def _configure_docx_footer(self, document: Document) -> None:
        workbook_name = self.base_name
        for section in document.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.clear()
            self._apply_docx_paragraph_format(paragraph)
            usable_width = section.page_width - section.left_margin - section.right_margin
            paragraph.paragraph_format.tab_stops.add_tab_stop(usable_width, alignment=WD_TAB_ALIGNMENT.RIGHT)
            self._apply_docx_top_border(paragraph)

            left_run = paragraph.add_run(f"Documentação Power BI pasta de Trabalho {workbook_name}")
            self._apply_docx_run_style(left_run, mono=False, bold=False, italic=False, font_size=8)
            tab_run = paragraph.add_run("\t")
            self._apply_docx_run_style(tab_run, mono=False, bold=False, italic=False, font_size=8)
            self._append_docx_field(paragraph, "PAGE", font_size=8)
            middle_run = paragraph.add_run(" de ")
            self._apply_docx_run_style(middle_run, mono=False, bold=False, italic=False, font_size=8)
            self._append_docx_field(paragraph, "NUMPAGES", font_size=8)

    def _apply_docx_top_border(self, paragraph: Any) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        top = p_bdr.find(qn("w:top"))
        if top is None:
            top = OxmlElement("w:top")
            p_bdr.append(top)
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "BFBFBF")

    def _apply_docx_bottom_border(self, paragraph: Any) -> None:
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = p_bdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_bdr.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")

    def _append_docx_field(self, paragraph: Any, field_name: str, font_size: float) -> None:
        run = paragraph.add_run()
        self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=font_size)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_name
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(placeholder)
        run._r.append(end)

    def _append_docx_toc(self, document: Document) -> None:
        toc_title = document.add_paragraph()
        self._apply_docx_paragraph_format(toc_title)
        title_run = toc_title.add_run("Índice")
        self._apply_docx_run_style(title_run, mono=False, bold=True, italic=False, font_size=12)

        toc_paragraph = document.add_paragraph()
        self._apply_docx_paragraph_format(toc_paragraph)
        self._append_docx_complex_field(
            toc_paragraph,
            'TOC \\o "1-3" \\h \\z \\u',
            placeholder_text="Atualize o índice no Word se necessário.",
            font_size=10,
        )

    def _append_docx_complex_field(
        self,
        paragraph: Any,
        instruction: str,
        placeholder_text: str,
        font_size: float,
    ) -> None:
        run = paragraph.add_run()
        self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=font_size)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        placeholder = OxmlElement("w:t")
        placeholder.text = placeholder_text
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(placeholder)
        run._r.append(end)

    def _docx_heading_style_for_block(self, style: str) -> str | None:
        if style == "section":
            return "Heading 1"
        if style == "subsection":
            return "Heading 2"
        if style == "subsubsection":
            return "Heading 3"
        return None

    def _rtf_paragraph(self, text: str, style: str = "body", level: int = 0, mono: bool = False) -> str:
        escaped = self._rtf_escape(text)
        indent = 360 * max(level, 0)
        if style == "title":
            return rf"\pard\sa240\sb160\f0\fs32\b {escaped}\b0\par"
        if style == "subtitle":
            return rf"\pard\sa120\sb160\i\f0\fs20 {escaped}\i0\par"
        if style == "section":
            return rf"\pard\sa200\sb120\li{indent}\f0\fs28\b {escaped}\b0\par"
        if style == "subsection":
            return rf"\pard\sa160\sb80\li{indent}\f0\fs24\b {escaped}\b0\par"
        if style == "subsubsection":
            return rf"\pard\sa120\sb60\li{indent}\f0\fs22\b {escaped}\b0\par"
        if style == "body_bold":
            return rf"\pard\sa80\sb40\li{indent}\f0\fs21\b {escaped}\b0\par"
        font_id = 1 if mono else 0
        return rf"\pard\sa60\sb20\li{indent}\f{font_id}\fs20 {escaped}\par"

    def _rtf_bullet(self, text: str, level: int = 0, mono: bool = False) -> str:
        indent = 360 * max(level, 0)
        hanging = 180
        escaped = self._rtf_escape(text)
        font_id = 1 if mono else 0
        return rf"\pard\sa40\sb20\li{indent}\fi-{hanging}\f{font_id}\fs20 \'95\tab {escaped}\par"

    def _rtf_code_block(self, text: str, level: int = 0) -> str:
        indent = 360 * max(level, 0)
        escaped = self._rtf_escape(text)
        return rf"\pard\sa60\sb40\li{indent}\f1\fs20 {escaped}\par"

    def _rtf_escape(self, text: Any) -> str:
        value = "" if text is None else str(text)
        output: list[str] = []
        for char in value:
            if char == "\\":
                output.append(r"\\")
            elif char == "{":
                output.append(r"\{")
            elif char == "}":
                output.append(r"\}")
            elif char == "\n":
                output.append(r"\line ")
            elif char == "\t":
                output.append(r"\tab ")
            else:
                codepoint = ord(char)
                if 32 <= codepoint <= 126:
                    output.append(char)
                else:
                    signed = codepoint if codepoint <= 32767 else codepoint - 65536
                    output.append(rf"\u{signed}?")
        return "".join(output)


def load_config(config_path: str | Path | None = None) -> dict[str, Any]:
    config_file = Path(config_path) if config_path is not None else PROJECT_ROOT / "config" / "config.json"
    if not config_file.exists():
        raise FileNotFoundError(
            f"Nenhum caminho informado e o arquivo '{config_file}' não foi encontrado."
        )
    try:
        config_data = json.loads(config_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"O arquivo '{config_file}' não contém um JSON válido.") from exc
    if not isinstance(config_data, dict):
        raise ValueError(f"O arquivo '{config_file}' deve conter um objeto JSON na raiz.")
    return config_data


def load_powerbi_path_from_config(config_path: str | Path | None = None) -> str:
    config_data = load_config(config_path)
    powerbi_path = config_data.get("powerbi_path")
    if not powerbi_path:
        config_file = Path(config_path) if config_path is not None else PROJECT_ROOT / "config" / "config.json"
        raise ValueError(f"O arquivo '{config_file}' não possui a chave obrigatória 'powerbi_path'.")
    if not isinstance(powerbi_path, str):
        raise ValueError("A chave 'powerbi_path' do config.json deve ser uma string.")
    if not powerbi_path.strip():
        raise ValueError("A chave 'powerbi_path' do config.json não pode estar vazia.")
    return powerbi_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Gera documentação de relatórios Power BI (.pbix)."
    )
    parser.add_argument(
        "filepath",
        nargs="?",
        help="Caminho completo do arquivo Power BI (.pbix).",
    )
    parser.add_argument(
        "--format",
        dest="output_format",
        choices=["all", "markdown", "json", "excel", "rtf", "docx"],
        default="all",
        help="Formato principal da documentação a gerar. Em `markdown`, o script gera `.md` e `.rtf`. O mapa de estrutura do pacote é sempre gerado; em `docx`, o relatório Word é exportado diretamente.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    try:
        file_path = args.filepath or load_powerbi_path_from_config()
        documenter = PowerBIDoc(file_path, output_format=args.output_format)
        written_files = documenter.write_outputs()
        print(f"Arquivo fonte: {documenter.source_path}")
        print(f"Diretório de saída: {documenter.output_dir}")
        print("Arquivos gerados:")
        for path in written_files:
            print(f"- {path}")
    except (FileNotFoundError, ValueError, zipfile.BadZipFile) as exc:
        print(f"Erro: {exc}")
        print(
            "Uso: python PowerBI_doc.py <caminho_do_arquivo.pbix> "
            "[--format all|markdown|json|excel|rtf|docx]"
        )
        sys.exit(1)


if __name__ == "__main__":
    main()

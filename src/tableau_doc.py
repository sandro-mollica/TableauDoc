# -*- coding: utf-8 -*-
"""
Gera documentação técnica de workbooks Tableau (`.twb` e `.twbx`).

Fluxo principal:
- lê um único workbook Tableau por execução;
- extrai o XML do workbook e, quando necessário, o conteúdo do pacote `.twbx`;
- monta um conjunto consolidado de metadados sobre fontes, planilhas, dashboards,
  parâmetros, cálculos e elementos visuais;
- gera artefatos de documentação em Markdown, RTF, DOCX, JSON e/ou Excel;
- opcionalmente complementa a leitura com um ou mais `.tdsx` externos definidos no
  `config/config.json`.

Exemplos:
- `python3 Tableau_doc.py --format all`
- `python3 Tableau_doc.py /caminho/arquivo.twbx --format markdown`
- `python3 Tableau_doc.py /caminho/arquivo.twbx --format json`
- `python3 Tableau_doc.py /caminho/arquivo.twbx --format excel`
- `python3 Tableau_doc.py /caminho/arquivo.twbx --format rtf`
- `python3 Tableau_doc.py /caminho/arquivo.twbx --format docx`
"""

from __future__ import annotations

import argparse
import base64
import json
import re
import shutil
import sys
import zipfile
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
import xml.etree.ElementTree as ET

import pandas as pd
from docx import Document
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

SCRIPT_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SCRIPT_DIR.parent
DEFAULT_OUTPUT_ROOT = PROJECT_ROOT / "data"
TEMPORARY_OUTPUT_NAMES = {
    "package_contents",
    ".tmp",
    "tmp",
    "temp",
}

# Fontes do documento RTF. Ajuste aqui caso queira personalizar a tipografia.
RTF_BODY_FONT_NAME = "Arial"
RTF_MONO_FONT_NAME = "Courier New"
DOCX_BODY_FONT_NAME = RTF_BODY_FONT_NAME
DOCX_MONO_FONT_NAME = RTF_MONO_FONT_NAME
TABLEAU_IMPLICIT_FONT_LABEL = "Fonte padrão do Tableau (não explicitada no XML)"
SHOW_WORKSHEET_SHELF_COLUMNS_IN_RTF = False
SHOW_FILTER_GROUP_VALUES_IN_RTF = False
ENABLE_EXTERNAL_TDSX_LOOKUP = True

# PRINTS_DE_PROGRESSO: comente esta constante e as chamadas de `_log_progress`
# se quiser silenciar rapidamente as mensagens de acompanhamento.
ENABLE_PROGRESS_PRINTS = True


def sanitize_filename(value: str) -> str:
    """Converte um texto em nome de arquivo seguro, preservando legibilidade."""
    safe = re.sub(r"[^\w\-. ]+", "_", value, flags=re.UNICODE).strip()
    return safe or "arquivo"


def decode_tableau_text(value: str | None) -> str | None:
    """Normaliza entidades XML frequentes nas fórmulas do Tableau."""
    if value is None:
        return None

    return (
        value.replace("&quot;", '"')
        .replace("&apos;", "'")
        .replace("&#10;", "\n")
        .replace("&#13;", "\r")
        .replace("&lt;", "<")
        .replace("&gt;", ">")
        .replace("&amp;", "&")
    )


def clean_brackets(value: str | None) -> str | None:
    """Remove colchetes do nome interno do Tableau, quando existirem."""
    if value is None:
        return None
    return value.replace("[", "").replace("]", "")


def clean_display_label(value: str | None) -> str | None:
    """Limpa ruídos visuais comuns em labels extraídos do XML."""
    if value is None:
        return None
    text = value.replace("\xa0", " ").strip()
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\s+([|,:;])", r"\1", text)
    text = re.sub(r"^\:+", "", text)
    return text.strip()


def is_color_like(value: str | None) -> bool:
    """Retorna True apenas para valores que parecem cores ou nomes de paleta úteis."""
    if not value:
        return False
    text = clean_display_label(value) or ""
    lower = text.lower()
    if re.fullmatch(r"#(?:[0-9a-fA-F]{3,8})", text):
        return True
    if lower.startswith(("rgb(", "rgba(", "hsl(", "hsla(")):
        return True
    if "palette" in lower:
        return True
    named_colors = {
        "black", "white", "red", "green", "blue", "yellow", "orange", "gray", "grey",
        "brown", "pink", "cyan", "magenta", "teal", "navy", "maroon", "olive", "lime",
        "silver", "gold", "beige", "transparent",
    }
    return lower in named_colors


def unique_ordered(values: list[Any]) -> list[Any]:
    """Remove duplicados preservando a ordem de aparição."""
    output = []
    seen = set()
    for value in values:
        marker = json.dumps(value, ensure_ascii=False, sort_keys=True) if isinstance(value, (dict, list)) else value
        if marker in seen:
            continue
        seen.add(marker)
        output.append(value)
    return output


def compact_json(value: Any) -> str:
    """Serializa estruturas aninhadas em uma linha legível."""
    if value is None:
        return "-"
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, sort_keys=True)
    return str(value)


def format_yes_no(value: Any) -> str:
    """Normaliza flags do XML para `sim` ou `não`."""
    if isinstance(value, bool):
        return "sim" if value else "não"
    if value is None:
        return "não"
    text = str(value).strip().lower()
    return "sim" if text in {"true", "1", "yes", "sim"} else "não"


def _log_progress(message: str) -> None:
    """Imprime mensagens destacadas de progresso durante a execução."""
    if not ENABLE_PROGRESS_PRINTS:
        return
    print(f"[TableauDoc][progresso] {message}")


def normalize_whitespace(value: str | None) -> str | None:
    """Compacta quebras e espaços sem destruir o conteúdo técnico."""
    if value is None:
        return None
    normalized = value.replace("\r\n", "\n").replace("\r", "\n").strip()
    return normalized or None


def normalize_lookup_token(value: str | None) -> str | None:
    """Normaliza nomes para comparação flexível entre datasource e arquivo externo."""
    if not value:
        return None
    text = value.strip().casefold()
    text = re.sub(r"\s+", "_", text)
    text = re.sub(r"[^\w.-]+", "_", text)
    return text.strip("_") or None


def element_path_with_indices(root: ET.Element, target: ET.Element) -> str:
    """
    Monta um caminho XPath simples e legível até um elemento.
    O `xml.etree` não mantém ponteiro para o pai, então percorremos a árvore.
    """
    target_id = id(target)

    def walk(node: ET.Element, trail: list[str]) -> list[str] | None:
        if id(node) == target_id:
            return trail + [node.tag]

        child_counts: Counter[str] = Counter()
        for child in list(node):
            child_counts[child.tag] += 1
            siblings_same_tag = [item for item in list(node) if item.tag == child.tag]
            if len(siblings_same_tag) > 1:
                index = siblings_same_tag.index(child) + 1
                child_name = f"{child.tag}[{index}]"
            else:
                child_name = child.tag

            result = walk(child, trail + [node.tag, child_name])
            if result is not None:
                return result
        return None

    if id(root) == target_id:
        return f"/{root.tag}"

    result = walk(root, [])
    if result is None:
        return ""

    normalized = []
    skip_next_root = False
    for idx, part in enumerate(result):
        if idx == 0 and part == root.tag:
            normalized.append(part)
            skip_next_root = True
            continue
        if skip_next_root and part.startswith(f"{root.tag}/"):
            continue
        normalized.append(part)
    return "/" + "/".join(normalized)


class TableauDoc:
    """Carrega, extrai e documenta um workbook Tableau."""

    MAP_DEFINITIONS = [
        {
            "section": "workbook",
            "label": "Raiz do workbook",
            "xpath": ".",
            "json_path": "$.workbook",
            "description": "Metadados gerais do workbook, versão e build.",
        },
        {
            "section": "preferences",
            "label": "Preferências e paletas",
            "xpath": "./preferences",
            "json_path": "$.preferences",
            "description": "Preferências globais, incluindo paletas de cor customizadas.",
        },
        {
            "section": "styles",
            "label": "Estilos globais",
            "xpath": "./style/style-rule",
            "json_path": "$.styles.global_rules[*]",
            "description": "Regras de estilo globais, como fonte padrão e títulos.",
        },
        {
            "section": "datasources",
            "label": "Datasources",
            "xpath": "./datasources/datasource",
            "json_path": "$.datasources[*]",
            "description": "Fontes de dados, conexões, relações, colunas e metadados.",
        },
        {
            "section": "parameters",
            "label": "Parâmetros",
            "xpath": "./datasources/datasource[@name='Parameters']/column",
            "json_path": "$.parameters[*]",
            "description": "Parâmetros definidos no workbook.",
        },
        {
            "section": "calculations",
            "label": "Campos calculados",
            "xpath": ".//column[calculation] | .//calculation[@column]",
            "json_path": "$.calculations[*]",
            "description": "Campos calculados, fórmulas e dependências.",
        },
        {
            "section": "worksheets",
            "label": "Worksheets",
            "xpath": "./worksheets/worksheet",
            "json_path": "$.worksheets[*]",
            "description": "Planilhas, filtros, encodings, marks e estilos locais.",
        },
        {
            "section": "dashboards",
            "label": "Dashboards",
            "xpath": "./dashboards/dashboard",
            "json_path": "$.dashboards[*]",
            "description": "Painéis, zonas, objetos, filtros expostos e layout.",
        },
        {
            "section": "stories",
            "label": "Stories",
            "xpath": "./stories/story",
            "json_path": "$.stories[*]",
            "description": "Histórias do Tableau, quando presentes.",
        },
        {
            "section": "windows",
            "label": "Windows e cards",
            "xpath": "./windows/window",
            "json_path": "$.windows[*]",
            "description": "Estado da UI, cartões de filtro, cor, marks e legendas.",
        },
        {
            "section": "thumbnails",
            "label": "Miniaturas",
            "xpath": "./thumbnails/thumbnail",
            "json_path": "$.thumbnails[*]",
            "description": "Pré-visualizações embutidas em base64.",
        },
    ]

    def __init__(
        self,
        source_path: str | Path,
        output_format: str = "all",
        external_tdsx_paths: list[str | Path] | None = None,
    ) -> None:
        self.source_path = Path(source_path).expanduser().resolve()
        if not self.source_path.exists():
            raise FileNotFoundError(f"Arquivo não encontrado: {self.source_path}")
        if self.source_path.suffix.lower() not in {".twb", ".twbx"}:
            raise ValueError("O arquivo informado deve ter extensão .twb ou .twbx.")

        self.output_format = output_format.lower()
        self.base_name = self.source_path.stem
        self.output_dir = DEFAULT_OUTPUT_ROOT / self.base_name
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.generated_at = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        self.external_tdsx_paths = []
        for path in external_tdsx_paths or []:
            candidate = Path(path).expanduser()
            if not candidate.exists():
                continue
            resolved = candidate.resolve()
            if resolved not in self.external_tdsx_paths:
                self.external_tdsx_paths.append(resolved)
        _log_progress(
            f"Iniciando documentação de `{self.source_path.name}` com formato `{self.output_format}`."
        )
        _log_progress(f"Diretório de saída preparado em `{self.output_dir}`.")
        if self.external_tdsx_paths:
            _log_progress(
                "Busca externa de `.tdsx` habilitada em: "
                + ", ".join(f"`{path}`" for path in self.external_tdsx_paths)
            )

        self.workbook_name_in_package: str | None = None
        self.workbook_xml_bytes: bytes | None = None
        self.package_manifest: list[dict[str, Any]] = []
        self.assets_written: list[str] = []
        self.caption_lookup: dict[str, str] = {}
        self.caption_lookup_by_clean: dict[str, str] = {}
        self.datasource_caption_lookup: dict[str, str] = {}
        self.datasource_caption_lookup_by_clean: dict[str, str] = {}

        self.root = self._load_root_and_extract_contents()
        self.metadata = self._build_metadata()
        self._augment_summary_metrics()

    def _load_root_and_extract_contents(self) -> ET.Element:
        """Lê o XML principal e salva o conteúdo bruto do workbook no diretório de saída."""
        suffix = self.source_path.suffix.lower()
        _log_progress(f"Lendo arquivo de origem `{self.source_path}`.")

        if suffix == ".twb":
            _log_progress("Arquivo `.twb` detectado; copiando workbook para a pasta de trabalho.")
            self.workbook_name_in_package = self.source_path.name
            self.workbook_xml_bytes = self.source_path.read_bytes()
            extracted_dir = self.output_dir / "package_contents"
            if extracted_dir.exists():
                shutil.rmtree(extracted_dir)
            extracted_dir.mkdir(parents=True, exist_ok=True)
            shutil.copy2(self.source_path, extracted_dir / self.source_path.name)
            self.package_manifest.append(
                {
                    "path": self.source_path.name,
                    "kind": "workbook",
                    "size_bytes": self.source_path.stat().st_size,
                }
            )
            _log_progress("XML principal carregado a partir do arquivo `.twb`.")
            return ET.fromstring(self.workbook_xml_bytes)

        _log_progress("Arquivo `.twbx` detectado; extraindo conteúdo do pacote.")
        with zipfile.ZipFile(self.source_path, "r") as archive:
            extracted_dir = self.output_dir / "package_contents"
            if extracted_dir.exists():
                shutil.rmtree(extracted_dir)
            extracted_dir.mkdir(parents=True, exist_ok=True)
            for member in archive.infolist():
                destination = extracted_dir / member.filename
                if member.is_dir():
                    destination.mkdir(parents=True, exist_ok=True)
                    continue
                if member.filename.lower().endswith(".hyper"):
                    continue
                destination.parent.mkdir(parents=True, exist_ok=True)
                with archive.open(member, "r") as src, destination.open("wb") as dst:
                    shutil.copyfileobj(src, dst)

            for info in archive.infolist():
                self.package_manifest.append(
                    {
                        "path": info.filename,
                        "kind": self._classify_package_member(info.filename),
                        "size_bytes": info.file_size,
                    }
                )

            twb_names = sorted(
                name for name in archive.namelist() if name.lower().endswith(".twb")
            )
            if not twb_names:
                raise ValueError("O arquivo .twbx não contém um workbook .twb.")

            self.workbook_name_in_package = twb_names[0]
            self.workbook_xml_bytes = archive.read(twb_names[0])
            _log_progress(
                f"Pacote extraído com sucesso; workbook interno identificado como `{self.workbook_name_in_package}`."
            )
            return ET.fromstring(self.workbook_xml_bytes)

    def _classify_package_member(self, package_path: str) -> str:
        """Classifica membros do pacote por tipo básico para documentação."""
        lower = package_path.lower()
        if lower.endswith(".twb"):
            return "workbook"
        if lower.endswith(".hyper"):
            return "extract"
        if lower.endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg")):
            return "image"
        if lower.endswith((".csv", ".xlsx", ".xls", ".txt", ".json", ".xml")):
            return "data_or_support"
        return "other"

    def _build_metadata(self) -> dict[str, Any]:
        """Consolida o workbook em uma estrutura única para saída JSON/MD/Excel."""
        _log_progress("Iniciando extração de metadados do workbook.")
        workbook_info = {
            "source_file_name": self.source_path.name,
            "source_file_path": str(self.source_path),
            "source_file_type": self.source_path.suffix.lower(),
            "source_file_last_modified": datetime.fromtimestamp(
                self.source_path.stat().st_mtime
            ).isoformat(timespec="seconds"),
            "output_directory": str(self.output_dir),
            "package_workbook_name": self.workbook_name_in_package,
            "attributes": dict(self.root.attrib),
            "repository_location": self._extract_repository_location(self.root),
        }

        preferences = self._extract_preferences()
        _log_progress("Preferências globais extraídas.")
        styles = self._extract_global_styles()
        _log_progress("Estilos globais extraídos.")
        datasources = self._extract_datasources()
        _log_progress(f"Fontes de dados extraídas: {len(datasources)}.")
        self._enrich_datasources_with_external_tdsx(datasources)
        self._build_caption_lookup(datasources)
        parameters = self._extract_parameters(datasources)
        _log_progress(f"Parâmetros identificados: {len(parameters)}.")
        calculations = self._extract_calculations(datasources, parameters)
        _log_progress(f"Campos calculados identificados: {len(calculations)}.")
        worksheets = self._extract_worksheets()
        _log_progress(f"Planilhas extraídas: {len(worksheets)}.")
        dashboards = self._extract_dashboards(worksheets)
        _log_progress(f"Dashboards extraídos: {len(dashboards)}.")
        self._enrich_usage(parameters, calculations, worksheets, dashboards)
        self._enrich_datasource_field_usage(datasources, worksheets, calculations, parameters)
        _log_progress("Relações de uso entre parâmetros, cálculos, planilhas e dashboards enriquecidas.")
        stories = self._extract_stories()
        windows = self._extract_windows()
        thumbnails = self._extract_thumbnails()
        colors = self._extract_visual_tokens(token_type="color")
        fonts = self._extract_visual_tokens(token_type="font")
        _log_progress("Tokens visuais, stories, janelas e miniaturas processados.")

        return {
            "workbook": workbook_info,
            "package_manifest": self.package_manifest,
            "preferences": preferences,
            "styles": styles,
            "datasources": datasources,
            "parameters": parameters,
            "calculations": calculations,
            "worksheets": worksheets,
            "dashboards": dashboards,
            "stories": stories,
            "windows": windows,
            "thumbnails": thumbnails,
            "visual_tokens": {
                "colors": colors,
                "fonts": fonts,
            },
            "summary": {
                "datasource_count": len(datasources),
                "parameter_count": len(parameters),
                "calculation_count": len(calculations),
                "worksheet_count": len(worksheets),
                "dashboard_count": len(dashboards),
                "story_count": len(stories),
                "window_count": len(windows),
                "thumbnail_count": len(thumbnails),
                "package_member_count": len(self.package_manifest),
            },
        }

    def _augment_summary_metrics(self) -> None:
        """Complementa o resumo com contadores derivados usados nos relatórios."""
        summary = self.metadata.setdefault("summary", {})
        unused_objects = self._collect_unused_objects()
        summary["worksheet_independent_count"] = self._count_independent_worksheets()
        summary["unused_calculation_count"] = len(unused_objects["calculations"])
        summary["unused_parameter_count"] = len(unused_objects["parameters"])

    def _count_independent_worksheets(self) -> int:
        """Conta worksheets que não aparecem associadas a nenhum dashboard."""
        associated_worksheets = {
            worksheet_name
            for dashboard in self.metadata.get("dashboards", [])
            for worksheet_name in dashboard.get("worksheet_members", [])
        }
        return sum(
            1
            for worksheet in self.metadata.get("worksheets", [])
            if worksheet.get("name") not in associated_worksheets
        )

    def _build_caption_lookup(self, datasources: list[dict[str, Any]]) -> None:
        """Monta um mapa global entre nomes internos e captions amigáveis."""
        self.caption_lookup = {}
        self.caption_lookup_by_clean = {}
        self.datasource_caption_lookup = {}
        self.datasource_caption_lookup_by_clean = {}
        for datasource in datasources:
            datasource_name = datasource.get("name")
            datasource_caption = datasource.get("caption") or datasource_name
            if datasource_name and datasource_caption:
                self.datasource_caption_lookup[datasource_name] = datasource_caption
                cleaned_ds_name = normalize_lookup_token(clean_brackets(datasource_name))
                if cleaned_ds_name:
                    self.datasource_caption_lookup_by_clean[cleaned_ds_name] = datasource_caption
            if datasource_caption:
                cleaned_ds_caption = normalize_lookup_token(clean_brackets(datasource_caption))
                if cleaned_ds_caption:
                    self.datasource_caption_lookup_by_clean[cleaned_ds_caption] = datasource_caption
            for column in datasource.get("columns", []):
                name = column.get("name")
                caption = column.get("caption") or clean_brackets(name)
                if name and caption:
                    self.caption_lookup[name] = caption
                    self.caption_lookup_by_clean[clean_brackets(name) or ""] = caption
            for record in datasource.get("metadata_records", []):
                local_name = record.get("local_name")
                caption = record.get("caption") or clean_brackets(local_name)
                if local_name and caption:
                    self.caption_lookup[local_name] = caption
                    self.caption_lookup_by_clean[clean_brackets(local_name) or ""] = caption

    def _extract_repository_location(self, element: ET.Element) -> dict[str, Any] | None:
        repository = element.find("./repository-location")
        return dict(repository.attrib) if repository is not None else None

    def _extract_preferences(self) -> dict[str, Any]:
        preferences = self.root.find("./preferences")
        if preferences is None:
            return {"preferences": [], "color_palettes": []}

        pref_rows = [dict(pref.attrib) for pref in preferences.findall("./preference")]
        palettes = []
        for palette in preferences.findall("./color-palette"):
            palettes.append(
                {
                    "name": palette.get("name"),
                    "type": palette.get("type"),
                    "custom": palette.get("custom"),
                    "colors": [color.text for color in palette.findall("./color") if color.text],
                }
            )
        return {"preferences": pref_rows, "color_palettes": palettes}

    def _extract_global_styles(self) -> dict[str, Any]:
        style = self.root.find("./style")
        return {"global_rules": self._parse_style_rules(style) if style is not None else []}

    def _extract_datasources(self) -> list[dict[str, Any]]:
        datasources = []
        for datasource in self.root.findall("./datasources/datasource"):
            columns = []
            for column in datasource.findall("./column"):
                columns.append(self._parse_column(column))

            metadata_records = []
            for record in datasource.findall(".//metadata-records/metadata-record"):
                metadata_records.append(self._parse_metadata_record(record))

            calculations = []
            for calc in datasource.findall(".//connection/calculations/calculation"):
                calculations.append(
                    {
                        "column": calc.get("column"),
                        "formula": decode_tableau_text(calc.get("formula")),
                    }
                )

            connections = []
            for connection in datasource.findall("./connection"):
                connections.append(self._parse_connection(connection))

            extracts = self._extract_extracts(datasource)
            for connection in connections:
                connection["mode"] = self._infer_connection_mode(connection, extracts)
                connection["hyper_paths"] = self._extract_hyper_paths(extracts)

            datasource_info = {
                "name": datasource.get("name"),
                "caption": datasource.get("caption"),
                "version": datasource.get("version"),
                "inline": datasource.get("inline"),
                "hasconnection": datasource.get("hasconnection"),
                "repository_location": self._extract_repository_location(datasource),
                "connections": connections,
                "extracts": extracts,
                "relationship_maps": self._extract_relationship_map_from_root(root=datasource, source_label="twb/twbx"),
                "columns": columns,
                "metadata_records": metadata_records,
                "connection_calculations": calculations,
                "aliases_enabled": datasource.find("./aliases") is not None and datasource.find("./aliases").get("enabled"),
                "object_count": len(datasource.findall("./object-graph/object")),
            }
            datasources.append(datasource_info)
        return datasources

    def _extract_extracts(self, datasource: ET.Element) -> list[dict[str, Any]]:
        """Extrai metadados básicos dos blocos de extração do datasource."""
        extracts: list[dict[str, Any]] = []
        for extract in datasource.findall("./extract"):
            extract_info = dict(extract.attrib)
            extract_connection = extract.find("./connection")
            extract_info["connection"] = dict(extract_connection.attrib) if extract_connection is not None else {}
            extracts.append(extract_info)
        return extracts

    def _infer_connection_mode(
        self,
        connection: dict[str, Any],
        extracts: list[dict[str, Any]],
    ) -> str:
        """Classifica a conexão como live, extração ou combinação de ambos."""
        attributes = connection.get("attributes") or {}
        connection_class = (attributes.get("class") or "").lower()

        if connection_class == "hyper":
            return "Extração"

        extract_connections = [item.get("connection", {}) for item in extracts]
        has_hyper_extract = any(
            (extract_connection.get("class") or "").lower() == "hyper"
            or str(extract_connection.get("dbname") or "").lower().endswith(".hyper")
            for extract_connection in extract_connections
        )
        has_enabled_extract = any(str(item.get("enabled")).lower() == "true" for item in extracts)
        has_disabled_extract = any(str(item.get("enabled")).lower() == "false" for item in extracts)

        if has_hyper_extract and has_enabled_extract:
            return "Live + extração habilitada"
        if has_hyper_extract and has_disabled_extract:
            return "Live + extração desabilitada"
        return "Live"

    def _extract_hyper_paths(self, extracts: list[dict[str, Any]]) -> list[str]:
        """Retorna os caminhos de arquivos .hyper associados aos blocos de extração."""
        paths = []
        for extract in extracts:
            connection = extract.get("connection", {}) or {}
            dbname = connection.get("dbname")
            if dbname and str(dbname).lower().endswith(".hyper"):
                paths.append(str(dbname))
        return unique_ordered(paths)

    def _enrich_datasources_with_external_tdsx(self, datasources: list[dict[str, Any]]) -> None:
        """Procura `.tdsx` externos e anexa SQL e relacionamentos quando disponíveis."""
        if not self.external_tdsx_paths:
            return

        for datasource in datasources:
            matches = self._find_matching_external_tdsx(datasource)
            datasource["external_tdsx_matches"] = [str(path) for path in matches]
            if not matches:
                continue

            _log_progress(
                f"Buscando SQL externo para datasource `{datasource.get('caption') or datasource.get('name')}`."
            )
            external_sql_relations = []
            external_relationship_maps = []
            for tdsx_path in matches:
                external_sql_relations.extend(self._extract_custom_sql_from_tdsx(tdsx_path))
                external_relationship_maps.extend(self._extract_relationship_map_from_tdsx(tdsx_path))

            datasource["external_custom_sql_relations"] = external_sql_relations
            datasource["external_relationship_maps"] = external_relationship_maps

    def _find_matching_external_tdsx(self, datasource: dict[str, Any]) -> list[Path]:
        """Encontra possíveis `.tdsx` externos com base no nome da fonte."""
        if datasource.get("name") == "Parameters" or not datasource.get("connections"):
            return []

        datasource_name = datasource.get("name") or ""
        datasource_name_base = datasource_name.split(".", 1)[0] if "." in datasource_name else datasource_name
        repo = datasource.get("repository_location") or {}
        repo_name_candidates = [
            repo.get("id"),
            repo.get("path"),
            repo.get("name"),
        ]
        tokens = {
            normalize_lookup_token(datasource.get("caption")),
            normalize_lookup_token(datasource_name),
            normalize_lookup_token(datasource_name_base),
            *(normalize_lookup_token(value) for value in repo_name_candidates),
        }
        tokens = {token for token in tokens if token}

        matches: list[Path] = []
        for root in self.external_tdsx_paths:
            candidates = [root] if root.is_file() and root.suffix.lower() == ".tdsx" else root.rglob("*.tdsx")
            for path in candidates:
                path_tokens = {
                    normalize_lookup_token(path.stem),
                    normalize_lookup_token(path.name),
                }
                path_tokens = {token for token in path_tokens if token}
                if tokens.intersection(path_tokens):
                    matches.append(path)
        matches = unique_ordered(matches)
        if matches:
            return matches

        all_candidates: list[Path] = []
        for root in self.external_tdsx_paths:
            all_candidates.extend(
                [root] if root.is_file() and root.suffix.lower() == ".tdsx" else list(root.rglob("*.tdsx"))
            )
        all_candidates = unique_ordered(all_candidates)
        if len(all_candidates) == 1 and len((datasource.get("connections") or [])) > 0:
            return all_candidates
        return []

    def _extract_custom_sql_from_tdsx(self, tdsx_path: Path) -> list[dict[str, Any]]:
        """Extrai blocos de SQL customizado de um `.tdsx` externo."""
        records: list[dict[str, Any]] = []
        try:
            with zipfile.ZipFile(tdsx_path, "r") as archive:
                tds_names = [name for name in archive.namelist() if name.lower().endswith(".tds")]
                for tds_name in tds_names:
                    root = ET.fromstring(archive.read(tds_name))
                    for relation in root.findall(".//relation"):
                        relation_text = normalize_whitespace("".join(relation.itertext()))
                        custom_sql = self._extract_relation_sql(relation, relation_text)
                        if not custom_sql:
                            continue
                        records.append(
                            {
                                "tdsx_path": str(tdsx_path),
                                "tds_path": tds_name,
                                "relation_name": relation.get("name"),
                                "relation_type": relation.get("type"),
                                "custom_sql": custom_sql,
                            }
                        )
        except (zipfile.BadZipFile, ET.ParseError, FileNotFoundError):
            return []
        return records

    def _extract_relationship_map_from_tdsx(self, tdsx_path: Path) -> list[dict[str, Any]]:
        """Extrai o mapa de relacionamentos entre tabelas a partir da tag `relationships`."""
        records: list[dict[str, Any]] = []
        try:
            with zipfile.ZipFile(tdsx_path, "r") as archive:
                tds_names = [name for name in archive.namelist() if name.lower().endswith(".tds")]
                for tds_name in tds_names:
                    root = ET.fromstring(archive.read(tds_name))
                    for record in self._extract_relationship_map_from_root(root=root, source_label="tdsx"):
                        record["tdsx_path"] = str(tdsx_path)
                        record["tds_path"] = tds_name
                        records.append(record)
        except (zipfile.BadZipFile, ET.ParseError, FileNotFoundError):
            return []
        return records

    def _extract_relationship_map_from_root(
        self,
        root: ET.Element,
        source_label: str,
    ) -> list[dict[str, Any]]:
        """Extrai relacionamentos de um bloco XML que contenha object-graph e relationships."""
        object_labels = {}
        for obj in root.findall(".//object-graph//object"):
            object_id = obj.get("id") or obj.get("object-id")
            if object_id:
                object_labels[object_id] = obj.get("caption") or obj.get("name") or object_id

        records: list[dict[str, Any]] = []
        for relationships in root.findall(".//relationships"):
            for relationship in relationships.findall("./relationship"):
                first = relationship.find("./first-end-point")
                second = relationship.find("./second-end-point")
                expression = relationship.find("./expression")

                first_id = first.get("object-id") if first is not None else None
                second_id = second.get("object-id") if second is not None else None
                first_label = object_labels.get(first_id, first_id or "(sem origem)")
                second_label = object_labels.get(second_id, second_id or "(sem destino)")
                expression_text = self._relationship_expression_to_text(expression)
                from_field, to_field = self._relationship_link_fields(expression)

                records.append(
                    {
                        "source": source_label,
                        "from_object_id": first_id,
                        "to_object_id": second_id,
                        "from_label": first_label,
                        "to_label": second_label,
                        "from_field": from_field,
                        "to_field": to_field,
                        "from_attributes": self._humanize_relationship_endpoint_attributes(
                            first.attrib if first is not None else {},
                            object_labels,
                        ),
                        "to_attributes": self._humanize_relationship_endpoint_attributes(
                            second.attrib if second is not None else {},
                            object_labels,
                        ),
                        "expression": expression_text,
                    }
                )
        return records

    def _humanize_relationship_endpoint_attributes(
        self,
        attributes: dict[str, Any],
        object_labels: dict[str, str],
    ) -> dict[str, Any]:
        """Substitui object-id técnico por nome amigável da tabela/objeto."""
        humanized = dict(attributes)
        object_id = humanized.pop("object-id", None)
        if object_id:
            humanized["objeto"] = object_labels.get(object_id, object_id)
        return humanized

    def _relationship_expression_to_text(self, expression: ET.Element | None) -> str:
        """Converte uma expressão XML de relacionamento em texto legível."""
        if expression is None:
            return "-"

        children = list(expression)
        op = expression.get("op")
        if not children:
            return op or "-"

        child_texts = [self._relationship_expression_to_text(child) for child in children]
        if len(child_texts) == 1:
            return child_texts[0]
        if op:
            separator = f" {op} "
            return separator.join(child_texts)
        return " | ".join(child_texts)

    def _relationship_link_fields(self, expression: ET.Element | None) -> tuple[str | None, str | None]:
        """Extrai os campos de ligação principais de cada lado do relacionamento."""
        if expression is None:
            return None, None

        children = list(expression)
        if expression.get("op") == "=" and len(children) >= 2:
            return (
                self._relationship_expression_to_text(children[0]),
                self._relationship_expression_to_text(children[1]),
            )
        if len(children) >= 2:
            return (
                self._relationship_expression_to_text(children[0]),
                self._relationship_expression_to_text(children[1]),
            )
        return None, None

    def _parse_connection(self, connection: ET.Element) -> dict[str, Any]:
        relation_rows = []
        for relation in connection.findall(".//relation"):
            relation_info = dict(relation.attrib)
            relation_text = normalize_whitespace("".join(relation.itertext()))
            custom_sql = self._extract_relation_sql(relation, relation_text)
            relation_info["text"] = relation_text
            relation_info["custom_sql"] = custom_sql
            relation_info["has_custom_sql"] = bool(custom_sql)
            relation_rows.append(relation_info)

        named_connections = []
        for named in connection.findall("./named-connections/named-connection"):
            named_connections.append(dict(named.attrib))

        return {
            "attributes": dict(connection.attrib),
            "relations": relation_rows,
            "named_connections": named_connections,
        }

    def _extract_relation_sql(self, relation: ET.Element, relation_text: str | None) -> str | None:
        """Tenta localizar SQL customizado embutido em relations do Tableau."""
        candidate_values = [
            relation.get("formula"),
            relation.get("query"),
            relation.get("command"),
            relation.get("sql"),
            relation_text,
        ]

        name = (relation.get("name") or "").lower()
        relation_type = (relation.get("type") or "").lower()
        for candidate in candidate_values:
            text = normalize_whitespace(decode_tableau_text(candidate) if candidate else candidate)
            if not text:
                continue
            lowered = text.lower()
            looks_like_sql = any(
                token in lowered
                for token in ["select ", "\nselect ", "with ", "\nwith ", " from ", "\nfrom ", "union ", "join "]
            )
            if looks_like_sql or relation_type == "text" or "custom sql" in name:
                return text
        return None

    def _parse_column(self, column: ET.Element) -> dict[str, Any]:
        aliases = [
            {"key": alias.get("key"), "value": alias.get("value")}
            for alias in column.findall("./aliases/alias")
        ]
        members = [member.get("value") for member in column.findall("./members/member") if member.get("value")]
        bins = []
        for bin_element in column.findall("./calculation/bin"):
            bins.append(
                {
                    "value": bin_element.get("value"),
                    "default_name": bin_element.get("default-name"),
                    "members": [value.text for value in bin_element.findall("./value") if value.text],
                }
            )

        calculation = column.find("./calculation")
        return {
            "name": column.get("name"),
            "caption": column.get("caption"),
            "role": column.get("role"),
            "datatype": column.get("datatype"),
            "type": column.get("type"),
            "default_type": column.get("default-type"),
            "default_format": column.get("default-format"),
            "aggregation": column.get("aggregation"),
            "hidden": column.get("hidden"),
            "is_used": column.get("is-used") or column.get("is_used"),
            "value": column.get("value"),
            "param_domain_type": column.get("param-domain-type"),
            "semantic_role": column.get("semantic-role"),
            "source_field": column.get("source-field"),
            "aliases": aliases,
            "members": members,
            "bins": bins,
            "formula": decode_tableau_text(calculation.get("formula")) if calculation is not None else None,
            "calculation_class": calculation.get("class") if calculation is not None else None,
        }

    def _parse_metadata_record(self, record: ET.Element) -> dict[str, Any]:
        attributes = {}
        for attribute in record.findall("./attributes/attribute"):
            key = attribute.get("name")
            if key:
                attributes[key] = decode_tableau_text(attribute.text or "")

        return {
            "class": record.get("class"),
            "caption": record.findtext("./caption"),
            "local_name": record.findtext("./local-name"),
            "remote_name": record.findtext("./remote-name"),
            "parent_name": record.findtext("./parent-name"),
            "local_type": record.findtext("./local-type"),
            "aggregation": record.findtext("./aggregation"),
            "attributes": attributes,
        }

    def _extract_parameters(self, datasources: list[dict[str, Any]]) -> list[dict[str, Any]]:
        for datasource in datasources:
            if datasource.get("name") == "Parameters":
                return datasource["columns"]
        return []

    def _extract_calculations(
        self,
        datasources: list[dict[str, Any]],
        parameters: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        friendly_names = {}
        for parameter in parameters:
            if parameter.get("name") and parameter.get("caption"):
                friendly_names[parameter["name"]] = parameter["caption"]

        records = []
        seen = set()

        for datasource in datasources:
            datasource_name = datasource.get("caption") or datasource.get("name")
            if datasource.get("name") == "Parameters":
                continue
            for column in datasource.get("columns", []):
                if not column.get("formula"):
                    continue
                if column.get("name") and column.get("caption"):
                    friendly_names[column["name"]] = column["caption"]

                key = ("column", datasource_name, column.get("name"), column.get("formula"))
                if key in seen:
                    continue
                seen.add(key)
                records.append(
                    {
                        "datasource": datasource_name,
                        "origin": "column",
                        "name": column.get("name"),
                        "caption": column.get("caption") or clean_brackets(column.get("name")),
                        "aliases": column.get("aliases", []),
                        "role": column.get("role"),
                        "datatype": column.get("datatype"),
                        "type": column.get("type"),
                        "hidden": column.get("hidden"),
                        "formula": column.get("formula"),
                    }
                )

            for calc in datasource.get("connection_calculations", []):
                key = ("connection", datasource_name, calc.get("column"), calc.get("formula"))
                if key in seen:
                    continue
                seen.add(key)
                records.append(
                    {
                        "datasource": datasource_name,
                        "origin": "connection.calculations",
                        "name": calc.get("column"),
                        "caption": clean_brackets(calc.get("column")),
                        "aliases": [],
                        "role": None,
                        "datatype": None,
                        "type": None,
                        "hidden": None,
                        "formula": calc.get("formula"),
                    }
                )

            for record in datasource.get("metadata_records", []):
                formula = record.get("attributes", {}).get("formula")
                if not formula:
                    continue
                key = ("metadata", datasource_name, record.get("local_name"), formula)
                if key in seen:
                    continue
                seen.add(key)
                records.append(
                    {
                        "datasource": datasource_name,
                        "origin": "metadata-record",
                        "name": record.get("local_name"),
                        "caption": record.get("caption") or clean_brackets(record.get("local_name")),
                        "aliases": [],
                        "role": None,
                        "datatype": record.get("local_type"),
                        "type": record.get("class"),
                        "hidden": None,
                        "formula": formula.strip('"'),
                    }
                )

        for record in records:
            if record.get("name") and record.get("caption"):
                friendly_names[record["name"]] = record["caption"]
                cleaned_name = clean_brackets(record["name"])
                if cleaned_name:
                    friendly_names[f"[{cleaned_name}]"] = record["caption"]
            record["formula_pretty"] = self._replace_internal_names(record.get("formula"), friendly_names)
            record["codigo"] = record["formula_pretty"] or record.get("formula")

        by_datasource = {}
        for record in records:
            by_datasource.setdefault(record["datasource"], []).append(record)

        for datasource_records in by_datasource.values():
            caption_by_name = {}
            for record in datasource_records:
                if record.get("name") and record.get("caption"):
                    caption_by_name[record["name"]] = record["caption"]

            for record in datasource_records:
                impacts = []
                target = f"[{record['caption']}]"
                for candidate in datasource_records:
                    if candidate is record:
                        continue
                    formula = candidate.get("formula_pretty") or ""
                    if target in formula:
                        impacts.append(candidate["caption"])
                record["impacts"] = sorted(unique_ordered([impact for impact in impacts if impact]), key=str.lower)

        return records

    def _replace_internal_names(self, formula: str | None, friendly_names: dict[str, str]) -> str | None:
        if not formula or not friendly_names:
            return formula

        pattern = re.compile(
            "|".join(re.escape(key) for key in sorted(friendly_names.keys(), key=len, reverse=True))
        )
        return pattern.sub(lambda match: f"[{friendly_names[match.group(0)]}]", formula)

    def _enrich_usage(
        self,
        parameters: list[dict[str, Any]],
        calculations: list[dict[str, Any]],
        worksheets: list[dict[str, Any]],
        dashboards: list[dict[str, Any]],
    ) -> None:
        """Mapeia uso de parâmetros e cálculos em worksheets e dashboards."""
        worksheet_xml = {
            worksheet.get("name"): ET.tostring(element, encoding="unicode")
            for worksheet, element in zip(
                worksheets,
                self.root.findall("./worksheets/worksheet"),
            )
        }
        dashboard_members = {
            dashboard["name"]: set(dashboard.get("worksheet_members", []))
            for dashboard in dashboards
        }

        for parameter in parameters:
            used_in_worksheets = sorted(
                name
                for name, xml in worksheet_xml.items()
                if self._xml_contains_reference(xml, parameter.get("name"), parameter.get("caption"))
            )
            used_in_dashboards = sorted(
                dashboard_name
                for dashboard_name, members in dashboard_members.items()
                if members.intersection(used_in_worksheets)
            )
            parameter["used_in_worksheets"] = used_in_worksheets
            parameter["used_in_dashboards"] = used_in_dashboards
            parameter["is_used"] = bool(used_in_worksheets or used_in_dashboards)
            parameter["source_field_details"] = self._parse_source_field_reference(parameter.get("source_field"))

        for calculation in calculations:
            used_in_worksheets = sorted(
                name
                for name, xml in worksheet_xml.items()
                if self._xml_contains_reference(xml, calculation.get("name"), calculation.get("caption"))
            )
            used_in_dashboards = sorted(
                dashboard_name
                for dashboard_name, members in dashboard_members.items()
                if members.intersection(used_in_worksheets)
            )
            calculation["used_in_worksheets"] = used_in_worksheets
            calculation["used_in_dashboards"] = used_in_dashboards
            calculation["is_used"] = bool(used_in_worksheets or used_in_dashboards)

    def _enrich_datasource_field_usage(
        self,
        datasources: list[dict[str, Any]],
        worksheets: list[dict[str, Any]],
        calculations: list[dict[str, Any]],
        parameters: list[dict[str, Any]],
    ) -> None:
        """Completa o indicador `is_used` dos campos da fonte com fallback baseado em referências reais."""
        worksheet_xml = {
            worksheet.get("name"): ET.tostring(element, encoding="unicode")
            for worksheet, element in zip(
                worksheets,
                self.root.findall("./worksheets/worksheet"),
            )
        }
        calculation_texts = [
            calculation.get("formula_pretty") or calculation.get("formula") or ""
            for calculation in calculations
        ]
        parameter_texts = [
            " ".join(
                filter(
                    None,
                    [
                        parameter.get("source_field"),
                        compact_json(parameter.get("source_field_details")) if parameter.get("source_field_details") else "",
                    ],
                )
            )
            for parameter in parameters
        ]
        reference_haystacks = list(worksheet_xml.values()) + calculation_texts + parameter_texts

        for datasource in datasources:
            for column in datasource.get("columns", []):
                if column.get("is_used") is not None:
                    continue
                tokens = self._reference_tokens(column.get("name"), self._display_object_name(column))
                column["is_used"] = any(
                    token and token in haystack
                    for token in tokens
                    for haystack in reference_haystacks
                )

    def _xml_contains_reference(self, xml_text: str, internal_name: str | None, caption: str | None) -> bool:
        """Busca referências de um campo no XML de uma worksheet."""
        candidates = []
        if internal_name:
            candidates.append(internal_name)
            clean_name = clean_brackets(internal_name)
            if clean_name:
                candidates.append(f"[{clean_name}]")
        if caption:
            candidates.append(f"[{caption}]")
        return any(candidate and candidate in xml_text for candidate in unique_ordered(candidates))

    def _parse_source_field_reference(self, value: str | None) -> dict[str, str] | None:
        """Quebra a referência source-field em datasource e campo, quando possível."""
        if not value:
            return None
        match = re.match(r"^\[(?P<datasource>.+?)\]\.\[(?P<field>.+?)\]$", value)
        if not match:
            return {"raw": value}
        return {
            "datasource": match.group("datasource"),
            "field": match.group("field"),
            "raw": value,
        }

    def _humanize_datasource_reference(self, value: str | None) -> str | None:
        """Converte o nome técnico do datasource para caption amigável, quando disponível."""
        if not value:
            return None
        if value in self.datasource_caption_lookup:
            return self.datasource_caption_lookup[value]
        normalized = normalize_lookup_token(clean_brackets(value))
        if normalized and normalized in self.datasource_caption_lookup_by_clean:
            return self.datasource_caption_lookup_by_clean[normalized]
        base_name = value.split(".", 1)[0] if "." in value else value
        normalized_base = normalize_lookup_token(clean_brackets(base_name))
        if normalized_base and normalized_base in self.datasource_caption_lookup_by_clean:
            return self.datasource_caption_lookup_by_clean[normalized_base]
        return value

    def _worksheet_calculation_labels(self, worksheet_name: str) -> list[str]:
        """
        Retorna os cálculos usados em uma planilha sem duplicar o mesmo campo calculado.
        A deduplicação usa preferencialmente o nome interno para não colidir campos
        distintos que compartilhem o mesmo caption.
        """
        labels_by_key: dict[str, str] = {}
        fallback_index = 0
        for calculation in self.metadata["calculations"]:
            if worksheet_name not in calculation.get("used_in_worksheets", []):
                continue

            label = calculation.get("caption") or clean_brackets(calculation.get("name")) or "(sem nome)"
            internal_name = calculation.get("name")
            key = internal_name or f"fallback::{label}::{calculation.get('datasource')}::{fallback_index}"
            if key not in labels_by_key:
                labels_by_key[key] = label
            fallback_index += 1

        return sorted(labels_by_key.values(), key=str.lower)

    def _custom_sql_by_connection(self, datasource: dict[str, Any]) -> list[dict[str, Any]]:
        """
        Retorna no máximo um SQL por conexão do datasource.
        Se houver o mesmo SQL repetido em várias relations da mesma conexão,
        apenas a primeira ocorrência é mantida.
        """
        records: list[dict[str, Any]] = []
        for index, connection in enumerate(datasource.get("connections", []), start=1):
            seen_sql: set[str] = set()
            for relation in connection.get("relations", []):
                custom_sql = relation.get("custom_sql")
                if not custom_sql or custom_sql in seen_sql:
                    continue
                seen_sql.add(custom_sql)
                records.append(
                    {
                        "connection_index": index,
                        "relation_name": relation.get("name") or relation.get("table") or f"relation_{index}",
                        "custom_sql": custom_sql,
                    }
                )
                break
        return records

    def _dedupe_external_custom_sql_relations(
        self,
        relations: list[dict[str, Any]],
    ) -> list[dict[str, Any]]:
        """Remove SQLs externos repetidos preservando a primeira ocorrência."""
        unique_relations: list[dict[str, Any]] = []
        seen_sql: set[str] = set()
        for relation in relations:
            custom_sql = relation.get("custom_sql")
            if not custom_sql or custom_sql in seen_sql:
                continue
            seen_sql.add(custom_sql)
            unique_relations.append(relation)
        return unique_relations

    def _dedupe_relationship_maps(self, relations: list[dict[str, Any]]) -> list[dict[str, Any]]:
        """Remove relacionamentos duplicados e consolida as origens encontradas."""
        unique_relations_by_key: dict[str, dict[str, Any]] = {}
        for relation in relations:
            key = json.dumps(
                {
                    "from_object_id": relation.get("from_object_id"),
                    "to_object_id": relation.get("to_object_id"),
                    "expression": relation.get("expression"),
                },
                ensure_ascii=False,
                sort_keys=True,
            )
            if key not in unique_relations_by_key:
                merged = dict(relation)
                source = relation.get("source")
                merged["sources"] = [source] if source else []
                unique_relations_by_key[key] = merged
                continue

            merged = unique_relations_by_key[key]
            source = relation.get("source")
            if source and source not in merged["sources"]:
                merged["sources"].append(source)
            if relation.get("tdsx_path") and not merged.get("tdsx_path"):
                merged["tdsx_path"] = relation.get("tdsx_path")
            if relation.get("tds_path") and not merged.get("tds_path"):
                merged["tds_path"] = relation.get("tds_path")
        return list(unique_relations_by_key.values())

    def _reference_tokens(self, name: str | None, caption: str | None) -> list[str]:
        """Monta variações de referência para busca em fórmulas e metadados."""
        tokens = []
        if name:
            tokens.append(name)
            clean_name = clean_brackets(name)
            if clean_name:
                tokens.append(f"[{clean_name}]")
        if caption:
            tokens.append(f"[{caption}]")
            tokens.append(caption)
        return unique_ordered([token for token in tokens if token])

    def _object_identifier(self, item: dict[str, Any], fallback_prefix: str) -> str:
        """Gera um identificador estável para cálculos e parâmetros."""
        name = item.get("name")
        if name:
            return str(name)
        caption = item.get("caption") or clean_brackets(item.get("name")) or "(sem nome)"
        return f"{fallback_prefix}::{caption}"

    def _collect_unused_objects(self) -> dict[str, list[dict[str, Any]]]:
        """
        Identifica cálculos, parâmetros e fontes de dados sem uso efetivo.

        Regras:
        - cálculo/parâmetro não pode ser usado em planilhas ou dashboards
        - cálculo/parâmetro não pode ser referenciado por outro cálculo ou parâmetro
        - datasource não pode ser usado em planilhas/dashboards
        - datasource não pode ser referenciado por cálculo/parâmetro fora da lista de não usados
        """
        calculations = self.metadata["calculations"]
        parameters = self.metadata["parameters"]
        datasources = self.metadata["datasources"]

        calc_ids = {self._object_identifier(item, "calc"): item for item in calculations}
        param_ids = {self._object_identifier(item, "param"): item for item in parameters}

        calc_referenced_by: dict[str, set[str]] = {key: set() for key in calc_ids}
        param_referenced_by: dict[str, set[str]] = {key: set() for key in param_ids}

        calc_tokens = {
            key: self._reference_tokens(item.get("name"), item.get("caption") or clean_brackets(item.get("name")))
            for key, item in calc_ids.items()
        }
        param_tokens = {
            key: self._reference_tokens(item.get("name"), item.get("caption") or clean_brackets(item.get("name")))
            for key, item in param_ids.items()
        }

        containers: list[tuple[str, str]] = []
        for key, item in calc_ids.items():
            containers.append((f"calc::{key}", item.get("formula_pretty") or item.get("formula") or ""))
        for key, item in param_ids.items():
            source = item.get("source_field") or ""
            source_details = compact_json(item.get("source_field_details")) if item.get("source_field_details") else ""
            members = " ".join(item.get("members", []))
            containers.append((f"param::{key}", " ".join([source, source_details, members]).strip()))

        for container_id, text in containers:
            if not text:
                continue
            for key, tokens in calc_tokens.items():
                if container_id == f"calc::{key}":
                    continue
                if any(token and token in text for token in tokens):
                    calc_referenced_by[key].add(container_id)
            for key, tokens in param_tokens.items():
                if container_id == f"param::{key}":
                    continue
                if any(token and token in text for token in tokens):
                    param_referenced_by[key].add(container_id)

        unused_calculation_ids = {
            key
            for key, item in calc_ids.items()
            if not item.get("used_in_worksheets")
            and not item.get("used_in_dashboards")
            and not calc_referenced_by[key]
        }
        unused_parameter_ids = {
            key
            for key, item in param_ids.items()
            if not item.get("used_in_worksheets")
            and not item.get("used_in_dashboards")
            and not param_referenced_by[key]
        }

        unused_calculations = [calc_ids[key] for key in unused_calculation_ids]
        unused_parameters = [param_ids[key] for key in unused_parameter_ids]

        active_datasource_refs = set()
        for item in calculations:
            item_id = self._object_identifier(item, "calc")
            if item_id in unused_calculation_ids:
                continue
            if item.get("datasource"):
                active_datasource_refs.add(item.get("datasource"))
        for item in parameters:
            item_id = self._object_identifier(item, "param")
            if item_id in unused_parameter_ids:
                continue
            source_details = item.get("source_field_details") or {}
            if source_details.get("datasource"):
                active_datasource_refs.add(source_details.get("datasource"))

        worksheet_datasources = {
            self._humanize_datasource_reference(value) or value
            for worksheet in self.metadata["worksheets"]
            for value in worksheet.get("datasource_dependencies", [])
        }
        dashboard_datasources = {
            self._humanize_datasource_reference(value) or value
            for dashboard in self.metadata["dashboards"]
            for value in dashboard.get("datasource_dependencies", [])
        }

        unused_datasources = []
        for datasource in datasources:
            ds_name = datasource.get("caption") or datasource.get("name") or "(sem nome)"
            if ds_name in worksheet_datasources or ds_name in dashboard_datasources:
                continue
            if ds_name in active_datasource_refs or datasource.get("name") in active_datasource_refs:
                continue
            unused_datasources.append(datasource)

        return {
            "calculations": sorted(
                unused_calculations,
                key=lambda item: (item.get("caption") or clean_brackets(item.get("name")) or "").lower(),
            ),
            "parameters": sorted(
                unused_parameters,
                key=lambda item: (item.get("caption") or clean_brackets(item.get("name")) or "").lower(),
            ),
            "datasources": sorted(
                unused_datasources,
                key=lambda item: (item.get("caption") or item.get("name") or "").lower(),
            ),
        }

    def _infer_datasource_type(self, datasource: dict[str, Any]) -> str:
        """Classifica a fonte como lógica ou física com base na estrutura das relations."""
        relationship_maps = self._dedupe_relationship_maps(
            (datasource.get("relationship_maps") or []) + (datasource.get("external_relationship_maps") or [])
        )
        relation_types = {
            (relation.get("type") or "").lower()
            for connection in datasource.get("connections", [])
            for relation in connection.get("relations", [])
        }
        if relationship_maps or "collection" in relation_types:
            return "lógico"
        return "físico (join)"

    def _extract_worksheets(self) -> list[dict[str, Any]]:
        worksheets = []
        for worksheet in self.root.findall("./worksheets/worksheet"):
            worksheet_name = worksheet.get("name")
            formatted_title = self._extract_formatted_text(worksheet.find("./layout-options/title/formatted-text"))
            datasource_dependencies = [
                dependency.get("datasource")
                for dependency in worksheet.findall(".//datasource-dependencies")
                if dependency.get("datasource")
            ]
            filters = [self._parse_filter(filter_element) for filter_element in worksheet.findall(".//filter")]
            encodings = self._parse_encodings(worksheet)
            marks = [mark.get("class") for mark in worksheet.findall(".//mark") if mark.get("class")]
            style_rules = self._parse_style_rules(worksheet)

            worksheets.append(
                {
                    "name": worksheet_name,
                    "title": formatted_title,
                    "datasource_dependencies": unique_ordered(datasource_dependencies),
                    "filters": filters,
                    "encodings": encodings,
                    "marks": unique_ordered(marks),
                    "style_rules": style_rules,
                    "colors_used": unique_ordered(self._collect_colors(worksheet)),
                    "fonts_used": unique_ordered(self._collect_fonts(worksheet)),
                    "referenced_columns": unique_ordered(self._collect_columns_from_dependencies(worksheet)),
                    "shelf_columns": [column.text for column in worksheet.findall(".//slices/column") if column.text],
                }
            )
        return worksheets

    def _extract_dashboards(self, worksheets: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
        dashboards = []
        worksheets = worksheets or []
        worksheet_names = {worksheet.get("name") for worksheet in self.root.findall("./worksheets/worksheet")}

        for dashboard in self.root.findall("./dashboards/dashboard"):
            zones = self._parse_zones(dashboard.find("./zones"))
            dashboard_name = dashboard.get("name")
            sheets = []
            exposed_filters = []
            for zone in zones:
                if zone.get("name") in worksheet_names:
                    sheets.append(zone["name"])
                if zone.get("type_v2") == "filter":
                    exposed_filters.append(
                        {
                            "zone_id": zone.get("id"),
                            "worksheet_or_dashboard": zone.get("name"),
                            "param": zone.get("param"),
                            "mode": zone.get("mode"),
                            "values": zone.get("values"),
                            "field_label": self._humanize_field_reference(zone.get("param")),
                            "is_context": False,
                        }
                    )

            device_layouts = []
            for layout in dashboard.findall("./devicelayouts/devicelayout"):
                device_layouts.append(
                    {
                        "name": layout.get("name"),
                        "auto_generated": layout.get("auto-generated"),
                        "zones": self._parse_zones(layout.find("./zones")),
                    }
                )

            worksheet_filters = []
            for worksheet_name in unique_ordered([sheet for sheet in sheets if sheet]):
                worksheet = next((item for item in worksheets if item.get("name") == worksheet_name), None)
                if worksheet:
                    worksheet_filters.extend(worksheet.get("filters", []))

            combined_filters = []
            for item in exposed_filters + worksheet_filters:
                field_label = item.get("field_label") or self._humanize_field_reference(item.get("column") or item.get("param"))
                combined_filters.append(
                    {
                        "field_label": field_label,
                        "class": item.get("class"),
                        "is_context": bool(item.get("is_context")),
                        "visibility": "exposto no painel" if "zone_id" in item else "interno da planilha",
                    }
                )
            combined_filters = sorted(
                unique_ordered(combined_filters),
                key=lambda item: (item.get("field_label") or "").lower(),
            )

            worksheet_colors = []
            worksheet_fonts = []
            for worksheet_name in unique_ordered([sheet for sheet in sheets if sheet]):
                worksheet = next((item for item in worksheets if item.get("name") == worksheet_name), None)
                if worksheet:
                    worksheet_colors.extend(worksheet.get("colors_used", []))
                    worksheet_fonts.extend(worksheet.get("fonts_used", []))

            dashboards.append(
                {
                    "name": dashboard_name,
                    "attributes": dict(dashboard.attrib),
                    "datasource_dependencies": unique_ordered(
                        [
                            dependency.get("datasource")
                            for dependency in dashboard.findall("./datasource-dependencies")
                            if dependency.get("datasource")
                        ]
                    ),
                    "zones": zones,
                    "worksheet_members": unique_ordered([sheet for sheet in sheets if sheet]),
                    "filters_exposed": exposed_filters,
                    "filters_used": combined_filters,
                    "device_layouts": device_layouts,
                    "colors_used": unique_ordered(self._collect_colors(dashboard) + worksheet_colors),
                    "fonts_used": unique_ordered(self._collect_fonts(dashboard) + worksheet_fonts),
                }
            )
        return dashboards

    def _extract_stories(self) -> list[dict[str, Any]]:
        stories = []
        for story in self.root.findall("./stories/story"):
            stories.append(
                {
                    "name": story.get("name"),
                    "attributes": dict(story.attrib),
                    "style_rules": self._parse_style_rules(story),
                    "colors_used": unique_ordered(self._collect_colors(story)),
                    "fonts_used": unique_ordered(self._collect_fonts(story)),
                }
            )
        return stories

    def _extract_windows(self) -> list[dict[str, Any]]:
        windows = []
        for window in self.root.findall("./windows/window"):
            cards = []
            for card in window.findall(".//card"):
                cards.append(dict(card.attrib))
            windows.append(
                {
                    "class": window.get("class"),
                    "name": window.get("name"),
                    "maximized": window.get("maximized"),
                    "cards": cards,
                }
            )
        return windows

    def _extract_thumbnails(self) -> list[dict[str, Any]]:
        thumbnails = []
        thumbnails_dir = self.output_dir / "thumbnails"
        if thumbnails_dir.exists():
            shutil.rmtree(thumbnails_dir)
        thumbnails_dir.mkdir(parents=True, exist_ok=True)

        for index, thumbnail in enumerate(self.root.findall("./thumbnails/thumbnail"), start=1):
            name = thumbnail.get("name") or f"thumbnail_{index}"
            file_name = f"{sanitize_filename(name)}.png"
            output_path = thumbnails_dir / file_name
            if thumbnail.text:
                output_path.write_bytes(base64.decodebytes(thumbnail.text.encode()))
                self.assets_written.append(str(output_path))
            thumbnails.append(
                {
                    "name": name,
                    "output_path": str(output_path),
                    "has_embedded_data": bool(thumbnail.text),
                }
            )
        return thumbnails

    def _extract_visual_tokens(self, token_type: str) -> list[dict[str, Any]]:
        """Varre o XML inteiro em busca de cores ou fontes com contexto e XPath."""
        results = []
        for element in self.root.iter():
            if token_type == "color":
                format_attr = element.get("attr")
                format_value = element.get("value")
                run_color = element.get("fontcolor")
                if format_attr and "color" in format_attr and format_value:
                    if is_color_like(format_value):
                        results.append(
                            {
                                "value": clean_display_label(format_value),
                                "context_tag": element.tag,
                                "attribute": format_attr,
                                "xpath": self._best_effort_xpath(element),
                            }
                        )
                if run_color:
                    if is_color_like(run_color):
                        results.append(
                            {
                                "value": clean_display_label(run_color),
                                "context_tag": element.tag,
                                "attribute": "fontcolor",
                                "xpath": self._best_effort_xpath(element),
                            }
                        )
            else:
                if element.get("attr") == "font-family" and element.get("value"):
                    results.append(
                        {
                            "value": element.get("value"),
                            "context_tag": element.tag,
                            "attribute": "font-family",
                            "xpath": self._best_effort_xpath(element),
                        }
                    )
                for attr_name in ("fontname", "font-family"):
                    if element.get(attr_name):
                        results.append(
                            {
                                "value": element.get(attr_name),
                                "context_tag": element.tag,
                                "attribute": attr_name,
                                "xpath": self._best_effort_xpath(element),
                            }
                        )
        return unique_ordered(results)

    def _parse_filter(self, filter_element: ET.Element) -> dict[str, Any]:
        members = []
        for groupfilter in filter_element.findall(".//groupfilter"):
            entry = {key: value for key, value in groupfilter.attrib.items()}
            if groupfilter.text and groupfilter.text.strip():
                entry["text"] = groupfilter.text.strip()
            members.append(entry)

        column_ref = filter_element.get("column")
        return {
            "class": filter_element.get("class"),
            "column": column_ref,
            "field_label": self._humanize_field_reference(column_ref),
            "filter_group": filter_element.get("filter-group"),
            "included_values": filter_element.get("included-values"),
            "is_context": self._is_context_filter(filter_element),
            "groupfilters": members,
        }

    def _is_context_filter(self, filter_element: ET.Element) -> bool:
        """Tenta identificar marcadores de contexto associados a um filtro."""
        serialized = ET.tostring(filter_element, encoding="unicode")
        return "context" in serialized.lower()

    def _humanize_field_reference(self, value: str | None) -> str | None:
        """Converte referências internas do Tableau em um nome mais legível."""
        if not value:
            return None
        original = value

        direct_candidates = re.findall(r"\[[^\]]+\]", original)
        for candidate in direct_candidates:
            if candidate in self.caption_lookup:
                return clean_display_label(self.caption_lookup[candidate])
            candidate_clean = clean_brackets(candidate)
            if candidate_clean in self.caption_lookup_by_clean:
                return clean_display_label(self.caption_lookup_by_clean[candidate_clean])

        text = original
        text = re.sub(r"^\[[^\]]+\]\.", "", text)
        text = text.replace("[", "").replace("]", "")
        text = re.sub(r"^(sum|avg|min|max|cnt|usr|none):", "", text)
        text = re.sub(r":[a-z]{1,3}$", "", text)
        text = re.sub(r"^\:+", "", text)

        if text in self.caption_lookup_by_clean:
            return clean_display_label(self.caption_lookup_by_clean[text])

        calc_match = re.search(r"(Calculation_\d+)", text)
        if calc_match and calc_match.group(1) in self.caption_lookup_by_clean:
            return clean_display_label(self.caption_lookup_by_clean[calc_match.group(1)])

        special_cases = {
            "Measure Names": "Measure Names",
            "Measure Values": "Measure Values",
            ":Measure Names": "Measure Names",
            ":Measure Values": "Measure Values",
        }
        if text in special_cases:
            return special_cases[text]
        if "Measure Names" in text:
            return "Measure Names"
        if "Measure Values" in text:
            return "Measure Values"

        text = re.sub(r"\s+\(\s*copy\s*\)\s*$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\(\s*cópia\s*\)\s*$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\(\s*local copy\s*\)\s*$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\(\s*local\s*\)\s*$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\|\s+snapshot$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\|\s+current year$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\|\s+next year$", "", text, flags=re.IGNORECASE)
        text = re.sub(r"\s+\(\s*prod_base_[^)]+\)$", "", text, flags=re.IGNORECASE)

        return clean_display_label(text)

    def _parse_encodings(self, element: ET.Element) -> list[dict[str, Any]]:
        encodings = []
        for encoding in element.findall(".//encodings/*"):
            record = {"tag": encoding.tag}
            record.update(encoding.attrib)
            encodings.append(record)
        for encoding in element.findall(".//style-rule/encoding"):
            record = {"tag": "style-encoding"}
            record.update(encoding.attrib)
            encodings.append(record)
        return encodings

    def _parse_style_rules(self, parent: ET.Element | None) -> list[dict[str, Any]]:
        if parent is None:
            return []
        rules = []
        for rule in parent.findall(".//style-rule"):
            formats = []
            for fmt in rule.findall("./format"):
                formats.append(dict(fmt.attrib))
            encodings = []
            for encoding in rule.findall("./encoding"):
                encodings.append(dict(encoding.attrib))
            rules.append(
                {
                    "element": rule.get("element"),
                    "formats": formats,
                    "encodings": encodings,
                }
            )
        return rules

    def _extract_formatted_text(self, formatted_text: ET.Element | None) -> list[dict[str, Any]]:
        if formatted_text is None:
            return []
        runs = []
        for run in formatted_text.findall("./run"):
            runs.append(
                {
                    "text": "".join(run.itertext()).strip(),
                    "attributes": dict(run.attrib),
                }
            )
        return runs

    def _collect_colors(self, element: ET.Element) -> list[str]:
        colors = []
        for item in element.iter():
            if item.get("attr") and "color" in item.get("attr", "") and item.get("value"):
                if is_color_like(item.get("value")):
                    colors.append(clean_display_label(item.get("value")))
            if item.get("fontcolor"):
                if is_color_like(item.get("fontcolor")):
                    colors.append(clean_display_label(item.get("fontcolor")))
        return colors

    def _collect_fonts(self, element: ET.Element) -> list[str]:
        fonts = []
        for item in element.iter():
            if item.get("attr") == "font-family" and item.get("value"):
                fonts.append(item.get("value"))
            if item.get("fontname"):
                fonts.append(item.get("fontname"))
        return fonts

    def _display_fonts(self, fonts: list[str] | None) -> list[str]:
        """Retorna fontes para exibição, com fallback quando o XML não explicita a tipografia."""
        values = sorted(set(fonts or []), key=str.lower)
        return values or [TABLEAU_IMPLICIT_FONT_LABEL]

    def _format_aliases(self, aliases: list[dict[str, Any]] | None) -> list[str]:
        """Converte aliases em linhas curtas e legíveis."""
        rows = []
        for alias in aliases or []:
            key = alias.get("key")
            value = alias.get("value")
            if key and value:
                rows.append(f"{key} -> {value}")
            elif value:
                rows.append(str(value))
            elif key:
                rows.append(str(key))
        return unique_ordered(rows)

    def _display_object_name(self, item: dict[str, Any], allow_clean_brackets: bool = True) -> str:
        """Retorna o nome mais amigável disponível para um objeto do Tableau."""
        if item.get("caption"):
            return str(item["caption"])
        if allow_clean_brackets:
            cleaned_name = clean_brackets(item.get("name"))
            if cleaned_name:
                return cleaned_name
        if item.get("name"):
            return str(item["name"])
        return "(sem nome)"

    def _display_real_name(self, item: dict[str, Any]) -> str:
        """Retorna o nome técnico/original do objeto para auditoria."""
        return item.get("name") or self._display_object_name(item)

    def _is_measure_names_column(self, column: dict[str, Any]) -> bool:
        """Identifica campos técnicos de Measure Names para escondê-los da listagem principal."""
        internal_name = clean_brackets(column.get("name")) or ""
        display_name = self._display_object_name(column)
        candidates = {internal_name.strip().lower(), display_name.strip().lower()}
        return any(value in {":measure names", "measure names"} for value in candidates)

    def _display_column_label(self, column: dict[str, Any]) -> str:
        """Monta o rótulo do campo com nome interno apenas quando ele agrega informação."""
        display_name = self._display_object_name(column)
        internal_name = clean_brackets(column.get("name"))
        if not internal_name:
            return display_name
        if internal_name.strip().lower() == display_name.strip().lower():
            return display_name
        return f"{display_name} ({column.get('name')})"

    def _group_datasource_columns(self, datasource: dict[str, Any]) -> list[tuple[str, list[dict[str, Any]]]]:
        """Agrupa campos por tabela de origem usando `metadata-records` quando disponível."""
        parent_name_by_column: dict[str, str] = {}
        for record in datasource.get("metadata_records", []):
            local_name = record.get("local_name")
            parent_name = clean_brackets(record.get("parent_name"))
            if local_name and parent_name and local_name not in parent_name_by_column:
                parent_name_by_column[local_name] = parent_name

        groups: dict[str, list[dict[str, Any]]] = {}
        group_order: list[str] = []
        for column in datasource.get("columns", []):
            if self._is_measure_names_column(column):
                continue
            group_name = parent_name_by_column.get(column.get("name")) or "Tabela não identificada"
            if group_name not in groups:
                groups[group_name] = []
                group_order.append(group_name)
            groups[group_name].append(column)

        return [(group_name, groups[group_name]) for group_name in group_order]

    def _collect_columns_from_dependencies(self, element: ET.Element) -> list[str]:
        columns = []
        for column in element.findall(".//datasource-dependencies/column"):
            columns.append(column.get("caption") or clean_brackets(column.get("name")))
        return [item for item in columns if item]

    def _parse_zones(self, zones_root: ET.Element | None) -> list[dict[str, Any]]:
        if zones_root is None:
            return []

        parsed = []

        def visit(zone: ET.Element, parent_id: str | None = None) -> None:
            zone_style = []
            for fmt in zone.findall("./zone-style/format"):
                zone_style.append(dict(fmt.attrib))

            record = {
                "id": zone.get("id"),
                "name": zone.get("name"),
                "type_v2": zone.get("type-v2"),
                "param": zone.get("param"),
                "mode": zone.get("mode"),
                "values": zone.get("values"),
                "x": zone.get("x"),
                "y": zone.get("y"),
                "w": zone.get("w"),
                "h": zone.get("h"),
                "parent_id": parent_id,
                "style": zone_style,
            }
            parsed.append(record)

            for child_zone in zone.findall("./zone"):
                visit(child_zone, zone.get("id"))

        for zone in zones_root.findall("./zone"):
            visit(zone)
        return parsed

    def _best_effort_xpath(self, element: ET.Element) -> str:
        try:
            return element_path_with_indices(self.root, element)
        except Exception:
            return ""

    def generate_xpath_json_map(self) -> tuple[Path, Path]:
        """Gera o documento humano-legível e a versão JSON do mapa XPath/JSON."""
        _log_progress("Gerando mapa XPath/JSON.")
        map_rows = []
        for definition in self.MAP_DEFINITIONS:
            count = self._count_xpath_matches(definition["xpath"])
            map_rows.append(
                {
                    **definition,
                    "match_count": count,
                }
            )

        map_json_path = self.output_dir / "mapa_XPath_JSON.json"
        map_md_path = self.output_dir / "mapa_XPath_JSON.md"
        map_json_path.write_text(
            json.dumps(
                {
                    "source_file": str(self.source_path),
                    "output_directory": str(self.output_dir),
                    "map_entries": map_rows,
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )

        lines = [
            f"# mapa_XPath_JSON - {self.source_path.name}",
            "",
            "## Como Ler",
            "",
            "- `XPath`: local esperado no XML do Tableau.",
            "- `JSON Path`: local correspondente na saída JSON gerada pelo script.",
            "- `Ocorrências`: quantidade encontrada no arquivo analisado.",
            "",
            "## Mapa",
            "",
            "| Seção | Label | XPath | JSON Path | Ocorrências | Descrição |",
            "|---|---|---|---|---:|---|",
        ]

        for row in map_rows:
            lines.append(
                f"| {row['section']} | {row['label']} | `{row['xpath']}` | `{row['json_path']}` | {row['match_count']} | {row['description']} |"
            )

        map_md_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        _log_progress("Mapa XPath/JSON gerado com sucesso.")
        return map_md_path, map_json_path

    def _count_xpath_matches(self, xpath: str) -> int:
        if xpath == ".":
            return 1
        if "|" in xpath:
            parts = [part.strip() for part in xpath.split("|")]
            count = 0
            for part in parts:
                count += len(self.root.findall(part))
            return count
        return len(self.root.findall(xpath))

    def write_outputs(self) -> list[Path]:
        """Escreve o mapa e os artefatos solicitados para o diretório de saída."""
        written_files: list[Path] = []
        _log_progress("Iniciando escrita dos artefatos finais.")

        workbook_xml_path = self.output_dir / f"{self.base_name}.xml"
        workbook_xml_path.write_bytes(self.workbook_xml_bytes or b"")
        written_files.append(workbook_xml_path)
        _log_progress(f"Arquivo XML salvo em `{workbook_xml_path.name}`.")

        map_md_path, map_json_path = self.generate_xpath_json_map()
        written_files.extend([map_md_path, map_json_path])

        if self.output_format in {"all", "json"}:
            written_files.append(self._write_json())
        if self.output_format in {"all", "markdown"}:
            written_files.append(self._write_markdown())
            written_files.append(self._write_rtf())
        if self.output_format == "rtf":
            written_files.append(self._write_rtf())
        if self.output_format in {"all", "docx"}:
            written_files.append(self._write_docx())
        if self.output_format in {"all", "excel"}:
            written_files.append(self._write_excel())

        manifest_path = self.output_dir / f"{self.base_name}_manifest.json"
        manifest_path.write_text(
            json.dumps(
                {
                    "source_file": str(self.source_path),
                    "output_directory": str(self.output_dir),
                    "generated_files": [str(path) for path in written_files],
                    "package_manifest": self.package_manifest,
                    "thumbnails_written": self.metadata["thumbnails"],
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        written_files.append(manifest_path)
        self._cleanup_temporary_outputs()
        _log_progress(f"Processamento concluído. Total de arquivos gerados: {len(written_files)}.")
        return written_files

    def _cleanup_temporary_outputs(self) -> None:
        """
        Remove artefatos temporários usados apenas durante o processamento.

        Mantém os arquivos finais de documentação e miniaturas, removendo
        diretórios intermediários temporários.
        """
        _log_progress("Removendo artefatos temporários de processamento.")
        for name in TEMPORARY_OUTPUT_NAMES:
            path = self.output_dir / name
            if not path.exists():
                continue
            if path.is_dir():
                shutil.rmtree(path, ignore_errors=True)
            else:
                try:
                    path.unlink()
                except FileNotFoundError:
                    pass

    def _write_json(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.json"
        _log_progress(f"Gerando arquivo JSON `{output_path.name}`.")
        output_path.write_text(
            json.dumps(self.metadata, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        return output_path

    def _write_markdown(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.md"
        _log_progress(f"Gerando arquivo Markdown `{output_path.name}`.")
        summary = self.metadata["summary"]
        lines = [
            f"# Documentação do Workbook Tableau - {self.source_path.name}",
            "",
            f"Relatório gerado em {self.generated_at}",
            "",
            "## Resumo",
            "",
            f"- Caminho de origem: `{self.source_path}`",
            f"- Tipo do arquivo: `{self.source_path.suffix.lower()}`",
            f"- Última alteração do arquivo: `{self.metadata['workbook']['source_file_last_modified']}`",
            f"- Datasources: {summary['datasource_count']}",
            f"- Parâmetros: {summary['parameter_count']} ({summary['unused_parameter_count']} não usados e não referenciados)",
            f"- Campos calculados: {summary['calculation_count']} ({summary['unused_calculation_count']} não usados e não referenciados)",
            f"- Worksheets: {summary['worksheet_count']} ({summary['worksheet_independent_count']} independentes)",
            f"- Dashboards: {summary['dashboard_count']}",
            "",
            "## Workbook",
            "",
        ]

        for key, value in self.metadata["workbook"]["attributes"].items():
            lines.append(f"- {key}: `{value}`")

        lines.extend(self._build_datasources_markdown())
        lines.extend(self._build_dashboards_markdown())
        lines.extend(self._build_visual_tokens_markdown())
        lines.extend(self._build_preferences_markdown())
        lines.extend(self._build_parameters_markdown())
        lines.extend(self._build_calculations_markdown())
        lines.extend(self._build_unused_objects_markdown())

        output_path.write_text("\n".join(lines) + "\n", encoding="utf-8")
        return output_path

    def _write_rtf(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.rtf"
        _log_progress(f"Gerando arquivo RTF `{output_path.name}`.")
        body = self._build_rtf_document()
        header = (
            r"{\rtf1\ansi\deff0"
            rf"{{\fonttbl{{\f0 {self._rtf_escape(RTF_BODY_FONT_NAME)};}}{{\f1 {self._rtf_escape(RTF_MONO_FONT_NAME)};}}}}"
            r"\viewkind4\uc1"
        )
        output_path.write_text(header + body + "}", encoding="utf-8")
        return output_path

    def _write_docx(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.docx"
        _log_progress(f"Gerando arquivo DOCX `{output_path.name}`.")
        document = Document()
        self._configure_docx_document(document)
        blocks = self._build_document_blocks()
        title_block_count = 2 if len(blocks) >= 2 else len(blocks)
        for block in blocks[:title_block_count]:
            self._append_docx_block(document, block)
        if blocks:
            spacer = document.add_paragraph()
            self._apply_docx_paragraph_format(spacer)
            self._append_docx_toc(document)
            document.add_page_break()
        for block in blocks[title_block_count:]:
            self._append_docx_block(document, block)
        document.save(output_path)
        return output_path

    def _build_rtf_document(self) -> str:
        return "".join(self._render_rtf_block(block) for block in self._build_document_blocks())

    def _build_document_blocks(self) -> list[dict[str, Any]]:
        parts: list[dict[str, Any]] = []
        summary = self.metadata["summary"]
        workbook = self.metadata["workbook"]

        parts.append(self._doc_paragraph(f"Documentação do Workbook Tableau - {self.source_path.name}", style="title"))
        parts.append(self._doc_paragraph(f"Relatório gerado em {self.generated_at}", style="subtitle"))

        parts.append(self._doc_paragraph("Seção de dados gerais", style="section"))
        parts.append(self._doc_paragraph("Resumo", style="subsection"))
        parts.append(self._doc_bullet(f"Caminho de origem: {self.source_path}"))
        parts.append(self._doc_bullet(f"Tipo do arquivo: {self.source_path.suffix.lower()}"))
        parts.append(self._doc_bullet(f"Última alteração do arquivo: {workbook['source_file_last_modified']}"))
        parts.append(self._doc_bullet(f"Datasources: {summary['datasource_count']}"))
        parts.append(
            self._doc_bullet(
                f"Parâmetros: {summary['parameter_count']} ({summary['unused_parameter_count']} não usados e não referenciados)"
            )
        )
        parts.append(
            self._doc_bullet(
                f"Campos calculados: {summary['calculation_count']} ({summary['unused_calculation_count']} não usados e não referenciados)"
            )
        )
        parts.append(
            self._doc_bullet(
                f"Worksheets: {summary['worksheet_count']} ({summary['worksheet_independent_count']} independentes)"
            )
        )
        parts.append(self._doc_bullet(f"Dashboards: {summary['dashboard_count']}"))
        parts.append(self._doc_bullet(f"Itens do pacote: {summary['package_member_count']}"))

        parts.append(self._doc_paragraph("Workbook", style="subsection"))
        parts.append(self._doc_bullet(f"Arquivo no pacote: {workbook.get('package_workbook_name') or '-'}"))
        for key, value in workbook.get("attributes", {}).items():
            parts.append(self._doc_bullet(f"{key}: {value}"))

        parts.append(self._doc_paragraph("Seção Fontes de dados", style="section"))
        parts.append(self._doc_paragraph("Fontes de dados", style="subsection"))
        datasources = self.metadata["datasources"]
        if not datasources:
            parts.append(self._doc_bullet("Nenhuma fonte de dados identificada."))
        for datasource in datasources:
            ds_name = datasource.get("caption") or datasource.get("name") or "(sem nome)"
            datasource_type = self._infer_datasource_type(datasource)
            package_custom_sql_by_connection = self._custom_sql_by_connection(datasource)
            external_custom_sql_relations = self._dedupe_external_custom_sql_relations(
                datasource.get("external_custom_sql_relations", [])
            )
            relationship_maps = self._dedupe_relationship_maps(
                (datasource.get("relationship_maps") or []) + (datasource.get("external_relationship_maps") or [])
            )
            parts.append(self._doc_paragraph(ds_name, style="subsubsection"))
            parts.append(self._doc_bullet(f"Tipo: {datasource_type}", level=1))
            parts.append(self._doc_bullet(f"Versão: {datasource.get('version')}", level=1))
            parts.append(self._doc_bullet(f"Inline: {datasource.get('inline')}", level=1))
            parts.append(self._doc_bullet(f"Has connection: {datasource.get('hasconnection')}", level=1))
            parts.append(self._doc_bullet(f"Quantidade de conexões: {len(datasource.get('connections', []))}", level=1))
            parts.append(self._doc_bullet(f"Quantidade de campos: {len(datasource.get('columns', []))}", level=1))
            parts.append(self._doc_bullet(f"Quantidade de metadata records: {len(datasource.get('metadata_records', []))}", level=1))
            parts.append(self._doc_bullet(f"Quantidade de cálculos na conexão: {len(datasource.get('connection_calculations', []))}", level=1))
            parts.append(self._doc_bullet(f"Quantidade de objetos: {datasource.get('object_count')}", level=1))
            parts.append(self._doc_paragraph("Campos da fonte de dados", style="body_bold", level=1))
            grouped_columns = self._group_datasource_columns(datasource)
            if grouped_columns:
                for table_name, columns in grouped_columns:
                    parts.append(self._doc_paragraph(table_name, style="body_bold", level=2))
                    for column in columns:
                        parts.append(self._doc_bullet(self._display_column_label(column), level=3))
                        parts.append(self._doc_bullet(f"Datatype: {column.get('datatype') or '-'}", level=4))
                        parts.append(self._doc_bullet(f"Role: {column.get('role') or '-'}", level=4))
                        parts.append(self._doc_bullet(f"Em uso: {format_yes_no(column.get('is_used'))}", level=4))
                        parts.append(self._doc_bullet(f"Oculto: {format_yes_no(column.get('hidden'))}", level=4))
                        alias_rows = self._format_aliases(column.get("aliases", []))
                        if alias_rows:
                            parts.extend(
                                self._doc_list_block(
                                    "Aliases",
                                    alias_rows,
                                    level=4,
                                    empty_text="nenhum alias",
                                )
                            )
            else:
                parts.append(self._doc_bullet("Nenhum campo identificado.", level=2))
            repo = datasource.get("repository_location") or {}
            if repo:
                parts.append(self._doc_paragraph("Localização publicada", style="body_bold", level=1))
                for key, value in repo.items():
                    parts.append(self._doc_bullet(f"{key}: {value}", level=2))
            for index, connection in enumerate(datasource.get("connections", []), start=1):
                parts.append(self._doc_paragraph(f"Conexão {index}", style="body_bold", level=1))
                parts.append(self._doc_bullet(f"Atributos: {compact_json(connection.get('attributes'))}", level=2))
                parts.append(self._doc_bullet(f"Modo da conexão: {connection.get('mode') or '-'}", level=2))
                if connection.get("hyper_paths"):
                    parts.extend(
                        self._doc_list_block(
                            "Arquivos .hyper associados",
                            connection.get("hyper_paths", []),
                            level=2,
                            empty_text="-",
                        )
                    )
                connection_diagram = self._build_connection_diagram(connection, relationship_maps=relationship_maps)
                if connection_diagram:
                    parts.append(self._doc_bullet("Estrutura da conexão:", level=2))
                    parts.append(self._doc_code_block(connection_diagram, level=3))
                connection_sql = next(
                    (item for item in package_custom_sql_by_connection if item.get("connection_index") == index),
                    None,
                )
                if connection_sql:
                    parts.append(self._doc_bullet("SQL customizado:", level=2))
                    parts.append(self._doc_code_block(connection_sql["custom_sql"], level=3))
            if not package_custom_sql_by_connection:
                parts.append(self._doc_bullet("SQL customizado encontrado no pacote: não", level=1))
            if external_custom_sql_relations:
                parts.append(self._doc_bullet("SQL customizado encontrado em `.tdsx` externo: sim", level=1))
                for relation in external_custom_sql_relations:
                    relation_name = relation.get("relation_name") or "(sem nome)"
                    parts.append(
                        self._doc_bullet(
                            f"Origem externa: {relation.get('tdsx_path')} | relation: {relation_name}",
                            level=2,
                        )
                    )
                    parts.append(self._doc_bullet("SQL externo:", level=2))
                    parts.append(self._doc_code_block(relation["custom_sql"], level=3))
            elif datasource.get("external_tdsx_matches"):
                parts.append(self._doc_bullet("SQL customizado encontrado em `.tdsx` externo: não", level=1))
            if relationship_maps:
                parts.append(self._doc_paragraph("Mapa de relacionamentos", style="body_bold", level=1))
                for relation in relationship_maps:
                    from_label = relation.get("from_label") or "(sem origem)"
                    to_label = relation.get("to_label") or "(sem destino)"
                    parts.append(self._doc_bullet(f"{from_label} -> {to_label}", level=2))
                    parts.append(self._doc_bullet(f"Condição: {relation.get('expression') or '-'}", level=3))
                    if relation.get("from_field"):
                        parts.append(self._doc_bullet(f"Campo de ligação na origem: {relation.get('from_field')}", level=3))
                    if relation.get("to_field"):
                        parts.append(self._doc_bullet(f"Campo de ligação no destino: {relation.get('to_field')}", level=3))
                    relation_sources = relation.get("sources") or ([relation.get("source")] if relation.get("source") else [])
                    if relation_sources:
                        parts.append(self._doc_bullet(f"Origem: {', '.join(relation_sources)}", level=3))
                    from_attrs = relation.get("from_attributes") or {}
                    to_attrs = relation.get("to_attributes") or {}
                    if from_attrs:
                        parts.append(self._doc_bullet(f"Atributos de origem: {compact_json(from_attrs)}", level=3))
                    if to_attrs:
                        parts.append(self._doc_bullet(f"Atributos de destino: {compact_json(to_attrs)}", level=3))

        parts.append(self._doc_paragraph("Seção de Dashboards", style="section"))
        parts.append(self._doc_paragraph("Dashboards", style="subsection"))
        dashboards = self.metadata["dashboards"]
        if not dashboards:
            parts.append(self._doc_bullet("Nenhum dashboard encontrado."))
        for dashboard in dashboards:
            dashboard_name = dashboard.get("name") or "(sem nome)"
            parts.append(self._doc_paragraph(dashboard_name, style="subsubsection"))

            parts.append(self._doc_paragraph("Planilhas", style="body_bold", level=1))
            worksheet_members = dashboard.get("worksheet_members") or []
            if not worksheet_members:
                parts.append(self._doc_bullet("Nenhuma planilha identificada.", level=2))
            for worksheet_name in worksheet_members:
                worksheet = next((item for item in self.metadata["worksheets"] if item.get("name") == worksheet_name), None)
                calculations_used = self._worksheet_calculation_labels(worksheet_name)
                parameters_used = sorted(
                    [
                        param.get("caption") or clean_brackets(param.get("name")) or "(sem nome)"
                        for param in self.metadata["parameters"]
                        if worksheet_name in param.get("used_in_worksheets", [])
                    ],
                    key=str.lower,
                )
                parts.append(self._doc_paragraph(worksheet_name, style="body_bold", level=2))
                if worksheet and worksheet.get("title"):
                    parts.append(self._doc_bullet(f"Título formatado: {compact_json(worksheet.get('title'))}", level=3))
                parts.extend(
                    self._doc_list_block(
                        "Datasources dependentes",
                        [
                            self._humanize_datasource_reference(value) or value
                            for value in (worksheet.get("datasource_dependencies", []) if worksheet else [])
                        ],
                        level=3,
                        empty_text="-",
                    )
                )
                parts.append(self._doc_bullet("Campos referenciados:", level=3))
                referenced_columns = worksheet.get("referenced_columns", []) if worksheet else []
                if referenced_columns:
                    for referenced_column in referenced_columns:
                        parts.append(self._doc_bullet(referenced_column, level=4))
                else:
                    parts.append(self._doc_bullet("Nenhum campo referenciado identificado.", level=4))
                if SHOW_WORKSHEET_SHELF_COLUMNS_IN_RTF:
                    parts.append(
                        self._doc_bullet(
                            f"Campos em shelves: {', '.join(worksheet.get('shelf_columns', [])) if worksheet else '-'}",
                            level=3,
                        )
                    )

                parts.append(self._doc_paragraph("Campos calculados", style="body_bold", level=3))
                if calculations_used:
                    for calculation_name in calculations_used:
                        parts.append(self._doc_bullet(calculation_name, level=4))
                else:
                    parts.append(self._doc_bullet("Nenhum campo calculado associado.", level=4))

                parts.append(self._doc_paragraph("Filtros", style="body_bold", level=3))
                worksheet_filters = worksheet.get("filters", []) if worksheet else []
                if worksheet_filters:
                    for filter_item in worksheet_filters:
                        field_label = filter_item.get("field_label") or filter_item.get("column") or "(sem nome)"
                        parts.append(
                            self._doc_bullet(
                                f"{field_label} | tipo: {filter_item.get('class') or '-'} | contexto: {'sim' if filter_item.get('is_context') else 'não'}",
                                level=4,
                            )
                        )
                        if SHOW_FILTER_GROUP_VALUES_IN_RTF and filter_item.get("groupfilters"):
                            parts.append(
                                self._doc_bullet(
                                    f"Valores/grupos: {compact_json(filter_item.get('groupfilters'))}",
                                    level=5,
                                )
                            )
                else:
                    parts.append(self._doc_bullet("Nenhum filtro identificado.", level=4))

                parts.append(self._doc_paragraph("Parâmetros", style="body_bold", level=3))
                if parameters_used:
                    for parameter_name in parameters_used:
                        parts.append(self._doc_bullet(parameter_name, level=4))
                else:
                    parts.append(self._doc_bullet("Nenhum parâmetro associado.", level=4))

            parts.append(self._doc_paragraph("Filtros expostos no painel", style="body_bold", level=1))
            exposed_filters = dashboard.get("filters_exposed") or []
            if exposed_filters:
                for filter_item in exposed_filters:
                    field_label = filter_item.get("field_label") or filter_item.get("param") or "(sem nome)"
                    parts.append(
                        self._doc_bullet(
                            f"{field_label} | modo: {filter_item.get('mode') or '-'} | valores: {filter_item.get('values') or '-'}",
                            level=2,
                        )
                    )
            else:
                parts.append(self._doc_bullet("Nenhum filtro exposto identificado.", level=2))

            parts.append(self._doc_paragraph("Zonas de layout", style="body_bold", level=1))
            if dashboard.get("zones"):
                for zone in dashboard["zones"]:
                    zone_summary = (
                        f"id={zone.get('id')} | nome={zone.get('name') or '-'} | tipo={zone.get('type_v2') or '-'} "
                        f"| x={zone.get('x') or '-'} | y={zone.get('y') or '-'} | w={zone.get('w') or '-'} | h={zone.get('h') or '-'}"
                    )
                    parts.append(self._doc_bullet(zone_summary, level=2))
            else:
                parts.append(self._doc_bullet("Nenhuma zona identificada.", level=2))

            parts.append(self._doc_paragraph("Layout de devices", style="body_bold", level=1))
            if dashboard.get("device_layouts"):
                for layout in dashboard["device_layouts"]:
                    parts.append(
                        self._doc_bullet(
                            f"{layout.get('name') or '(sem nome)'} | auto generated: {layout.get('auto_generated') or '-'} | zonas: {len(layout.get('zones', []))}",
                            level=2,
                        )
                    )
            else:
                parts.append(self._doc_bullet("Nenhum layout de device identificado.", level=2))

            parts.append(self._doc_paragraph("Cores usadas", style="body_bold", level=1))
            colors_used = sorted(set(dashboard.get("colors_used", [])), key=str.lower)
            if colors_used:
                for color in colors_used:
                    parts.append(self._doc_bullet(color, level=2))
            else:
                parts.append(self._doc_bullet("Nenhuma cor identificada.", level=2))

            parts.append(self._doc_paragraph("Fontes usadas", style="body_bold", level=1))
            fonts_used = self._display_fonts(dashboard.get("fonts_used", []))
            for font in fonts_used:
                parts.append(self._doc_bullet(font, level=2))

        parts.append(self._doc_paragraph("Seção Visual", style="section"))
        parts.append(self._doc_paragraph("Tokens Visuais", style="subsection"))

        parts.append(self._doc_paragraph("Cores", style="body_bold", level=1))
        visual_colors = sorted(
            {item["value"] for item in self.metadata["visual_tokens"]["colors"] if item.get("value")},
            key=str.lower,
        )
        if visual_colors:
            for color in visual_colors:
                parts.append(self._doc_bullet(color, level=2))
        else:
            parts.append(self._doc_bullet("Nenhuma cor identificada.", level=2))

        parts.append(self._doc_paragraph("Paletas de Cor", style="body_bold", level=1))
        palettes = self.metadata["preferences"]["color_palettes"]
        if palettes:
            for palette in palettes:
                palette_name = palette.get("name") or "(sem nome)"
                parts.append(
                    self._doc_bullet(
                        f"{palette_name} | tipo: {palette.get('type') or '-'} | custom: {palette.get('custom') or '-'}",
                        level=2,
                    )
                )
                for color in palette.get("colors", []):
                    parts.append(self._doc_bullet(color, level=3, mono=self._looks_like_hex_color(color)))
        else:
            parts.append(self._doc_bullet("Nenhuma paleta de cor identificada.", level=2))

        parts.append(self._doc_paragraph("Fontes", style="body_bold", level=1))
        visual_fonts = self._display_fonts(
            [item["value"] for item in self.metadata["visual_tokens"]["fonts"] if item.get("value")]
        )
        for font in visual_fonts:
            parts.append(self._doc_bullet(font, level=2))

        parts.append(self._doc_paragraph("Preferências e Paletas", style="body_bold", level=1))
        preferences = self.metadata["preferences"]["preferences"]
        if preferences:
            for pref in preferences:
                parts.append(self._doc_bullet(f"{pref.get('name')}: {pref.get('value')}", level=2))
        else:
            parts.append(self._doc_bullet("Nenhuma preferência global identificada.", level=2))
        if palettes:
            for palette in palettes:
                parts.append(
                    self._doc_bullet(
                        f"Paleta {palette.get('name') or '(sem nome)'} com {len(palette.get('colors', []))} cores",
                        level=2,
                    )
                )

        parts.append(self._doc_paragraph("Seção Parâmetros", style="section"))
        parts.append(self._doc_paragraph("Parâmetros", style="subsection"))
        parameters = self.metadata["parameters"]
        if not parameters:
            parts.append(self._doc_bullet("Nenhum parâmetro encontrado."))
        for parameter in parameters:
            name = parameter.get("caption") or clean_brackets(parameter.get("name")) or "(sem nome)"
            parts.append(self._doc_paragraph(name, style="subsubsection"))
            parts.append(self._doc_bullet(f"Nome interno: {parameter.get('name')}", level=1))
            parts.append(self._doc_bullet(f"Datatype: {parameter.get('datatype')}", level=1))
            parts.append(self._doc_bullet(f"Type: {parameter.get('type')}", level=1))
            parts.append(self._doc_bullet(f"Role: {parameter.get('role')}", level=1))
            parts.append(self._doc_bullet(f"Domínio: {parameter.get('param_domain_type')}", level=1))
            parts.append(self._doc_bullet(f"Valor atual/default: {parameter.get('value')}", level=1))
            parts.extend(
                self._doc_list_block(
                    "Membros",
                    parameter.get("members", []),
                    level=1,
                    empty_text="-",
                )
            )
            source_field = parameter.get("source_field_details")
            if source_field:
                parts.append(self._doc_bullet(f"Preenchido por campo da fonte: {compact_json(source_field)}", level=1))
            else:
                parts.append(self._doc_bullet("Preenchido por campo da fonte: não", level=1))
            used_dashboards = sorted(parameter.get("used_in_dashboards", []), key=str.lower)
            used_worksheets = sorted(parameter.get("used_in_worksheets", []), key=str.lower)
            parts.extend(
                self._doc_list_block(
                    "Usado nos dashboards",
                    used_dashboards,
                    level=1,
                    empty_text="não identificado",
                )
            )
            parts.extend(
                self._doc_list_block(
                    "Usado nas planilhas",
                    used_worksheets,
                    level=1,
                    empty_text="não identificado",
                )
            )

        parts.append(self._doc_paragraph("Relação de Campos Calculados", style="section"))
        parts.append(self._doc_paragraph("Relação de Campos Calculados", style="subsection"))
        calculations = self.metadata["calculations"]
        if not calculations:
            parts.append(self._doc_bullet("Nenhum campo calculado encontrado."))
        grouped_calculations: dict[str, list[dict[str, Any]]] = {}
        for calculation in calculations:
            grouped_calculations.setdefault(calculation.get("datasource") or "(sem datasource)", []).append(calculation)
        for datasource_name, items in grouped_calculations.items():
            parts.append(self._doc_paragraph(datasource_name, style="subsubsection"))
            for item in items:
                caption = self._display_object_name(item)
                parts.append(self._doc_paragraph(caption, style="body_bold", level=1))
                parts.append(self._doc_bullet(f"Nome interno: {item.get('name')}", level=2))
                parts.append(self._doc_bullet(f"Origem: {item.get('origin')}", level=2))
                parts.append(self._doc_bullet(f"Role: {item.get('role') or '-'}", level=2))
                parts.append(self._doc_bullet(f"Datatype: {item.get('datatype') or '-'}", level=2))
                parts.append(self._doc_bullet(f"Type: {item.get('type') or '-'}", level=2))
                parts.append(self._doc_bullet(f"Em uso: {format_yes_no(item.get('is_used'))}", level=2))
                parts.append(self._doc_bullet(f"Oculto: {format_yes_no(item.get('hidden'))}", level=2))
                alias_rows = self._format_aliases(item.get("aliases", []))
                if alias_rows:
                    parts.extend(
                        self._doc_list_block(
                            "Aliases",
                            alias_rows,
                            level=2,
                            empty_text="nenhum alias",
                        )
                    )
                impacts = sorted(item.get("impacts", []), key=str.lower)
                if impacts:
                    parts.extend(
                        self._doc_list_block(
                            "Impacta / é referenciado por",
                            impacts,
                            level=2,
                            empty_text="nenhum outro campo calculado",
                        )
                    )
                used_in_dashboards = sorted(item.get("used_in_dashboards", []), key=str.lower)
                used_in_worksheets = sorted(item.get("used_in_worksheets", []), key=str.lower)
                if used_in_dashboards:
                    parts.extend(
                        self._doc_list_block(
                            "Usado nos dashboards",
                            used_in_dashboards,
                            level=2,
                            empty_text="não identificado",
                        )
                    )
                if used_in_worksheets:
                    parts.extend(
                        self._doc_list_block(
                            "Usado nas planilhas",
                            used_in_worksheets,
                            level=2,
                            empty_text="não identificado",
                        )
                    )
                parts.append(self._doc_bullet("Código:", level=2))
                parts.append(self._doc_code_block(item.get("codigo") or "-", level=3))

        unused_objects = self._collect_unused_objects()
        parts.append(self._doc_paragraph("Objetos não usados", style="section"))
        parts.append(self._doc_paragraph("Objetos não usados", style="subsection"))

        parts.append(self._doc_paragraph("Campos calculados", style="body_bold", level=1))
        if unused_objects["calculations"]:
            for item in unused_objects["calculations"]:
                caption = self._display_object_name(item)
                parts.append(self._doc_bullet(caption, level=2))
                parts.append(self._doc_bullet(f"Oculto: {format_yes_no(item.get('hidden'))}", level=3))
        else:
            parts.append(self._doc_bullet("Nenhum campo calculado não usado identificado.", level=2))

        parts.append(self._doc_paragraph("Parâmetros", style="body_bold", level=1))
        if unused_objects["parameters"]:
            for item in unused_objects["parameters"]:
                caption = self._display_object_name(item)
                parts.append(self._doc_bullet(caption, level=2))
        else:
            parts.append(self._doc_bullet("Nenhum parâmetro não usado identificado.", level=2))

        parts.append(self._doc_paragraph("Fontes de dados", style="body_bold", level=1))
        if unused_objects["datasources"]:
            for item in unused_objects["datasources"]:
                caption = self._display_object_name(item, allow_clean_brackets=False)
                parts.append(self._doc_bullet(caption, level=2))
        else:
            parts.append(self._doc_bullet("Nenhuma fonte de dados não usada identificada.", level=2))

        return parts

    def _doc_paragraph(self, text: str, style: str = "body", level: int = 0, mono: bool = False) -> dict[str, Any]:
        return {"type": "paragraph", "text": text, "style": style, "level": level, "mono": mono}

    def _doc_bullet(self, text: str, level: int = 0, mono: bool = False) -> dict[str, Any]:
        return {"type": "bullet", "text": text, "level": level, "mono": mono}

    def _doc_code_block(self, text: str, level: int = 0) -> dict[str, Any]:
        return {"type": "code", "text": text, "level": level, "mono": True}

    def _doc_list_block(
        self,
        label: str,
        values: list[Any],
        level: int = 0,
        empty_text: str = "-",
        mono: bool = False,
    ) -> list[dict[str, Any]]:
        parts = [self._doc_bullet(f"{label}:", level=level)]
        if values:
            for value in values:
                parts.append(self._doc_bullet(str(value), level=level + 1, mono=mono))
        else:
            parts.append(self._doc_bullet(empty_text, level=level + 1, mono=mono))
        return parts

    def _render_rtf_block(self, block: dict[str, Any]) -> str:
        if block["type"] == "paragraph":
            return self._rtf_paragraph(
                block["text"],
                style=block.get("style", "body"),
                level=block.get("level", 0),
                mono=block.get("mono", False),
            )
        if block["type"] == "bullet":
            return self._rtf_bullet(
                block["text"],
                level=block.get("level", 0),
                mono=block.get("mono", False),
            )
        return self._rtf_code_block(block["text"], level=block.get("level", 0))

    def _configure_docx_document(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        self._set_docx_font_family(normal_style.font, DOCX_BODY_FONT_NAME)
        normal_style.font.size = Pt(10)
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.line_spacing = 1.0
        normal_paragraph_format.space_before = Pt(0)
        normal_paragraph_format.space_after = Pt(6)
        document.core_properties.title = f"Documentação do Workbook Tableau - {self.source_path.name}"
        self._configure_docx_header(document)
        self._configure_docx_footer(document)

    def _append_docx_block(self, document: Document, block: dict[str, Any]) -> None:
        block_type = block["type"]
        if block_type == "bullet":
            paragraph = document.add_paragraph()
            self._apply_docx_paragraph_format(paragraph)
            level = block.get("level", 0)
            paragraph.paragraph_format.left_indent = Pt(18 * max(level, 0))
            paragraph.paragraph_format.first_line_indent = Pt(-12)
            run = paragraph.add_run(f"• {block['text']}")
            self._apply_docx_run_style(
                run,
                mono=block.get("mono", False),
                bold=False,
                italic=False,
                font_size=10,
            )
            return
        if block_type == "code":
            paragraph = document.add_paragraph()
            self._apply_docx_paragraph_format(paragraph)
            paragraph.paragraph_format.left_indent = Pt(18 * max(block.get("level", 0), 0))
            run = paragraph.add_run(block["text"])
            self._apply_docx_run_style(run, mono=True, bold=False, italic=False, font_size=10)
            return

        paragraph = document.add_paragraph()
        self._apply_docx_paragraph_format(paragraph)
        style = block.get("style", "body")
        level = block.get("level", 0)
        heading_style = self._docx_heading_style_for_block(style)
        if heading_style:
            paragraph.style = heading_style
        paragraph.paragraph_format.left_indent = Pt(18 * max(level, 0))
        font_size = 10
        bold = False
        italic = False
        if style == "title":
            font_size = 16
            bold = True
        elif style == "subtitle":
            font_size = 10
            italic = True
        elif style == "section":
            font_size = 14
            bold = True
            paragraph.paragraph_format.left_indent = Cm(-0.8)
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "subsection":
            font_size = 12
            bold = True
            paragraph.paragraph_format.left_indent = Cm(-0.7)
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "subsubsection":
            font_size = 11
            bold = True
            paragraph.paragraph_format.space_before = Pt(12)
        elif style == "body_bold":
            font_size = 10.5
            bold = True
        run = paragraph.add_run(block["text"])
        self._apply_docx_run_style(
            run,
            mono=block.get("mono", False),
            bold=bold,
            italic=italic,
            font_size=font_size,
        )

    def _apply_docx_run_style(
        self,
        run: Any,
        mono: bool,
        bold: bool,
        italic: bool,
        font_size: float,
    ) -> None:
        run.bold = bold
        run.italic = italic
        self._set_docx_font_family(
            run.font,
            DOCX_MONO_FONT_NAME if mono else DOCX_BODY_FONT_NAME,
        )
        run.font.size = Pt(font_size)

    def _set_docx_font_family(self, font: Any, font_name: str) -> None:
        """
        Define a fonte em todos os slots relevantes do DOCX.

        Apenas `font.name` pode não ser suficiente para o Word, que em alguns
        cenários mantém Calibri se `w:rFonts` não estiver totalmente preenchido.
        """
        font.name = font_name
        if getattr(font, "element", None) is None:
            return
        r_pr = font.element.get_or_add_rPr()
        r_fonts = r_pr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            r_fonts.set(qn(attr), font_name)

    def _apply_docx_paragraph_format(self, paragraph: Any) -> None:
        """Aplica o padrão de espaçamento solicitado para os parágrafos do Word."""
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 1.0
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)

    def _configure_docx_footer(self, document: Document) -> None:
        """Configura o rodapé com nome da pasta de trabalho e paginação dinâmica."""
        workbook_name = self.base_name
        for section in document.sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.clear()
            self._apply_docx_paragraph_format(paragraph)

            usable_width = section.page_width - section.left_margin - section.right_margin
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                usable_width,
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )
            self._apply_docx_top_border(paragraph)

            left_run = paragraph.add_run(f"Documentação Tableau pasta de Trabalho {workbook_name}")
            self._apply_docx_run_style(left_run, mono=False, bold=False, italic=False, font_size=8)

            tab_run = paragraph.add_run("\t")
            self._apply_docx_run_style(tab_run, mono=False, bold=False, italic=False, font_size=8)

            self._append_docx_field(paragraph, "PAGE", font_size=8)

            middle_run = paragraph.add_run(" de ")
            self._apply_docx_run_style(middle_run, mono=False, bold=False, italic=False, font_size=8)

            self._append_docx_field(paragraph, "NUMPAGES", font_size=8)

    def _configure_docx_header(self, document: Document) -> None:
        """Configura o cabeçalho com título centralizado e linha inferior e logos se configurados."""
        try:
            config_data = load_config()
            logo_empresa = config_data.get("logo_empresa")
            logo_cliente = config_data.get("logo_cliente")
        except Exception:
            logo_empresa = None
            logo_cliente = None

        for section in document.sections:
            header = section.header
            paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            paragraph.clear()
            self._apply_docx_paragraph_format(paragraph)
            
            paragraph.alignment = 0
            usable_width = section.page_width - section.left_margin - section.right_margin
            paragraph.paragraph_format.tab_stops.add_tab_stop(int(usable_width / 2), alignment=WD_TAB_ALIGNMENT.CENTER)
            paragraph.paragraph_format.tab_stops.add_tab_stop(int(usable_width), alignment=WD_TAB_ALIGNMENT.RIGHT)
            
            self._apply_docx_bottom_border(paragraph)

            if logo_cliente and Path(logo_cliente).is_file():
                run_cli = paragraph.add_run()
                run_cli.add_picture(logo_cliente, height=Cm(0.91))
            
            paragraph.add_run("\t")
            run = paragraph.add_run("Documentação Tableau")
            self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=8)
            
            paragraph.add_run("\t")
            if logo_empresa and Path(logo_empresa).is_file():
                run_emp = paragraph.add_run()
                run_emp.add_picture(logo_empresa, height=Cm(0.91))

    def _apply_docx_top_border(self, paragraph: Any) -> None:
        """Adiciona uma linha superior ao parágrafo, útil para o rodapé."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        top = p_bdr.find(qn("w:top"))
        if top is None:
            top = OxmlElement("w:top")
            p_bdr.append(top)
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")
        top.set(qn("w:space"), "1")
        top.set(qn("w:color"), "BFBFBF")

    def _apply_docx_bottom_border(self, paragraph: Any) -> None:
        """Adiciona uma linha inferior ao parágrafo, útil para o cabeçalho."""
        p_pr = paragraph._p.get_or_add_pPr()
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        bottom = p_bdr.find(qn("w:bottom"))
        if bottom is None:
            bottom = OxmlElement("w:bottom")
            p_bdr.append(bottom)
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "BFBFBF")

    def _append_docx_field(self, paragraph: Any, field_name: str, font_size: float) -> None:
        """Insere um campo dinâmico do Word, como PAGE ou NUMPAGES."""
        run = paragraph.add_run()
        self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=font_size)

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_name

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = "1"

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(placeholder)
        run._r.append(end)

    def _append_docx_toc(self, document: Document) -> None:
        """Insere um índice automático com os três níveis principais."""
        toc_title = document.add_paragraph()
        self._apply_docx_paragraph_format(toc_title)
        title_run = toc_title.add_run("Índice")
        self._apply_docx_run_style(title_run, mono=False, bold=True, italic=False, font_size=12)

        toc_paragraph = document.add_paragraph()
        self._apply_docx_paragraph_format(toc_paragraph)
        self._append_docx_complex_field(
            toc_paragraph,
            'TOC \\o "1-3" \\h \\z \\u',
            placeholder_text="Atualize o índice no Word se necessário.",
            font_size=10,
        )

    def _append_docx_complex_field(
        self,
        paragraph: Any,
        instruction: str,
        placeholder_text: str,
        font_size: float,
    ) -> None:
        """Insere um campo complexo do Word, como o TOC."""
        run = paragraph.add_run()
        self._apply_docx_run_style(run, mono=False, bold=False, italic=False, font_size=font_size)

        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction

        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")

        placeholder = OxmlElement("w:t")
        placeholder.text = placeholder_text

        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")

        run._r.append(begin)
        run._r.append(instr)
        run._r.append(separate)
        run._r.append(placeholder)
        run._r.append(end)

    def _docx_heading_style_for_block(self, style: str) -> str | None:
        """Mapeia os níveis principais do relatório para estilos reconhecidos pelo Word."""
        if style == "section":
            return "Heading 1"
        if style == "subsection":
            return "Heading 2"
        if style == "subsubsection":
            return "Heading 3"
        return None

    def _rtf_paragraph(self, text: str, style: str = "body", level: int = 0, mono: bool = False) -> str:
        escaped = self._rtf_escape(text)
        indent = 360 * max(level, 0)
        if style == "title":
            return rf"\pard\sa240\sb160\f0\fs32\b {escaped}\b0\par"
        if style == "subtitle":
            return rf"\pard\sa120\sb160\i\f0\fs20 {escaped}\i0\par"
        if style == "section":
            return rf"\pard\sa200\sb120\li{indent}\f0\fs28\b {escaped}\b0\par"
        if style == "subsection":
            return rf"\pard\sa160\sb80\li{indent}\f0\fs24\b {escaped}\b0\par"
        if style == "subsubsection":
            return rf"\pard\sa120\sb60\li{indent}\f0\fs22\b {escaped}\b0\par"
        if style == "body_bold":
            return rf"\pard\sa80\sb40\li{indent}\f0\fs21\b {escaped}\b0\par"
        font_id = 1 if mono else 0
        return rf"\pard\sa60\sb20\li{indent}\f{font_id}\fs20 {escaped}\par"

    def _rtf_bullet(self, text: str, level: int = 0, mono: bool = False) -> str:
        indent = 360 * max(level, 0)
        hanging = 180
        escaped = self._rtf_escape(text)
        font_id = 1 if mono else 0
        return rf"\pard\sa40\sb20\li{indent}\fi-{hanging}\f{font_id}\fs20 \'95\tab {escaped}\par"

    def _rtf_code_block(self, text: str, level: int = 0) -> str:
        indent = 360 * max(level, 0)
        escaped = self._rtf_escape(text)
        return rf"\pard\sa60\sb40\li{indent}\f1\fs20 {escaped}\par"

    def _rtf_list_block(
        self,
        label: str,
        values: list[Any],
        level: int = 0,
        empty_text: str = "-",
        mono: bool = False,
    ) -> list[str]:
        parts = [self._rtf_bullet(f"{label}:", level=level)]
        if values:
            for value in values:
                parts.append(self._rtf_bullet(str(value), level=level + 1, mono=mono))
        else:
            parts.append(self._rtf_bullet(empty_text, level=level + 1, mono=mono))
        return parts

    def _rtf_escape(self, text: Any) -> str:
        value = "" if text is None else str(text)
        output: list[str] = []
        for char in value:
            if char == "\\":
                output.append(r"\\")
            elif char == "{":
                output.append(r"\{")
            elif char == "}":
                output.append(r"\}")
            elif char == "\n":
                output.append(r"\line ")
            elif char == "\t":
                output.append(r"\tab ")
            else:
                codepoint = ord(char)
                if 32 <= codepoint <= 126:
                    output.append(char)
                else:
                    signed = codepoint if codepoint <= 32767 else codepoint - 65536
                    output.append(rf"\u{signed}?")
        return "".join(output)

    def _looks_like_hex_color(self, value: str | None) -> bool:
        if not value:
            return False
        return bool(re.fullmatch(r"#(?:[0-9a-fA-F]{3}|[0-9a-fA-F]{6}|[0-9a-fA-F]{8})", value.strip()))

    def _build_connection_diagram(
        self,
        connection: dict[str, Any],
        relationship_maps: list[dict[str, Any]] | None = None,
    ) -> str:
        """Monta um desenho textual simples de conexões e relacionamentos."""
        lines: list[str] = []
        attributes = connection.get("attributes") or {}
        connection_label = attributes.get("class") or "connection"
        lines.append(connection_label)

        named_connections = connection.get("named_connections") or []
        relations = connection.get("relations") or []
        relationship_maps = relationship_maps or []

        if named_connections:
            lines.append("|")
            lines.append("+-- named-connections")
            for named in named_connections:
                caption = named.get("caption") or named.get("name") or "(sem nome)"
                name = named.get("name")
                suffix = f" [{name}]" if name and name != caption else ""
                lines.append(f"|   +-- {caption}{suffix}")

        if relationship_maps:
            lines.append("|")
            lines.append("+-- relationships")
            for relation in relationship_maps:
                from_label = relation.get("from_label") or "(sem origem)"
                to_label = relation.get("to_label") or "(sem destino)"
                from_field = relation.get("from_field") or "-"
                to_field = relation.get("to_field") or "-"
                lines.append(f"|   +-- {from_label} -> {to_label}")
                lines.append(f"|       \\-- {from_field} = {to_field}")

        if relations:
            lines.append("|")
            lines.append("+-- relations")
            for relation in relations:
                relation_name = relation.get("name") or relation.get("table") or relation.get("type") or "(sem nome)"
                relation_type = relation.get("type") or "-"
                relation_connection = relation.get("connection")
                table_name = relation.get("table")

                detail_parts = [f"tipo={relation_type}"]
                if relation_connection:
                    detail_parts.append(f"connection={relation_connection}")
                if table_name and table_name != relation_name:
                    detail_parts.append(f"table={table_name}")

                lines.append(f"    +-- {relation_name}")
                lines.append(f"        \\-- {' | '.join(detail_parts)}")

        if len(lines) == 1:
            return ""
        return "\n".join(lines)

    def _build_datasources_markdown(self) -> list[str]:
        lines = ["", "## Fontes de Dados", ""]
        datasources = self.metadata["datasources"]
        if not datasources:
            lines.append("- Nenhuma fonte de dados identificada.")
            return lines

        for datasource in datasources:
            name = datasource.get("caption") or datasource.get("name") or "(sem nome)"
            datasource_type = self._infer_datasource_type(datasource)
            repo = datasource.get("repository_location") or {}
            is_published = bool(repo)
            custom_sql_relations = self._custom_sql_by_connection(datasource)
            external_custom_sql_relations = self._dedupe_external_custom_sql_relations(
                datasource.get("external_custom_sql_relations", [])
            )
            relationship_maps = self._dedupe_relationship_maps(
                (datasource.get("relationship_maps") or []) + (datasource.get("external_relationship_maps") or [])
            )
            lines.append(f"### {name}")
            lines.append("")
            lines.append(f"- Tipo: `{datasource_type}`")
            lines.append(f"- Versão: `{datasource.get('version')}`")
            lines.append(f"- Publicada: `{'sim' if is_published else 'não'}`")
            if datasource.get("connections"):
                for index, connection in enumerate(datasource.get("connections", []), start=1):
                    lines.append(f"- Conexão {index}: modo `{connection.get('mode') or '-'}`")
                    if connection.get("hyper_paths"):
                        lines.append(f"  - Arquivos .hyper associados: {', '.join(connection.get('hyper_paths', []))}")
            if is_published:
                lines.append(f"- Site: `{repo.get('site', '-')}`")
                lines.append(f"- Path: `{repo.get('path', '-')}`")
                lines.append(f"- ID: `{repo.get('id', '-')}`")
                lines.append(f"- Revision: `{repo.get('revision', '-')}`")
            lines.append(f"- Quantidade de campos mapeados: {len(datasource.get('columns', []))}")
            grouped_columns = self._group_datasource_columns(datasource)
            if grouped_columns:
                lines.append("- Campos da fonte de dados:")
                for table_name, columns in grouped_columns:
                    lines.append(f"  - {table_name}")
                    for column in columns:
                        lines.append(f"    - {self._display_column_label(column)}")
                        lines.append(f"      Datatype: `{column.get('datatype') or '-'}`")
                        lines.append(f"      Role: `{column.get('role') or '-'}`")
                        lines.append(f"      Em uso: `{format_yes_no(column.get('is_used'))}`")
                        lines.append(f"      Oculto: `{format_yes_no(column.get('hidden'))}`")
                        alias_rows = self._format_aliases(column.get('aliases', []))
                        if alias_rows:
                            lines.append("      Aliases:")
                            for alias in alias_rows:
                                lines.append(f"        - {alias}")
            if custom_sql_relations:
                lines.append("- SQL customizado encontrado no pacote: sim")
                for index, relation in enumerate(custom_sql_relations, start=1):
                    relation_name = relation.get("relation_name") or f"relation_{index}"
                    connection_index = relation.get("connection_index")
                    lines.append(f"- SQL customizado da conexão {connection_index}: `{relation_name}`")
                    lines.append("```sql")
                    lines.append(relation["custom_sql"])
                    lines.append("```")
            else:
                lines.append("- SQL customizado encontrado no pacote: não")
            if external_custom_sql_relations:
                lines.append("- SQL customizado encontrado em `.tdsx` externo: sim")
                for index, relation in enumerate(external_custom_sql_relations, start=1):
                    relation_name = relation.get("relation_name") or f"relation_{index}"
                    lines.append(
                        f"- SQL externo {index}: `{relation_name}` | origem: `{relation.get('tdsx_path')}`"
                    )
                    lines.append("```sql")
                    lines.append(relation["custom_sql"])
                    lines.append("```")
            elif datasource.get("external_tdsx_matches"):
                lines.append("- SQL customizado encontrado em `.tdsx` externo: não")
            if relationship_maps:
                lines.append("- Mapa de relacionamentos:")
                for relation in relationship_maps:
                    from_label = relation.get("from_label") or "(sem origem)"
                    to_label = relation.get("to_label") or "(sem destino)"
                    lines.append(f"  - {from_label} -> {to_label}")
                    lines.append(f"    Condição: {relation.get('expression') or '-'}")
                    if relation.get("from_field"):
                        lines.append(f"    Campo de ligação na origem: {relation.get('from_field')}")
                    if relation.get("to_field"):
                        lines.append(f"    Campo de ligação no destino: {relation.get('to_field')}")
                    relation_sources = relation.get("sources") or ([relation.get("source")] if relation.get("source") else [])
                    if relation_sources:
                        lines.append(f"    Origem: {', '.join(relation_sources)}")
                    if relation.get("from_attributes"):
                        lines.append(f"    Atributos de origem: {compact_json(relation.get('from_attributes'))}")
                    if relation.get("to_attributes"):
                        lines.append(f"    Atributos de destino: {compact_json(relation.get('to_attributes'))}")
            lines.append("")
        return lines

    def _build_preferences_markdown(self) -> list[str]:
        lines = ["", "## Preferências e Paletas", ""]
        palettes = self.metadata["preferences"]["color_palettes"]
        prefs = self.metadata["preferences"]["preferences"]

        if not prefs and not palettes:
            lines.append("- Nenhuma preferência global encontrada.")
            return lines

        for pref in prefs:
            lines.append(f"- Preferência `{pref.get('name')}` = `{pref.get('value')}`")

        if palettes:
            lines.append("")
            lines.append("### Paletas de Cor")
            lines.append("")
            for palette in palettes:
                colors = ", ".join(palette.get("colors", [])) or "-"
                lines.append(f"- `{palette.get('name')}` ({palette.get('type')}) -> {colors}")
        return lines

    def _build_parameters_markdown(self) -> list[str]:
        lines = ["", "## Parâmetros", ""]
        parameters = self.metadata["parameters"]
        if not parameters:
            lines.append("- Nenhum parâmetro encontrado.")
            return lines

        for parameter in parameters:
            name = parameter.get("caption") or clean_brackets(parameter.get("name")) or "(sem nome)"
            members = ", ".join(parameter.get("members", [])) or "-"
            lines.append(f"### {name}")
            lines.append("")
            lines.append(f"- Nome interno: `{parameter.get('name')}`")
            lines.append(f"- Datatype: `{parameter.get('datatype')}`")
            lines.append(f"- Type: `{parameter.get('type')}`")
            lines.append(f"- Role: `{parameter.get('role')}`")
            lines.append(f"- Domínio: `{parameter.get('param_domain_type')}`")
            lines.append(f"- Valor atual/default: `{parameter.get('value')}`")
            lines.append(f"- Membros: {members}")
            source_field = parameter.get("source_field_details")
            if source_field:
                if source_field.get("datasource") and source_field.get("field"):
                    lines.append(f"- Preenchido por campo da fonte: `{source_field['datasource']}` -> `{source_field['field']}`")
                else:
                    lines.append(f"- Preenchido por campo da fonte: `{source_field.get('raw')}`")
            else:
                lines.append("- Preenchido por campo da fonte: não")
            if parameter.get("used_in_dashboards"):
                lines.append("- Usado nos painéis:")
                for dashboard in sorted(parameter["used_in_dashboards"], key=str.lower):
                    lines.append(f"  - {dashboard}")
            else:
                lines.append("- Usado nos painéis: não está sendo usado em nenhum painel")
            if parameter.get("used_in_worksheets"):
                lines.append("- Usado nas planilhas:")
                for worksheet in sorted(parameter["used_in_worksheets"], key=str.lower):
                    lines.append(f"  - {worksheet}")
            else:
                lines.append("- Usado nas planilhas: não está sendo usado em nenhuma planilha")
            lines.append("")
        return lines

    def _build_calculations_markdown(self) -> list[str]:
        lines = ["## Campos Calculados", ""]
        calculations = self.metadata["calculations"]
        if not calculations:
            lines.append("- Nenhum campo calculado encontrado.")
            return lines

        grouped: dict[str, list[dict[str, Any]]] = {}
        for calculation in calculations:
            grouped.setdefault(calculation["datasource"] or "(sem datasource)", []).append(calculation)

        for datasource, items in grouped.items():
            lines.append(f"### {datasource}")
            lines.append("")
            for item in items:
                lines.append(f"- Campo: `{item.get('caption')}`")
                lines.append(f"- Nome interno: `{item.get('name')}`")
                lines.append(f"- Origem: `{item.get('origin')}`")
                lines.append(f"- Em uso: `{format_yes_no(item.get('is_used'))}`")
                lines.append(f"- Oculto: `{format_yes_no(item.get('hidden'))}`")
                alias_rows = self._format_aliases(item.get("aliases", []))
                if alias_rows:
                    lines.append("- Aliases:")
                    for alias in alias_rows:
                        lines.append(f"  - {alias}")
                codigo = item.get("codigo") or "-"
                lines.append(f"- Código: {codigo}")
                impacts = sorted(item.get("impacts", []), key=str.lower)
                if impacts:
                    lines.append("- Impacta / é referenciado por:")
                    for impact in impacts:
                        lines.append(f"  - {impact}")
                if item.get("used_in_dashboards"):
                    lines.append("- Usado nos painéis:")
                    for dashboard in sorted(item["used_in_dashboards"], key=str.lower):
                        lines.append(f"  - {dashboard}")
                if item.get("used_in_worksheets"):
                    lines.append("- Usado nas planilhas:")
                    for worksheet in sorted(item["used_in_worksheets"], key=str.lower):
                        lines.append(f"  - {worksheet}")
                lines.append("")
        return lines

    def _build_unused_objects_markdown(self) -> list[str]:
        lines = ["## Objetos não usados", ""]
        unused_objects = self._collect_unused_objects()

        lines.append("### Campos calculados")
        lines.append("")
        if unused_objects["calculations"]:
            for item in unused_objects["calculations"]:
                caption = self._display_object_name(item)
                lines.append(f"- {caption}")
                lines.append(f"  - Oculto: {format_yes_no(item.get('hidden'))}")
        else:
            lines.append("- Nenhum campo calculado não usado identificado.")

        lines.append("")
        lines.append("### Parâmetros")
        lines.append("")
        if unused_objects["parameters"]:
            for item in unused_objects["parameters"]:
                caption = self._display_object_name(item)
                lines.append(f"- {caption}")
        else:
            lines.append("- Nenhum parâmetro não usado identificado.")

        lines.append("")
        lines.append("### Fontes de dados")
        lines.append("")
        if unused_objects["datasources"]:
            for item in unused_objects["datasources"]:
                caption = self._display_object_name(item, allow_clean_brackets=False)
                lines.append(f"- {caption}")
        else:
            lines.append("- Nenhuma fonte de dados não usada identificada.")
        lines.append("")
        return lines

    def _build_dashboards_markdown(self) -> list[str]:
        lines = ["## Dashboards", ""]
        dashboards = self.metadata["dashboards"]
        if not dashboards:
            lines.append("- Nenhum dashboard encontrado.")
            return lines

        for dashboard in dashboards:
            lines.append(f"### {dashboard['name']}")
            lines.append("")
            if dashboard.get("worksheet_members"):
                lines.append("- Worksheets incluídas:")
                for member in sorted(dashboard["worksheet_members"], key=str.lower):
                    worksheet = next(
                        (item for item in self.metadata["worksheets"] if item.get("name") == member),
                        None,
                    )
                    datasources = worksheet.get("datasource_dependencies", []) if worksheet else []
                    lines.append(f"  - {member}")
                    if datasources:
                        friendly_datasources = sorted(
                            [self._humanize_datasource_reference(value) or value for value in datasources],
                            key=str.lower,
                        )
                        lines.append(f"    Fonte de dados: {', '.join(friendly_datasources)}")
                    else:
                        lines.append("    Fonte de dados: nenhuma identificada")
            else:
                lines.append("- Worksheets incluídas: nenhuma identificada")
            if dashboard.get("filters_used"):
                exposed_filters = [item for item in dashboard["filters_used"] if item.get("visibility") == "exposto no painel"]
                internal_filters = [item for item in dashboard["filters_used"] if item.get("visibility") == "interno da planilha"]
                if exposed_filters:
                    lines.append("- Filtros expostos no painel:")
                    for filter_item in exposed_filters:
                        context_label = "sim" if filter_item.get("is_context") else "não"
                        field_label = filter_item.get("field_label") or "(sem nome)"
                        filter_class = filter_item.get("class") or "-"
                        lines.append(f"  - {field_label} | tipo: {filter_class} | contexto: {context_label}")
                else:
                    lines.append("- Filtros expostos no painel: nenhum identificado")
                if internal_filters:
                    lines.append("- Filtros internos das planilhas:")
                    for filter_item in internal_filters:
                        context_label = "sim" if filter_item.get("is_context") else "não"
                        field_label = filter_item.get("field_label") or "(sem nome)"
                        filter_class = filter_item.get("class") or "-"
                        lines.append(f"  - {field_label} | tipo: {filter_class} | contexto: {context_label}")
                else:
                    lines.append("- Filtros internos das planilhas: nenhum identificado")
            else:
                lines.append("- Filtros expostos no painel: nenhum identificado")
                lines.append("- Filtros internos das planilhas: nenhum identificado")
            lines.append(f"- Zonas no layout: {len(dashboard.get('zones', []))}")
            lines.append(f"- Layouts de device: {len(dashboard.get('device_layouts', []))}")
            unique_colors = sorted(set(dashboard.get("colors_used", [])), key=str.lower)
            unique_fonts = self._display_fonts(dashboard.get("fonts_used", []))
            lines.append(f"- Cores usadas: {', '.join(unique_colors) or '-'}")
            lines.append(f"- Fontes usadas: {', '.join(unique_fonts) or '-'}")
            lines.append("")
        return lines

    def _build_visual_tokens_markdown(self) -> list[str]:
        lines = ["## Tokens Visuais", ""]
        colors = self.metadata["visual_tokens"]["colors"]
        fonts = self.metadata["visual_tokens"]["fonts"]
        unique_color_values = sorted({item["value"] for item in colors if item.get("value")}, key=str.lower)
        unique_font_values = self._display_fonts([item["value"] for item in fonts if item.get("value")])

        lines.append(f"- Total de cores únicas: {len(unique_color_values)}")
        lines.append(f"- Total de fontes únicas: {len(unique_font_values)}")
        lines.append("")
        lines.append("### Cores")
        lines.append("")
        if unique_color_values:
            for color in unique_color_values:
                lines.append(f"- `{color}`")
        else:
            lines.append("- Nenhuma cor identificada.")
        lines.append("")
        lines.append("### Fontes")
        lines.append("")
        for font in unique_font_values:
            lines.append(f"- `{font}`")
        lines.append("")
        return lines

    def _write_excel(self) -> Path:
        output_path = self.output_dir / f"{self.base_name}.xlsx"
        _log_progress(f"Gerando arquivo Excel `{output_path.name}`.")
        with pd.ExcelWriter(output_path) as writer:
            self._to_frame(self.metadata["parameters"]).to_excel(writer, sheet_name="Parameters", index=False)
            self._to_frame(self.metadata["calculations"]).to_excel(writer, sheet_name="Calculations", index=False)
            self._to_frame(self.metadata["datasources"], stringify_nested=True).to_excel(writer, sheet_name="Datasources", index=False)
            self._to_frame(self.metadata["worksheets"], stringify_nested=True).to_excel(writer, sheet_name="Worksheets", index=False)
            self._to_frame(self.metadata["dashboards"], stringify_nested=True).to_excel(writer, sheet_name="Dashboards", index=False)
            self._to_frame(self.metadata["visual_tokens"]["colors"]).to_excel(writer, sheet_name="Colors", index=False)
            self._to_frame(self.metadata["visual_tokens"]["fonts"]).to_excel(writer, sheet_name="Fonts", index=False)
            self._to_frame(self.package_manifest).to_excel(writer, sheet_name="PackageManifest", index=False)
        return output_path

    def _to_frame(self, records: list[dict[str, Any]], stringify_nested: bool = False) -> pd.DataFrame:
        if not records:
            return pd.DataFrame()
        df = pd.DataFrame(records)
        if stringify_nested:
            for column in df.columns:
                df[column] = df[column].apply(
                    lambda value: json.dumps(value, ensure_ascii=False, indent=2)
                    if isinstance(value, (dict, list))
                    else value
                )
        return df


def load_config(config_path: str | Path | None = None) -> dict[str, Any]:
    """Lê o arquivo de configuração padrão do projeto."""
    config_file = Path(config_path) if config_path is not None else PROJECT_ROOT / "config" / "config.json"
    if not config_file.exists():
        raise FileNotFoundError(
            f"Nenhum caminho informado e o arquivo '{config_file}' não foi encontrado."
        )

    try:
        config_data = json.loads(config_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"O arquivo '{config_file}' não contém um JSON válido.") from exc

    if not isinstance(config_data, dict):
        raise ValueError(f"O arquivo '{config_file}' deve conter um objeto JSON na raiz.")
    return config_data


def load_path_from_config(config_path: str | Path | None = None) -> str:
    """Lê o caminho padrão do arquivo Tableau a partir do config.json."""
    config_data = load_config(config_path)
    tableau_path = config_data.get("tableau_path")
    if not tableau_path:
        config_file = Path(config_path) if config_path is not None else PROJECT_ROOT / "config" / "config.json"
        raise ValueError(f"O arquivo '{config_file}' não possui a chave obrigatória 'tableau_path'.")
    if not isinstance(tableau_path, str):
        raise ValueError("A chave 'tableau_path' do config.json deve ser uma string.")
    if not tableau_path.strip():
        raise ValueError("A chave 'tableau_path' do config.json não pode estar vazia.")
    return tableau_path


def load_external_tdsx_paths_from_config(config_path: str | Path | None = None) -> list[str]:
    """Lê os caminhos opcionais de `.tdsx` externo definidos em config."""
    if not ENABLE_EXTERNAL_TDSX_LOOKUP:
        return []

    config_data = load_config(config_path)
    raw_paths = config_data.get("external_tdsx_paths", [])
    if raw_paths is None:
        return []
    if isinstance(raw_paths, str):
        raw_paths = [raw_paths]
    if not isinstance(raw_paths, list):
        raise ValueError(
            "A chave 'external_tdsx_paths' do config.json deve ser uma string ou uma lista de caminhos."
        )
    return [str(path) for path in raw_paths if str(path).strip()]


def parse_args() -> argparse.Namespace:
    """Define a interface de linha de comando."""
    parser = argparse.ArgumentParser(
        description="Gera documentação de workbooks Tableau (.twb/.twbx)."
    )
    parser.add_argument(
        "filepath",
        nargs="?",
        help="Caminho completo do arquivo Tableau (.twb ou .twbx).",
    )
    parser.add_argument(
        "--format",
        dest="output_format",
        choices=["all", "markdown", "json", "excel", "rtf", "docx"],
        default="all",
        help="Formato principal da documentação a gerar. Em `markdown`, o script gera `.md` e `.rtf`. O mapa XPath/JSON é sempre gerado; em `docx`, o relatório Word é exportado diretamente.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    try:
        file_path = args.filepath or load_path_from_config()
        external_tdsx_paths = load_external_tdsx_paths_from_config()
        documenter = TableauDoc(
            file_path,
            output_format=args.output_format,
            external_tdsx_paths=external_tdsx_paths,
        )
        written_files = documenter.write_outputs()

        print(f"Arquivo fonte: {documenter.source_path}")
        print(f"Diretório de saída: {documenter.output_dir}")
        print("Arquivos gerados:")
        for path in written_files:
            print(f"- {path}")
    except (FileNotFoundError, ValueError, ET.ParseError, zipfile.BadZipFile) as exc:
        print(f"Erro: {exc}")
        print(
            "Uso: python Tableau_doc.py <caminho_do_arquivo.twb|.twbx> "
            "[--format all|markdown|json|excel|rtf|docx]"
        )
        sys.exit(1)


if __name__ == "__main__":
    main()
